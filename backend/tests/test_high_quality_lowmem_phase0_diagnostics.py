import zipfile
from types import SimpleNamespace

from docx import Document

from app.core.recognizer_base import RecognizerResult
from app.core.pipeline_manager import PipelineManager
from app.recognizers.high_quality_lowmem_recognizer import HighQualityLowMemoryRecognizer
from app.rules.pipeline import RuleFirstPipeline
from app.rules.type_recognizers import TypeRuleRecognizers
from app.services.coverage_first.final_export import build_coverage_first_final_export_bundle
from app.services.default_numeric_masking import mask_default_numeric_text
from app.processors.document_exporter import DocumentExporter
from app.services.qwen_fragment_review_service import QwenFragmentReviewService
from app.services.risk_snippet_scheduler import RiskSnippetScheduler
from app.services.risk_snippet_scheduler import RiskSnippet
from app.services.recall_first_entity_merge_service import RecallFirstEntityMergeService
from app.services.lowmem_entity_utils import (
    build_recognition_view,
    docx_structure_unit_inventory,
    iter_docx_structure_units,
    resolve_docx_unit_spans,
    subject_noun_gate,
)
from app.processors.docx_xml_utils import extract_docx_visible_text_units
from app.services.contextual_desensitization_service import ContextualDesensitizationService
from app.workers.qwen_review_worker import (
    _apply_late_deterministic_adjudication_to_entities,
    _postprocess_final_entities,
    _validate_and_expand,
)


class _SyntheticUnitExtractor:
    def extract(self, text):
        target = "北京甲有限公司"
        start = text.index(target)
        return [
            RecognizerResult(
                entity_type="ORGANIZATION",
                start=start,
                end=start + len(target),
                score=0.91,
                text=target,
                source="synthetic_unit_extractor",
                metadata={},
            )
        ]


def _write_minimal_docx(path, document_xml, extra_parts=None):
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels)
        archive.writestr("word/document.xml", document_xml)
        for part_name, part_xml in dict(extra_parts or {}).items():
            archive.writestr(part_name, part_xml)


def test_docx_visible_text_backfills_complex_layout_nodes(tmp_path):
    docx_path = tmp_path / "complex_layout.docx"
    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>普通段落甲公司</w:t></w:r></w:p>
    <w:customLayout>
      <w:r><a:t>浮动布局乙公司</a:t></w:r>
    </w:customLayout>
  </w:body>
</w:document>"""
    _write_minimal_docx(docx_path, document_xml)

    text, metadata, units = extract_docx_visible_text_units(docx_path)

    assert "普通段落甲公司" in text
    assert "浮动布局乙公司" in text
    assert metadata["docx_coverage_backfill_unit_count"] == 1
    assert metadata["docx_unhandled_text_part_count"] == 0
    assert any("parser_coverage_backfill" in unit.flags for unit in units)


def test_docx_alternate_content_uses_single_effective_branch(tmp_path):
    docx_path = tmp_path / "alternate_content.docx"
    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
  <w:body>
    <mc:AlternateContent>
      <mc:Choice Requires="wps">
        <w:p><w:r><w:t>优先分支甲公司</w:t></w:r></w:p>
      </mc:Choice>
      <mc:Fallback>
        <w:p><w:r><w:t>回退分支乙公司</w:t></w:r></w:p>
      </mc:Fallback>
    </mc:AlternateContent>
  </w:body>
</w:document>"""
    _write_minimal_docx(docx_path, document_xml)

    text, metadata, units = extract_docx_visible_text_units(docx_path)

    assert "优先分支甲公司" in text
    assert "回退分支乙公司" not in text
    assert metadata["docx_uncovered_text_node_count"] == 0
    assert len([unit for unit in units if "分支" in unit.text]) == 1


def test_docx_visible_text_part_discovery_includes_non_whitelisted_word_xml(tmp_path):
    docx_path = tmp_path / "extra_visible_part.docx"
    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>正文甲公司</w:t></w:r></w:p>
  </w:body>
</w:document>"""
    extra_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:customTextStory xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:p><w:r><w:t>额外排版乙公司</w:t></w:r></w:p>
</w:customTextStory>"""
    _write_minimal_docx(
        docx_path,
        document_xml,
        extra_parts={"word/customLayout/story1.xml": extra_xml},
    )

    text, metadata, units = extract_docx_visible_text_units(docx_path)

    assert "正文甲公司" in text
    assert "额外排版乙公司" in text
    assert "word/customLayout/story1.xml" in metadata["docx_text_parts"]
    assert metadata["docx_unhandled_text_part_count"] == 0
    assert any(unit.part_name == "word/customLayout/story1.xml" for unit in units)


def test_recognition_view_metadata_reports_inline_space_removal():
    metadata = HighQualityLowMemoryRecognizer._build_recognition_view_metadata("甲 方在北京")

    assert metadata["recognition_view_original_length"] > metadata["recognition_view_sanitized_length"]
    assert metadata["recognition_view_removed_inline_space_count"] >= 1
    assert metadata["recognition_view_index_map_length"] == metadata["recognition_view_sanitized_length"]
    assert metadata["recognition_view_original_to_sanitized_length"] == metadata["recognition_view_original_length"]
    assert metadata["recognition_view_span_remap_fail_count"] == 0


def test_recognition_view_preserves_bidirectional_span_map():
    view = build_recognition_view("甲 方\n乙方")

    assert view.sanitized_text == "甲方\n乙方"
    assert view.sanitized_to_original == [0, 2, 3, 4, 5]
    assert view.original_to_sanitized[0] == 0
    assert view.original_to_sanitized[1] == -1
    assert view.original_to_sanitized[2] == 1


def test_docx_unit_diagnostic_metadata_counts_span_quality_without_reading_documents():
    recognizer = object.__new__(HighQualityLowMemoryRecognizer)
    text = "甲方北京甲公司\n乙方上海乙公司\n甲方北京甲公司"
    source_structure = {
        "docx_text_units": [
            {
                "text": "甲方北京甲公司",
                "start": 0,
                "end": 7,
                "container_type": "paragraph",
            },
            {
                "text": "乙方上海乙公司",
                "start": 99,
                "end": 106,
                "container_type": "table_cell",
            },
            {
                "text": "",
                "container_type": "header",
            },
            {
                "text": "不存在的单元",
                "container_type": "textbox",
            },
        ],
    }

    metadata = recognizer._build_docx_unit_diagnostic_metadata(text, source_structure)

    assert metadata["docx_unit_count"] == 4
    assert metadata["docx_unit_count_by_container"] == {
        "header": 1,
        "paragraph": 1,
        "table_cell": 1,
        "textbox": 1,
    }
    assert metadata["docx_unit_span_exact_count"] == 1
    assert metadata["docx_unit_span_mapped_count"] == 2
    assert metadata["docx_unit_span_mismatch_count"] == 2
    assert metadata["docx_unit_span_missing_count"] == 1
    assert metadata["docx_unit_span_duplicate_text_count"] == 1
    assert metadata["docx_unit_span_unresolved_count"] == 2


def test_docx_structure_unit_iterator_uses_raw_units_as_canonical_view():
    raw_unit = {
        "unit_id": "u1",
        "text": "北京甲有限公司",
        "start": 0,
        "end": 7,
        "container_type": "paragraph",
    }
    page_duplicate = dict(raw_unit)
    page_extra = {
        "unit_id": "page-only",
        "text": "上海乙有限公司",
        "start": 8,
        "end": 15,
        "container_type": "textbox",
    }
    source_structure = {
        "docx_text_units": [raw_unit],
        "pages": [{"units": [page_duplicate, page_extra]}],
    }

    units = list(iter_docx_structure_units(source_structure))
    inventory = docx_structure_unit_inventory(source_structure)

    assert [unit["unit_id"] for unit in units] == ["u1", "page-only"]
    assert inventory["raw_docx_text_unit_count"] == 1
    assert inventory["page_docx_unit_count"] == 2
    assert inventory["page_docx_unit_duplicate_raw_id_count"] == 1
    assert inventory["page_docx_unit_unique_extra_count"] == 1
    assert inventory["canonical_docx_structure_unit_count"] == 2


def test_docx_structure_unit_iterator_dedupes_page_view_without_unit_id_by_span_key():
    raw_unit = {
        "text": "北京甲有限公司",
        "start": 0,
        "end": 7,
        "container_type": "paragraph",
        "part_name": "word/document.xml",
    }
    page_duplicate = dict(raw_unit)
    repeated_other_place = {
        "text": "北京甲有限公司",
        "start": 8,
        "end": 15,
        "container_type": "paragraph",
        "part_name": "word/document.xml",
    }
    source_structure = {
        "docx_text_units": [raw_unit],
        "pages": [{"units": [page_duplicate, repeated_other_place]}],
    }

    units = list(iter_docx_structure_units(source_structure))
    inventory = docx_structure_unit_inventory(source_structure)

    assert [(unit.get("start"), unit.get("end")) for unit in units] == [(0, 7), (8, 15)]
    assert inventory["page_docx_unit_duplicate_raw_id_count"] == 1
    assert inventory["page_docx_unit_unique_extra_count"] == 1
    assert inventory["canonical_docx_structure_unit_count"] == 2


def test_docx_unit_span_resolution_uses_structural_order_for_duplicate_text():
    text = "名称\t北京甲公司\n名称\t北京甲公司"
    units = [
        {
            "text": "名称\t北京甲公司",
            "start": 99,
            "end": 108,
            "container_type": "table_cell",
            "order_index": 0,
        },
        {
            "text": "名称\t北京甲公司",
            "start": 99,
            "end": 108,
            "container_type": "table_cell",
            "order_index": 1,
        },
    ]

    resolved = resolve_docx_unit_spans(text, units)

    assert resolved[0]["_span_resolution"] == "ordered_forward"
    assert resolved[0]["_resolved_start"] == 0
    assert resolved[1]["_span_resolution"] == "ordered_forward"
    assert resolved[1]["_resolved_start"] == text.rindex("名称\t北京甲公司")


def test_docx_unit_span_resolution_uses_recognition_view_for_intrusive_spaces():
    text = "名称\t北 京甲公司"
    units = [
        {
            "text": "名称\t北京甲公司",
            "start": 99,
            "end": 107,
            "container_type": "table_cell",
        }
    ]

    resolved = resolve_docx_unit_spans(text, units)

    assert resolved[0]["_span_resolution"] == "sanitized_ordered_forward"
    assert text[resolved[0]["_resolved_start"] : resolved[0]["_resolved_end"]] == "名称\t北 京甲公司"


def test_rule_layer_recognizes_standard_company_full_name_with_intrusive_spaces():
    recognizer = TypeRuleRecognizers()
    text = "甲方：北 京 星河 科技 有限公司签署合同。"

    results = recognizer.recognize(text)

    assert any(
        item.entity_type == "ORGANIZATION"
        and item.text == "北 京 星河 科技 有限公司"
        and (item.metadata or {}).get("sanitized_rule_match")
        and (item.metadata or {}).get("normalized_text") == "北京星河科技有限公司"
        for item in results
    )


def test_rule_layer_recognizes_company_full_name_with_national_and_province_prefixes():
    recognizer = TypeRuleRecognizers()
    cases = [
        "甲方：中国星河科技有限公司签署合同。",
        "甲方：国家星河科技有限公司签署合同。",
        "甲方：广东省星河科技有限公司签署合同。",
        "甲方：内蒙古自治区星河科技有限公司签署合同。",
    ]

    for text in cases:
        results = recognizer.recognize(text)
        assert any(
            item.entity_type == "ORGANIZATION"
            and item.text.endswith("星河科技有限公司")
            and item.text in text
            and (item.metadata or {}).get("normalized_text") == item.text
            for item in results
        ), text


def test_rule_layer_recognizes_parenthesized_company_subject_without_parentheses():
    recognizer = TypeRuleRecognizers()
    text = "合同由（北京星河科技有限公司）负责履行。"

    results = recognizer.recognize(text)

    assert any(
        item.entity_type == "ORGANIZATION"
        and item.text == "北京星河科技有限公司"
        and text[item.start - 1] == "（"
        and text[item.end] == "）"
        for item in results
    )


def test_rule_pipeline_keeps_parenthesized_company_boundary_clean():
    pipeline = RuleFirstPipeline()
    text = "合同由（北京星河科技有限公司）负责履行。"

    result = pipeline.apply(text=text, results=[])
    org_texts = [item.text for item in result.results if item.entity_type == "ORGANIZATION"]

    assert "北京星河科技有限公司" in org_texts
    assert "（北京星河科技有限公司" not in org_texts


def test_rule_pipeline_recognizes_prefixed_company_full_names_without_loosening():
    pipeline = RuleFirstPipeline()
    cases = [
        "甲方：中国星河科技有限公司签署合同。",
        "甲方：国家星河科技有限公司签署合同。",
        "甲方：广东省星河科技有限公司签署合同。",
        "甲方：内蒙古自治区星河科技有限公司签署合同。",
    ]

    for text in cases:
        result = pipeline.apply(text=text, results=[])
        org_texts = [item.text for item in result.results if item.entity_type == "ORGANIZATION"]
        assert any(item.endswith("星河科技有限公司") for item in org_texts), text
        assert not any(item.startswith(("甲方", "：", "（", "(")) for item in org_texts), text


def test_rule_pipeline_splits_parallel_org_subjects_instead_of_merging():
    pipeline = RuleFirstPipeline()
    text = "星河和景岳共同负责履约。"

    result = pipeline.apply(text=text, results=[])
    org_texts = [item.text for item in result.results if item.entity_type == "ORGANIZATION"]

    assert "星河" in org_texts
    assert "景岳" in org_texts
    assert "星河和景岳" not in org_texts


def test_rule_pipeline_splits_person_org_bridge_subjects_without_merging():
    pipeline = RuleFirstPipeline()
    text = "张三担任北京星河科技有限公司法定代表人。"

    result = pipeline.apply(text=text, results=[])
    person_texts = [item.text for item in result.results if item.entity_type == "PERSON"]
    org_texts = [item.text for item in result.results if item.entity_type == "ORGANIZATION"]

    assert "张三" in person_texts
    assert "北京星河科技有限公司" in org_texts
    assert "张三担任北京星河科技有限公司" not in org_texts


def test_rule_pipeline_splits_person_org_bridge_subjects_for_short_company_suffixes():
    pipeline = RuleFirstPipeline()
    text = "张三担任乙有限公司法定代表人。张三是乙有限公司负责人。"

    result = pipeline.apply(text=text, results=[])
    person_texts = [item.text for item in result.results if item.entity_type == "PERSON"]
    org_texts = [item.text for item in result.results if item.entity_type == "ORGANIZATION"]

    assert "张三" in person_texts
    assert "乙有限公司" in org_texts
    assert "张三担任乙有限公司" not in org_texts
    assert "张三是乙有限公司" not in org_texts


def test_rule_pipeline_rejects_leading_bridge_verb_organization_pollution():
    pipeline = RuleFirstPipeline()
    text = "甲方是乙有限公司负责人。"

    result = pipeline.apply(text=text, results=[])
    org_texts = [item.text for item in result.results if item.entity_type == "ORGANIZATION"]

    assert "是乙有限公司" not in org_texts
    assert "甲方是乙有限公司" not in org_texts


def test_final_review_rejects_rule_organization_bridge_pollution():
    for value in ("是乙有限公司", "系乙有限公司", "担任乙有限公司", "张三是乙有限公司"):
        reason = QwenFragmentReviewService._deterministic_rule_organization_rejection_reason(
            {
                "text": value,
                "type": "ORGANIZATION",
                "source": "rule_organization",
                "metadata": {},
            }
        )
        assert reason == "deterministic_rule_org_left_context_pollution"


def test_final_review_rejects_bridge_pollution_from_any_organization_source():
    for source in ("qwen_fragment_review", "uie", "ner", "review_deterministic_decision"):
        for value in ("是乙有限公司", "张三是乙有限公司", "甲方是乙有限公司"):
            reason = QwenFragmentReviewService._deterministic_final_subject_rejection_reason(
                {
                    "text": value,
                    "type": "ORGANIZATION",
                    "source": source,
                    "metadata": {},
                }
            )
            assert reason == "deterministic_final_subject_left_context_pollution"


def test_subject_gate_rejects_strong_left_context_pollution_terms():
    for value in ("将乙有限公司", "把乙有限公司", "被乙有限公司", "担任乙有限公司", "任职乙有限公司", "是乙有限公司", "系乙有限公司", "为乙有限公司"):
        passed, reason = subject_noun_gate("ORGANIZATION", value, allow_short_org=True)
        assert not passed
        assert reason in {"leading_subject_linking_verb", "leading_function_prefix"}


def test_final_review_rejects_government_bridge_pollution_too():
    for value in ("是北京市朝阳区人民法院", "张三是北京市朝阳区人民法院", "担任北京市朝阳区人民法院"):
        reason = QwenFragmentReviewService._deterministic_final_subject_rejection_reason(
            {
                "text": value,
                "type": "GOVERNMENT",
                "source": "qwen_fragment_review",
                "metadata": {},
            }
        )
        assert reason == "deterministic_final_subject_left_context_pollution"


def test_rule_layer_label_value_does_not_emit_parenthesized_polluted_company_candidate():
    recognizer = TypeRuleRecognizers()
    text = "甲方：（广东省星河科技有限公司）签署合同。"

    results = [
        item
        for item in recognizer.recognize(text)
        if item.entity_type == "ORGANIZATION" and "星河科技有限公司" in item.text
    ]

    assert any(item.text == "广东省星河科技有限公司" for item in results)
    assert not any(item.text.startswith("（") for item in results)


def test_docx_structure_rule_pass_recognizes_spaced_company_inside_unit():
    recognizer = TypeRuleRecognizers()
    unit_text = "主体名称：北 京 星河 科技 有限公司"
    text = unit_text
    source_structure = {
        "docx_text_units": [
            {
                "unit_id": "u1",
                "text": unit_text,
                "start": 0,
                "end": len(unit_text),
                "container_type": "table_cell",
            }
        ]
    }

    results = recognizer.recognize(text, source_structure=source_structure)

    assert any(
        item.source == "rule_docx_structure"
        and item.entity_type == "ORGANIZATION"
        and item.text == "北 京 星河 科技 有限公司"
        and (item.metadata or {}).get("sanitized_rule_match")
        and (item.metadata or {}).get("normalized_text") == "北京星河科技有限公司"
        for item in results
    )


def test_docx_unit_model_pass_uses_resolved_spans_for_stale_unit_offsets():
    recognizer = HighQualityLowMemoryRecognizer()
    text = "前文\n甲方 北京甲有限公司\n后文"
    unit_text = "甲方 北京甲有限公司"
    source_structure = {
        "docx_text_units": [
            {
                "text": unit_text,
                "start": 99,
                "end": 99 + len(unit_text),
                "container_type": "paragraph",
                "unit_id": "u1",
            }
        ],
    }

    results = recognizer._extract_docx_unit_model_results(
        text=text,
        source_structure=source_structure,
        extractor=_SyntheticUnitExtractor(),
        source_name="docx_structure_uie",
    )
    metadata = recognizer._consume_docx_unit_model_metadata("docx_structure_uie")

    assert [item.text for item in results] == ["北京甲有限公司"]
    assert metadata["docx_structure_uie_unit_total_count"] == 1
    assert metadata["docx_structure_uie_unit_selected_count"] == 1
    assert metadata["docx_structure_uie_unit_skip_unresolved_span_count"] == 0
    assert metadata["docx_structure_uie_local_extract_call_count"] == 1


def test_docx_unit_model_pass_scans_paragraph_without_sensitive_cue():
    recognizer = object.__new__(HighQualityLowMemoryRecognizer)
    text = "普通段落北京甲有限公司继续履行。"
    source_structure = {
        "docx_text_units": [
            {
                "text": text,
                "start": 0,
                "end": len(text),
                "container_type": "paragraph",
                "unit_id": "p1",
            }
        ],
    }

    units = recognizer._select_docx_units_for_local_model_pass(
        text,
        source_structure,
        source_name="docx_structure_ner",
    )
    metadata = recognizer._consume_docx_unit_model_metadata("docx_structure_ner")

    assert len(units) == 1
    assert metadata["docx_structure_ner_unit_selected_count"] == 1
    assert metadata["docx_structure_ner_unit_skip_unsupported_container_count"] == 0


def test_docx_unit_model_pass_does_not_dedupe_repeated_text_at_different_resolved_spans():
    recognizer = object.__new__(HighQualityLowMemoryRecognizer)
    unit_text = "北京甲有限公司"
    text = f"{unit_text}\n{unit_text}"
    source_structure = {
        "docx_text_units": [
            {
                "text": unit_text,
                "start": 99,
                "end": 99 + len(unit_text),
                "container_type": "table_cell",
                "unit_id": "same-parser-id",
            },
            {
                "text": unit_text,
                "start": 99,
                "end": 99 + len(unit_text),
                "container_type": "table_cell",
                "unit_id": "same-parser-id",
            },
        ],
    }

    units = recognizer._select_docx_units_for_local_model_pass(
        text,
        source_structure,
        source_name="docx_structure_uie",
    )
    metadata = recognizer._consume_docx_unit_model_metadata("docx_structure_uie")

    assert len(units) == 2
    assert metadata["docx_structure_uie_unit_selected_count"] == 2
    assert metadata["docx_structure_uie_unit_skip_duplicate_count"] == 0


def test_docx_unit_model_pass_dedupes_page_view_duplicate_by_raw_unit_id():
    recognizer = object.__new__(HighQualityLowMemoryRecognizer)
    unit_text = "北京甲有限公司"
    text = unit_text
    raw_unit = {
        "text": unit_text,
        "start": 0,
        "end": len(unit_text),
        "container_type": "table_cell",
        "unit_id": "u1",
    }
    source_structure = {
        "docx_text_units": [raw_unit],
        "pages": [{"units": [dict(raw_unit)]}],
    }

    units = recognizer._select_docx_units_for_local_model_pass(
        text,
        source_structure,
        source_name="docx_structure_uie",
    )
    metadata = recognizer._consume_docx_unit_model_metadata("docx_structure_uie")

    assert len(units) == 1
    assert metadata["docx_structure_uie_unit_total_count"] == 1
    assert metadata["docx_structure_uie_unit_page_view_count"] == 1
    assert metadata["docx_structure_uie_unit_page_duplicate_raw_id_count"] == 1
    assert metadata["docx_structure_uie_unit_skip_duplicate_count"] == 0


def test_pipeline_subject_ledger_propagation_fills_unseen_repeated_occurrences():
    manager = object.__new__(PipelineManager)
    text = "甲方北京星河科技有限公司签订合同。北京星河科技有限公司负责履行。"
    first_start = text.index("北京星河科技有限公司")
    first_end = first_start + len("北京星河科技有限公司")
    second_start = text.rindex("北京星河科技有限公司")
    seed = RecognizerResult(
        entity_type="ORGANIZATION",
        start=first_start,
        end=first_end,
        score=0.91,
        text=text[first_start:first_end],
        source="high_quality_lowmem",
        metadata={
            "subject_ledger_status": "confirmed_subject",
            "subject_ledger_subject_status": "confirmed_subject",
            "subject_ledger_canonical_text": "北京星河科技有限公司",
            "subject_surfaces": ["北京星河科技有限公司"],
            "rule_first": {"action": "review", "risk_level": "medium"},
        },
    )

    expanded = manager._expand_repeated_mentions([seed], text)

    assert any(
        item.source == "propagate"
        and item.start == second_start
        and item.text == "北京星河科技有限公司"
        and (item.metadata or {}).get("propagated_from_subject_ledger")
        for item in expanded
    )


def test_rule_first_pipeline_emits_candidate_lifecycle_metadata():
    pipeline = RuleFirstPipeline()
    text = "甲方北京甲公司与乙方上海乙公司签订合同。"
    seed_results = [
        RecognizerResult(
            entity_type="ORGANIZATION",
            start=text.index("北京甲公司"),
            end=text.index("北京甲公司") + len("北京甲公司"),
            score=0.93,
            text="北京甲公司",
            source="synthetic_seed",
            metadata={"source_layer": "test_seed"},
        )
    ]

    result = pipeline.apply(text=text, results=seed_results, source_structure=None)
    metadata = result.metadata

    expected_keys = {
        "candidate_raw_count_by_source",
        "candidate_raw_count_by_type",
        "candidate_raw_count_by_source_layer",
        "rule_first_canonicalized_candidate_count",
        "rule_first_boundary_candidate_count",
        "candidate_after_boundary_count_by_source",
        "candidate_after_boundary_count_by_type",
        "candidate_after_boundary_count_by_source_layer",
        "candidate_after_gate_count_by_source",
        "candidate_after_gate_count_by_type",
        "candidate_after_gate_count_by_source_layer",
        "rule_first_gate_action_source_counts",
        "subject_admission_action_counts",
        "subject_admission_negative_reason_counts",
        "subject_admission_positive_reason_counts",
        "rule_first_boundary_repair_reason_counts",
        "rule_first_deduped_candidate_count",
        "rule_first_dedupe_removed_candidate_count",
        "subject_ledger_input_candidate_count",
        "candidate_ledger_total",
        "candidate_ledger_accepted",
        "candidate_ledger_review_only",
        "candidate_ledger_rejected",
        "candidate_ledger_rejected_total",
        "candidate_ledger_lost_before_subject_ledger",
        "candidate_ledger",
        "candidate_ledger_summary",
        "candidate_ledger_entry_count",
        "candidate_ledger_state_counts",
        "candidate_ledger_action_counts",
        "candidate_ledger_reason_counts",
    }

    assert expected_keys.issubset(metadata)
    assert metadata["rule_first_raw_candidate_count"] >= len(seed_results)
    assert metadata["rule_first_canonicalized_candidate_count"] <= metadata["rule_first_raw_candidate_count"]
    assert metadata["candidate_ledger_total"] == metadata["rule_first_boundary_candidate_count"]
    assert metadata["subject_ledger_input_candidate_count"] == metadata["rule_first_deduped_candidate_count"]
    assert metadata["candidate_ledger_entry_count"] >= metadata["candidate_ledger_total"]
    assert metadata["candidate_ledger_state_counts"]["raw"] == metadata["rule_first_raw_candidate_count"]
    assert metadata["candidate_ledger_state_counts"]["boundary_normalized"] == metadata["rule_first_boundary_candidate_count"]
    assert (
        metadata["candidate_ledger_state_counts"].get("subject_ledger_input", 0)
        == metadata["subject_ledger_input_candidate_count"]
    )
    assert metadata["subject_admission_action_counts"]
    assert any(
        isinstance((item.metadata or {}).get("subject_admission"), dict)
        for item in result.results
    )
    assert result.results

    ledger = metadata["candidate_ledger"]
    assert ledger["version"] == "candidate_ledger.v1"
    assert ledger["entries"]
    assert all("text" not in entry for entry in ledger["entries"])
    assert all("text_hash" in entry for entry in ledger["entries"])


def test_boundary_repair_risk_routes_through_subject_admission_metadata():
    pipeline = RuleFirstPipeline()
    text = "通过公司提交材料"
    seed_results = [
        RecognizerResult(
            entity_type="ORGANIZATION",
            start=0,
            end=len(text),
            score=0.81,
            text=text,
            source="synthetic_seed",
            metadata={"source_layer": "test_seed"},
        )
    ]

    result = pipeline.apply(text=text, results=seed_results, source_structure=None)
    all_candidates = [*result.results, *result.rejected_results]

    assert all_candidates
    assert all(not (item.metadata or {}).get("boundary_repair_rejected") for item in all_candidates)
    assert any(
        "subject_admission" in (item.metadata or {})
        for item in all_candidates
    )
    assert result.metadata["subject_admission_negative_reason_counts"]


def test_boundary_repair_strips_function_prefix_without_dropping_short_subject():
    pipeline = RuleFirstPipeline()
    text = "通过星河提交材料"
    seed_results = [
        RecognizerResult(
            entity_type="ORGANIZATION",
            start=0,
            end=len("通过星河"),
            score=0.81,
            text="通过星河",
            source="synthetic_seed",
            metadata={"source_layer": "test_seed"},
        )
    ]

    result = pipeline.apply(text=text, results=seed_results, source_structure=None)
    all_candidates = [*result.results, *result.rejected_results]

    assert any(item.text == "星河" for item in all_candidates)
    assert all(item.text != "通过星河" for item in result.results)
    repair_counts = result.metadata["rule_first_boundary_repair_reason_counts"]
    assert repair_counts.get("strip_leading_function_words", 0) + repair_counts.get("strip_leading_noise", 0) >= 1
    assert result.metadata["subject_admission_action_counts"]


def test_boundary_repair_suffix_anchor_does_not_hard_drop_on_admission_shape():
    pipeline = RuleFirstPipeline()
    text = "通知清theta公司员工提交材料"
    seed_results = [
        RecognizerResult(
            entity_type="ORGANIZATION",
            start=0,
            end=len("通知清theta公司"),
            score=0.81,
            text="通知清theta公司",
            source="synthetic_seed",
            metadata={"source_layer": "test_seed"},
        )
    ]

    result = pipeline.apply(text=text, results=seed_results, source_structure=None)
    all_candidates = [*result.results, *result.rejected_results]

    assert any(item.text == "清theta公司" for item in all_candidates)
    assert all(not (item.metadata or {}).get("boundary_repair_rejected") for item in all_candidates)
    repair_counts = result.metadata["rule_first_boundary_repair_reason_counts"]
    assert repair_counts.get("strip_leading_function_words", 0) + repair_counts.get("strip_leading_noise", 0) >= 1
    assert result.metadata["candidate_ledger_total"] >= 1


def test_short_org_recall_preserves_company_suffix_before_functional_tail():
    recognizer = TypeRuleRecognizers()
    text = "通知清theta公司提交材料"

    results = recognizer.recognize(text)

    assert any(
        item.entity_type == "ORGANIZATION"
        and item.text == "清theta公司"
        and item.source == "rule_organization_context"
        and (item.metadata or {}).get("trigger") == "right_boundary_short_org"
        for item in results
    )
    assert all(item.text != "通知清theta公司" for item in results)
    assert all(item.text != "清theta公司提交材料" for item in results)


def test_short_org_recall_keeps_company_suffix_before_role_tail():
    recognizer = TypeRuleRecognizers()
    text = "通过清theta公司员工提交材料"

    results = recognizer.recognize(text)

    assert any(
        item.entity_type == "ORGANIZATION"
        and item.text == "清theta公司"
        and (item.metadata or {}).get("right_boundary_cue") == "role_tail"
        for item in results
    )
    assert all("员工" not in item.text for item in results if item.entity_type == "ORGANIZATION")


def test_qwen_coverage_discovery_snippet_targets_uncovered_structure_unit():
    scheduler = RiskSnippetScheduler()
    text = "普通说明\n乙方星河公司负责付款"
    source_structure = {
        "docx_text_units": [
            {
                "unit_id": "u1",
                "text": "普通说明",
                "start": 0,
                "end": 4,
                "container_type": "paragraph",
                "order_index": 0,
            },
            {
                "unit_id": "u2",
                "text": "乙方星河公司负责付款",
                "start": 5,
                "end": len(text),
                "container_type": "paragraph",
                "order_index": 1,
            },
        ]
    }

    snippets = scheduler.build_snippets(
        text,
        [],
        source_structure=source_structure,
        max_snippets=20,
        max_chars_per_snippet=80,
    )

    discovery = [
        snippet
        for snippet in snippets
        if snippet.snippet_type == "qwen_coverage_discovery"
    ]
    assert len(discovery) == 1
    assert discovery[0].risk_reason == "qwen_discovery:uncovered_paragraph"
    assert discovery[0].target_entity["metadata"]["docx_unit_id"] == "u2"


def test_qwen_coverage_discovery_reviews_signal_units_even_when_partly_covered():
    scheduler = RiskSnippetScheduler()
    text = "乙方星河公司负责付款"
    source_structure = {
        "docx_text_units": [
            {
                "unit_id": "u1",
                "text": text,
                "start": 0,
                "end": len(text),
                "container_type": "paragraph",
            }
        ]
    }
    covered = [
        RecognizerResult(
            entity_type="ORGANIZATION",
            start=2,
            end=6,
            score=0.9,
            text="星河公司",
            source="synthetic_seed",
            metadata={},
        )
    ]

    snippets = scheduler.build_snippets(
        text,
        covered,
        source_structure=source_structure,
        max_snippets=20,
        max_chars_per_snippet=80,
    )

    discovery = [snippet for snippet in snippets if snippet.snippet_type == "qwen_coverage_discovery"]
    assert len(discovery) == 1
    assert discovery[0].target_entity["metadata"]["docx_unit_id"] == "u1"
    assert scheduler.last_metadata["coverage_discovery_signal_unit_count"] == 1
    assert scheduler.last_metadata["coverage_discovery_partial_unit_count"] == 1


def test_qwen_coverage_discovery_keeps_multi_subject_unit_with_partial_coverage():
    scheduler = RiskSnippetScheduler()
    text = "甲方星河公司与乙方景岳公司共同履行"
    source_structure = {
        "docx_text_units": [
            {
                "unit_id": "u1",
                "text": text,
                "start": 0,
                "end": len(text),
                "container_type": "paragraph",
            }
        ]
    }
    covered = [
        RecognizerResult(
            entity_type="ORGANIZATION",
            start=text.index("星河公司"),
            end=text.index("星河公司") + len("星河公司"),
            score=0.9,
            text="星河公司",
            source="synthetic_seed",
            metadata={},
        )
    ]

    snippets = scheduler.build_snippets(
        text,
        covered,
        source_structure=source_structure,
        max_snippets=20,
        max_chars_per_snippet=80,
    )

    discovery = [item for item in snippets if item.snippet_type == "qwen_coverage_discovery"]
    assert len(discovery) == 1
    assert discovery[0].target_entity["metadata"]["docx_unit_id"] == "u1"


def test_qwen_review_scheduler_keeps_coverage_discovery_snippets():
    snippets = [
        RiskSnippet("narrative_hotspot", "ordinary", 0, 10, "普通片段"),
        RiskSnippet("qwen_coverage_discovery", "qwen_discovery:uncovered_paragraph", 10, 20, "乙方星河公司负责付款"),
    ]

    scheduled = QwenFragmentReviewService._schedule_review_snippets(snippets, review_limit=1)

    assert any(snippet.snippet_type == "qwen_coverage_discovery" for snippet in scheduled)


def test_risk_snippet_scheduler_does_not_drop_discovery_behind_structure_cap():
    snippets = [
        RiskSnippet(
            "docx_table_cell_block",
            "docx_structure:table_cell",
            index * 10,
            index * 10 + 6,
            f"结构{index}公司",
        )
        for index in range(30)
    ]
    snippets.append(
        RiskSnippet(
            "qwen_coverage_discovery",
            "qwen_discovery:uncovered_table_cell",
            500,
            520,
            "乙方星河公司负责付款",
            target_entity={
                "type": "DISCOVERY",
                "start": 500,
                "end": 520,
                "source": "coverage_discovery",
                "metadata": {
                    "qwen_coverage_discovery": True,
                    "docx_unit_id": "u-discovery",
                    "docx_container_type": "table_cell",
                },
            },
        )
    )

    deduped = RiskSnippetScheduler._dedupe(snippets, max_count=12)

    assert any(snippet.snippet_type == "qwen_coverage_discovery" for snippet in deduped)


def test_standard_qwen_review_does_not_materialize_entities_from_model_payload(monkeypatch):
    service = QwenFragmentReviewService()
    text = "通过公司材料完成审核，乙方北京星河科技有限公司负责履行。"
    snippet = RiskSnippet("narrative_hotspot", "ordinary", 0, len(text), text)

    async def fake_review_with_runtime(*args, **kwargs):
        return '{"entities":[{"type":"ORGANIZATION","text":"公司材料完成"},{"type":"ORGANIZATION","text":"北京星河科技有限公司"}],"rejects":[],"entity_decisions":[]}'

    runtime = SimpleNamespace(model_id="test-review", backend="test", fallback=False, asset=None)
    monkeypatch.setattr(service, "_select_review_runtime", lambda: runtime)
    monkeypatch.setattr(service, "_review_thinking_enabled", lambda runtime: False)
    monkeypatch.setattr(service, "_review_with_runtime", fake_review_with_runtime)

    import asyncio

    result = asyncio.run(service.review(text, [snippet], existing_entities=[], max_snippets=1))

    assert result.raw_candidate_count == 2
    assert result.entities == []


def test_coverage_qwen_review_still_materializes_discovery_entities(monkeypatch):
    service = QwenFragmentReviewService()
    text = "乙方北京星河科技有限公司负责履行。"
    snippet = RiskSnippet(
        "qwen_coverage_discovery",
        "qwen_discovery:uncovered_paragraph",
        0,
        len(text),
        text,
    )

    async def fake_review_with_runtime(*args, **kwargs):
        return '{"entities":[{"type":"ORGANIZATION","text":"北京星河科技有限公司"}],"rejects":[],"entity_decisions":[]}'

    runtime = SimpleNamespace(model_id="test-review", backend="test", fallback=False, asset=None)
    monkeypatch.setattr(service, "_select_review_runtime", lambda: runtime)
    monkeypatch.setattr(service, "_review_thinking_enabled", lambda runtime: False)
    monkeypatch.setattr(service, "_review_with_runtime", fake_review_with_runtime)

    import asyncio

    result = asyncio.run(service.review(text, [snippet], existing_entities=[], max_snippets=1))

    assert result.raw_candidate_count == 1
    assert [item.text for item in result.entities] == ["北京星河科技有限公司"]


def test_coverage_qwen_review_uses_dedicated_discovery_prompt():
    service = QwenFragmentReviewService()
    snippet = RiskSnippet(
        "qwen_coverage_discovery",
        "qwen_discovery:uncovered_table_cell",
        0,
        18,
        "合同由（北京星河科技有限公司）负责履行",
        target_entity={
            "type": "DISCOVERY",
            "start": 0,
            "end": 18,
            "source": "coverage_discovery",
            "metadata": {
                "qwen_coverage_discovery": True,
                "docx_unit_id": "u1",
                "docx_container_type": "table_cell",
            },
        },
    )

    prompt = service._build_review_prompt_for_snippet(snippet, existing_entities=[])

    assert "最终查漏模型" in prompt
    assert "重点查漏" in prompt
    assert "括号内主体" in prompt
    assert "已识别实体" not in prompt


def test_missing_candidate_review_only_materializes_target_related_entities(monkeypatch):
    service = QwenFragmentReviewService()
    text = "通过公司材料完成审核，乙方北京星河科技有限公司负责履行。"
    target_start = text.index("北京星河科技有限公司")
    snippet = RiskSnippet(
        "missing_candidate_review",
        "rule_first:review_only_rejected",
        0,
        len(text),
        text,
        target_entity={
            "type": "ORGANIZATION",
            "text": "北京星河科技有限公司",
            "start": target_start,
            "end": target_start + len("北京星河科技有限公司"),
            "source": "rule_first_review_only",
            "metadata": {"missing_candidate_review": True},
        },
    )

    async def fake_review_with_runtime(*args, **kwargs):
        return '{"entities":[{"type":"ORGANIZATION","text":"公司材料完成"},{"type":"ORGANIZATION","text":"北京星河科技有限公司"}],"rejects":[],"entity_decisions":[]}'

    runtime = SimpleNamespace(model_id="test-review", backend="test", fallback=False, asset=None)
    monkeypatch.setattr(service, "_select_review_runtime", lambda: runtime)
    monkeypatch.setattr(service, "_review_thinking_enabled", lambda runtime: False)
    monkeypatch.setattr(service, "_review_with_runtime", fake_review_with_runtime)

    import asyncio

    result = asyncio.run(service.review(text, [snippet], existing_entities=[], max_snippets=1))

    assert result.raw_candidate_count == 2
    assert [item.text for item in result.entities] == ["北京星河科技有限公司"]
    assert all(item.text != "公司材料完成" for item in result.entities)
    assert result.metadata["missing_candidate_materialized_entity_count"] == 1


def test_weak_organization_reference_hits_are_collected_when_supported_subject_exists():
    service = ContextualDesensitizationService()
    text = "北京星河科技有限公司签署合同。该公司提交材料。"
    entity_memory = {
        "org:star": {
            "primary_type": "ORGANIZATION",
            "primary_text": "北京星河科技有限公司",
            "canonical_role": "ORGANIZATION",
            "variants": ["该公司"],
        }
    }
    hits = service._collect_weak_organization_reference_hits(
        text=text,
        entities=[
            {
                "type": "ORGANIZATION",
                "text": "北京星河科技有限公司",
                "start": 0,
                "end": len("北京星河科技有限公司"),
                "metadata": {"resolved_group_key": "org:star"},
            }
        ],
        entity_memory=entity_memory,
        covered=set(),
        covered_ranges=[],
    )

    assert hits
    assert any(hit["canonical_key"] == "org:star" for hit in hits)


def test_semantic_recovery_selection_protects_coverage_discovery_from_rule_first_budget():
    recognizer = HighQualityLowMemoryRecognizer()
    snippets = [
        RiskSnippet(
            "rule_first_review_block",
            "rule_first:review",
            index * 10,
            index * 10 + 6,
            f"主体{index}公司",
        )
        for index in range(40)
    ]
    snippets.append(
        RiskSnippet(
            "qwen_coverage_discovery",
            "qwen_discovery:uncovered_table_cell",
            1000,
            1010,
            "乙方星河公司负责付款",
        )
    )

    selected = recognizer._select_review_snippets_requiring_semantic_recovery(snippets, [])

    assert any(snippet.snippet_type == "qwen_coverage_discovery" for snippet in selected)


def test_qwen_discovery_materialization_marks_entities_and_counts_rejects():
    service = QwenFragmentReviewService()
    text = "乙方星河公司负责付款"
    snippet = RiskSnippet(
        "qwen_coverage_discovery",
        "qwen_discovery:uncovered_paragraph",
        0,
        len(text),
        text,
        target_entity={
            "type": "DISCOVERY",
            "start": 0,
            "end": len(text),
            "source": "coverage_discovery",
            "metadata": {
                "docx_unit_id": "u1",
                "docx_container_type": "paragraph",
                "span_resolution": "exact",
            },
        },
    )

    results = service._materialize_candidates(
        text,
        {
            "entities": [
                {"type": "ORGANIZATION", "text": "星河公司"},
                {"type": "ORGANIZATION", "text": "不存在公司"},
                {"type": "ORGANIZATION", "text": "通过"},
            ]
        },
        snippet,
    )

    assert len(results) == 1
    assert results[0].text == "星河公司"
    assert (results[0].metadata or {}).get("qwen_coverage_discovery") is True
    assert (results[0].metadata or {}).get("qwen_discovery_target") == {
        "docx_unit_id": "u1",
        "docx_container_type": "paragraph",
        "span_resolution": "exact",
    }
    assert service._last_materialize_span_miss_count == 1
    assert service._last_materialize_gate_reject_count == 1


def test_qwen_discovery_materialization_does_not_pollute_parenthesized_subject_left_boundary():
    service = QwenFragmentReviewService()
    text = "合同由（北京星河科技有限公司）负责履行"
    snippet = RiskSnippet(
        "qwen_coverage_discovery",
        "qwen_discovery:uncovered_paragraph",
        0,
        len(text),
        text,
    )

    results = service._materialize_candidates(
        text,
        {"entities": [{"type": "ORGANIZATION", "text": "北京星河科技有限公司"}]},
        snippet,
    )

    assert len(results) == 1
    assert results[0].text == "北京星河科技有限公司"
    assert text[results[0].start - 1] == "（"
    assert text[results[0].end] == "）"


def test_qwen_discovery_materialization_preserves_national_and_province_company_prefixes():
    service = QwenFragmentReviewService()
    cases = [
        "甲方：中国星河科技有限公司签署合同",
        "甲方：国家星河科技有限公司签署合同",
        "甲方：广东省星河科技有限公司签署合同",
        "甲方：内蒙古自治区星河科技有限公司签署合同",
    ]
    for text in cases:
        expected = text.split("：", 1)[1].split("签署", 1)[0]
        snippet = RiskSnippet(
            "qwen_coverage_discovery",
            "qwen_discovery:uncovered_paragraph",
            0,
            len(text),
            text,
        )

        results = service._materialize_candidates(
            text,
            {"entities": [{"type": "ORGANIZATION", "text": expected}]},
            snippet,
        )

        assert len(results) == 1
        assert results[0].text == expected


def test_final_export_keeps_qwen_discovery_entities_in_directory_and_mapping():
    text = "乙方星河公司负责付款"
    entities = [
        {
            "type": "ORGANIZATION",
            "text": "星河公司",
            "start": text.index("星河公司"),
            "end": text.index("星河公司") + len("星河公司"),
            "source": "qwen_fragment_review",
            "replacement": "公司1",
            "replacement_method": "contextual",
            "metadata": {
                "qwen_coverage_discovery": True,
                "source_layer": "llm_review",
                "subject_ledger_subject_id": "S1",
            },
        }
    ]

    bundle = build_coverage_first_final_export_bundle(entities=entities, source_text=text)

    assert bundle["enabled"] is True
    assert bundle["summary"]["final_desensitized_entity_input_count"] == 1
    assert bundle["summary"]["final_directory_occurrence_count"] == 1
    assert bundle["summary"]["final_mapping_entity_count"] == 1
    assert bundle["summary"]["qwen_discovery_desensitized_entity_count"] == 1
    assert bundle["summary"]["qwen_discovery_directory_occurrence_count"] == 1
    assert bundle["summary"]["qwen_discovery_mapping_entity_count"] == 1
    assert bundle["directory_rows"][0]["qwen_coverage_discovery"] is True
    assert bundle["mapping_entities"][0]["text"] == "星河公司"
    assert bundle["mapping_entities"][0]["metadata"]["qwen_coverage_discovery"] is True


def test_qwen_discovery_projection_counts_surviving_public_subjects():
    discovery = RecognizerResult(
        entity_type="COMPANY_NAME",
        start=2,
        end=6,
        score=0.9,
        text="星河公司",
        source="qwen_fragment_review",
        metadata={"qwen_coverage_discovery": True, "source_layer": "llm_review"},
    )

    projected = HighQualityLowMemoryRecognizer._project_default_public_results([discovery])
    counts = HighQualityLowMemoryRecognizer._qwen_discovery_projection_stage_counts(
        projected,
        prefix="qwen_discovery_projection_output",
    )

    assert len(projected) == 1
    assert projected[0].entity_type == "ORGANIZATION"
    assert (projected[0].metadata or {}).get("qwen_coverage_discovery") is True
    assert counts["qwen_discovery_projection_output_count"] == 1
    assert counts["qwen_discovery_projection_output_subject_count"] == 1


def test_review_worker_projection_happens_before_public_entity_filter():
    text = "乙方星河公司负责付款"
    discovery = RecognizerResult(
        entity_type="COMPANY_NAME",
        start=text.index("星河公司"),
        end=text.index("星河公司") + len("星河公司"),
        score=0.9,
        text="星河公司",
        source="qwen_fragment_review",
        metadata={"qwen_coverage_discovery": True, "source_layer": "llm_review"},
    )
    recognizer = HighQualityLowMemoryRecognizer()

    projected = recognizer._project_default_public_results([discovery])
    filtered = recognizer._filter_entities(projected, ["ORGANIZATION"])
    final_results = _validate_and_expand(text, filtered)

    assert len(final_results) == 1
    assert final_results[0].entity_type == "ORGANIZATION"
    assert (final_results[0].metadata or {}).get("qwen_coverage_discovery") is True


def test_prune_keeps_ambiguous_structure_short_org_with_review_marker():
    service = ContextualDesensitizationService()
    entities = [
        {
            "type": "ORGANIZATION",
            "text": "星河",
            "start": 2,
            "end": 4,
            "source": "rule_organization_context",
            "metadata": {
                "short_org_candidate": True,
                "source_layer": "structure",
                "trigger": "right_boundary_short_org",
                "subject_ledger_status": "ambiguous_short_subject",
            },
        }
    ]

    pruned = service._prune_invalid_entities(entities)

    assert len(pruned) == 1
    assert pruned[0]["text"] == "星河"
    assert pruned[0]["metadata"]["short_org_publication_review_required"] is True
    assert pruned[0]["metadata"]["requires_manual_review"] is True


def test_postprocess_filters_ambiguous_structure_short_org_from_public_entities():
    service = ContextualDesensitizationService()
    entities = [
        {
            "type": "ORGANIZATION",
            "text": "星河",
            "start": 2,
            "end": 4,
            "source": "rule_organization_context",
            "metadata": {
                "short_org_candidate": True,
                "source_layer": "structure",
                "trigger": "right_boundary_short_org",
                "subject_ledger_status": "ambiguous_short_subject",
                "short_org_publication_review_required": True,
                "requires_manual_review": True,
            },
        },
        {
            "type": "ORGANIZATION",
            "text": "北京星河科技有限公司",
            "start": 10,
            "end": 20,
            "source": "rule_organization",
            "metadata": {},
        },
    ]

    processed = _postprocess_final_entities(service, "甲方星河乙方北京星河科技有限公司", entities)

    assert [item["text"] for item in processed] == ["北京星河科技有限公司"]


def test_mapping_directory_prefers_final_mapping_layer_over_blocked_rewrite_subset():
    exporter = DocumentExporter()
    coverage_first_final_export = {
        "mapping_entities": [
            {
                "type": "ORGANIZATION",
                "text": "星河公司",
                "replacement": "公司1",
                "start": 2,
                "end": 6,
            },
            {
                "type": "PERSON",
                "text": "张三",
                "replacement": "人员1",
                "start": 10,
                "end": 12,
            },
        ],
        "rewrite_entries": [
            {
                "entity_type": "ORGANIZATION",
                "source_text": "星河公司",
                "replacement": "公司1",
                "start": 2,
                "end": 6,
                "verification_status": "pending_prewrite",
            },
            {
                "entity_type": "PERSON",
                "source_text": "张三",
                "replacement": "人员1",
                "start": 10,
                "end": 12,
                "verification_status": "blocked",
            },
        ],
    }

    mapping_entities = exporter._build_mapping_export_entities(
        entities=[],
        coverage_first_final_export=coverage_first_final_export,
    )

    assert [item["text"] for item in mapping_entities] == ["星河公司", "张三"]


def test_docx_export_applies_coverage_first_final_range_rewrite(tmp_path, monkeypatch):
    source_path = tmp_path / "source.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("甲方：")
    paragraph.add_run("北京 星河")
    paragraph.add_run("科技有限公司")
    paragraph.add_run("签署合同。")
    document.save(source_path)

    source_text = "甲方：北京星河科技有限公司签署合同。"
    subject = "北京星河科技有限公司"
    start = source_text.index(subject)
    entities = [
        {
            "type": "ORGANIZATION",
            "text": subject,
            "start": start,
            "end": start + len(subject),
            "replacement": "组织机构甲",
            "replacement_method": "contextual",
        }
    ]
    bundle = build_coverage_first_final_export_bundle(entities=entities, source_text=source_text)
    monkeypatch.setattr("app.processors.document_exporter.settings.OUTPUT_DIR", str(tmp_path))

    result = DocumentExporter().export(
        task_id="coverage-final-docx",
        source_path=str(source_path),
        original_filename="source.docx",
        source_text=source_text,
        source_metadata={},
        source_structure={},
        entities=entities,
        anonymized_text="甲方：组织机构甲签署合同。",
        coverage_first_final_export=bundle,
    )

    exported = Document(str(result["output_path"]))
    exported_text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
    assert "组织机构甲" in exported_text
    assert "北京星河科技有限公司" not in exported_text
    assert "北京 星河科技有限公司" not in exported_text
    assert result["docx_rewrite_method"] == "coverage_first_range_then_text"
    assert result["docx_range_rewrite_required_count"] == 1
    assert result["docx_range_rewrite_applied_count"] == 1
    assert result["docx_range_rewrite_unapplied_count"] == 0
    assert result["docx_precise_rewrite_complete"] is True


def test_final_export_does_not_merge_distinct_subjects_by_same_replacement():
    text = "甲方星河公司，乙方景岳公司。"
    entities = [
        {
            "type": "ORGANIZATION",
            "text": "星河公司",
            "start": text.index("星河公司"),
            "end": text.index("星河公司") + len("星河公司"),
            "replacement": "公司1",
            "replacement_method": "contextual",
        },
        {
            "type": "ORGANIZATION",
            "text": "景岳公司",
            "start": text.index("景岳公司"),
            "end": text.index("景岳公司") + len("景岳公司"),
            "replacement": "公司1",
            "replacement_method": "contextual",
        },
    ]

    bundle = build_coverage_first_final_export_bundle(entities=entities, source_text=text)

    assert bundle["summary"]["final_directory_subject_count"] == 2
    assert bundle["summary"]["final_replacement_reused_by_multi_subject_count"] == 1
    assert [row["canonical_text"] for row in bundle["directory_rows"]] == ["星河公司", "景岳公司"]


def test_final_export_excludes_date_amount_and_numeric_masking_preserves_date_year_only():
    text = "合同签订于2026年6月14日，金额10000元，编号AB12345。"
    entities = [
        {
            "type": "DATE",
            "text": "2026年6月14日",
            "start": text.index("2026年6月14日"),
            "end": text.index("2026年6月14日") + len("2026年6月14日"),
            "replacement": "2026年6月14日",
            "replacement_method": "preserve",
        },
        {
            "type": "AMOUNT",
            "text": "10000元",
            "start": text.index("10000元"),
            "end": text.index("10000元") + len("10000元"),
            "replacement": "10000元",
            "replacement_method": "preserve",
        },
    ]

    bundle = build_coverage_first_final_export_bundle(entities=entities, source_text=text)

    assert bundle["enabled"] is False
    assert bundle["summary"]["final_desensitized_entity_input_count"] == 0
    assert bundle["summary"]["final_directory_subject_count"] == 0
    assert mask_default_numeric_text(text) == "合同签订于2026年*月**日，金额*****元，编号AB*****。"


def test_merge_prefers_stronger_review_and_official_sources_over_generic_organization():
    service = RecallFirstEntityMergeService()
    generic = RecognizerResult(
        entity_type="ORGANIZATION",
        start=0,
        end=6,
        score=0.82,
        text="北京星河",
        source="rule_organization",
        metadata={},
    )
    contextual = RecognizerResult(
        entity_type="ORGANIZATION",
        start=0,
        end=6,
        score=0.8,
        text="北京星河",
        source="rule_organization_context",
        metadata={},
    )
    official = RecognizerResult(
        entity_type="GOVERNMENT",
        start=0,
        end=6,
        score=0.79,
        text="北京星河",
        source="rule_official_institution",
        metadata={},
    )
    qwen = RecognizerResult(
        entity_type="ORGANIZATION",
        start=0,
        end=6,
        score=0.81,
        text="北京星河",
        source="qwen_entity_decision",
        metadata={},
    )

    merged = service.merge([generic, contextual, official, qwen])

    assert len(merged) == 1
    assert merged[0].source == "rule_official_institution"
    assert merged[0].entity_type == "GOVERNMENT"


def test_merge_prefers_alias_backscan_review_over_generic_organization():
    service = RecallFirstEntityMergeService()
    generic = RecognizerResult(
        entity_type="ORGANIZATION",
        start=0,
        end=4,
        score=0.82,
        text="星河",
        source="rule_organization",
        metadata={},
    )
    alias_backscan = RecognizerResult(
        entity_type="ORGANIZATION",
        start=0,
        end=4,
        score=0.78,
        text="星河",
        source="alias_backscan_review",
        metadata={},
    )

    merged = service.merge([generic, alias_backscan])

    assert len(merged) == 1
    assert merged[0].source == "alias_backscan_review"


def test_merge_prefers_propagate_over_generic_organization():
    service = RecallFirstEntityMergeService()
    generic = RecognizerResult(
        entity_type="ORGANIZATION",
        start=0,
        end=6,
        score=0.82,
        text="北京星河",
        source="rule_organization",
        metadata={},
    )
    propagated = RecognizerResult(
        entity_type="ORGANIZATION",
        start=0,
        end=6,
        score=0.78,
        text="北京星河",
        source="propagate",
        metadata={"propagated_from_stable_seed": True},
    )

    merged = service.merge([generic, propagated])

    assert len(merged) == 1
    assert merged[0].source == "propagate"


def test_review_rejection_keeps_same_span_other_source_entities():
    result = RecognizerResult(
        entity_type="ORGANIZATION",
        start=0,
        end=6,
        score=0.8,
        text="北京星河",
        source="rule_official_institution",
        metadata={},
    )
    rejected = [
        {
            "text": "北京星河",
            "type": "ORGANIZATION",
            "start": 0,
            "end": 6,
            "source": "qwen_entity_decision",
            "reason": "reject_bad_candidate",
        }
    ]

    assert HighQualityLowMemoryRecognizer._matches_review_rejection(result, rejected) is False


def test_late_deterministic_adjudication_rejects_qwen_and_propagated_left_pollution():
    text = "甲方是乙有限公司负责人。"
    results = [
        RecognizerResult(
            entity_type="ORGANIZATION",
            start=0,
            end=len("甲方是乙有限公司"),
            score=0.87,
            text="甲方是乙有限公司",
            source="qwen_fragment_review",
            metadata={"source_layer": "llm_review"},
        ),
        RecognizerResult(
            entity_type="ORGANIZATION",
            start=2,
            end=len("甲方是乙有限公司"),
            score=0.84,
            text="乙有限公司",
            source="propagate",
            metadata={"propagated_from_subject_ledger": True},
        ),
    ]

    filtered, rejected_count, added_count = HighQualityLowMemoryRecognizer._apply_late_deterministic_adjudication(
        text,
        results,
    )

    assert rejected_count == 1
    assert added_count == 0
    assert [item.text for item in filtered] == ["乙有限公司"]


def test_review_worker_late_deterministic_adjudication_rejects_polluted_final_entities():
    text = "甲方是乙有限公司负责人。"
    recognizer = HighQualityLowMemoryRecognizer()
    entities = [
        {
            "type": "ORGANIZATION",
            "text": "甲方是乙有限公司",
            "start": 0,
            "end": len("甲方是乙有限公司"),
            "score": 0.87,
            "source": "qwen_fragment_review",
            "metadata": {"source_layer": "llm_review"},
        },
        {
            "type": "ORGANIZATION",
            "text": "乙有限公司",
            "start": 2,
            "end": len("甲方是乙有限公司"),
            "score": 0.84,
            "source": "propagate",
            "metadata": {"propagated_from_subject_ledger": True},
        },
    ]

    filtered, rejected_count, added_count = _apply_late_deterministic_adjudication_to_entities(
        text=text,
        entities=entities,
        recognizer=recognizer,
    )

    assert rejected_count == 1
    assert added_count == 0
    assert [item["text"] for item in filtered] == ["乙有限公司"]


def test_residual_repair_does_not_materialize_weak_generic_org_reference():
    service = ContextualDesensitizationService()
    text = "北京星河科技有限公司签署合同。该公司提交材料。"
    weak_ref = "该公司"
    start = text.index(weak_ref)

    additions = service._materialize_residual_hit_entities(
        text=text,
        residual_hits=[
            {
                "canonical_key": "org:star",
                "type": "ORGANIZATION",
                "variant": weak_ref,
                "start": start,
                "end": start + len(weak_ref),
                "weak_reference": True,
            }
        ],
        entity_memory={
            "org:star": {
                "primary_type": "ORGANIZATION",
                "primary_text": "北京星河科技有限公司",
                "canonical_role": "ORGANIZATION",
            }
        },
        existing_entities=[],
    )

    assert additions == []


def test_memory_residual_does_not_materialize_weak_generic_org_reference():
    service = ContextualDesensitizationService()
    text = "北京星河科技有限公司签署合同。该公司提交材料。"

    additions = service._materialize_residual_entities(
        text=text,
        entities=[],
        entity_memory={
            "org:star": {
                "primary_type": "ORGANIZATION",
                "primary_text": "北京星河科技有限公司",
                "canonical_role": "ORGANIZATION",
                "variants": ["该公司"],
            }
        },
        explicit_types=set(),
    )

    assert additions == []


def test_residual_repair_keeps_confirmed_full_subject_occurrence():
    service = ContextualDesensitizationService()
    text = "另见北京星河科技有限公司提交的资料。"
    subject = "北京星河科技有限公司"
    start = text.index(subject)

    additions = service._materialize_residual_hit_entities(
        text=text,
        residual_hits=[
            {
                "canonical_key": "org:star",
                "type": "ORGANIZATION",
                "variant": subject,
                "start": start,
                "end": start + len(subject),
            }
        ],
        entity_memory={
            "org:star": {
                "primary_type": "ORGANIZATION",
                "primary_text": subject,
                "canonical_role": "ORGANIZATION",
            }
        },
        existing_entities=[],
    )

    assert len(additions) == 1
    assert additions[0]["text"] == subject
    assert additions[0]["source"] == "residual_repair"
