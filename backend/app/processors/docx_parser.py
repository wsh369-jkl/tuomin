"""DOCX parser backed by python-docx."""

import asyncio
import logging
import zipfile
from typing import Any, Dict, List
from xml.etree import ElementTree as ET

from docx import Document

from app.processors.base_parser import BaseParser
from app.processors.docx_xml_utils import (
    DocxVisibleTextUnit,
    count_paragraph_page_breaks,
    count_paragraph_section_page_breaks,
    extract_docx_visible_text_units,
)

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    COMPACT_CHARS_PER_PAGE = 1800
    MIN_REASONABLE_COMPACT_CHARS_PER_PAGE = 120
    MAX_REASONABLE_COMPACT_CHARS_PER_PAGE = 5000

    async def parse(self, file_path: str, **kwargs: Any) -> Dict:
        logger.info("Start parsing DOCX: %s", file_path)
        try:
            return await asyncio.to_thread(self._parse_sync, file_path)
        except Exception as exc:
            logger.error("DOCX parsing failed: %s", exc)
            raise

    def _parse_sync(self, file_path: str) -> Dict:
        doc = Document(file_path)
        tracked_metadata: Dict[str, int | bool] = {
            "tracked_changes_detected": False,
            "tracked_insertions": 0,
            "tracked_deletions": 0,
            "tracked_deleted_text_nodes": 0,
        }

        try:
            estimated_page_count = self._estimate_page_count_from_package_hint(file_path)
            text, tracked_metadata, text_units = extract_docx_visible_text_units(
                file_path,
                estimated_page_count=estimated_page_count,
            )
        except Exception as exc:
            logger.warning(
                "DOCX XML extraction failed, falling back to python-docx paragraphs: %s",
                exc,
            )
            text_units = []
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            text = "\n".join(text_parts)

        structure = self._build_structure(file_path, text=text, text_units=text_units)
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "format": "docx",
            "file_path": file_path,
            "pages": int(structure.get("page_count", 1) or 1),
        }
        metadata.update(tracked_metadata)

        logger.info(
            "DOCX parsing complete, paragraphs=%s, text_length=%s, tracked_changes=%s",
            metadata["paragraphs"],
            len(text),
            metadata["tracked_changes_detected"],
        )
        return {
            "text": text,
            "metadata": metadata,
            "structure": structure,
        }

    def supports(self, file_extension: str) -> bool:
        return file_extension.lower() == ".docx"

    def _build_structure(
        self,
        file_path: str,
        *,
        text: str,
        text_units: List[DocxVisibleTextUnit] | None = None,
    ) -> Dict[str, Any]:
        page_count = self._estimate_page_count(file_path, text=text)
        units = list(text_units or [])
        if units:
            self._fill_missing_docx_page_numbers(units=units, page_count=page_count)
        pages = self._build_pages_from_units(text=text, units=units, page_count=page_count)
        if not pages:
            pages = self._build_fallback_pages(text=text, page_count=page_count)
        resolved_page_count = max(1, int(page_count or 0), len(pages))
        return {
            "pages": pages,
            "page_count": resolved_page_count,
            "docx_text_units": [unit.to_dict(include_fragments=True) for unit in units],
            "docx_text_extraction": "ooxml_visible_text_units" if units else "fallback_text",
        }

    def _estimate_page_count_from_package_hint(self, file_path: str) -> int:
        app_page_count = self._read_app_page_count(file_path)
        explicit_page_count = self._count_explicit_pages(file_path)
        if app_page_count > 0:
            return max(1, app_page_count, explicit_page_count)
        if explicit_page_count > 0:
            return max(1, explicit_page_count)
        return 0

    def _estimate_page_count(self, file_path: str, *, text: str) -> int:
        explicit_page_count = self._count_explicit_pages(file_path)
        app_page_count = self._read_app_page_count(file_path)
        heuristic_page_count = self._estimate_text_density_page_count(text)
        if app_page_count > 0:
            return max(1, int(app_page_count), int(explicit_page_count or 0))
        if explicit_page_count > 0:
            return max(1, int(explicit_page_count))
        return max(1, heuristic_page_count)

    def _estimate_text_density_page_count(self, text: str) -> int:
        compact_length = self._compact_text_length(text)
        if compact_length <= 0:
            return 1
        return max(1, round(compact_length / float(self.COMPACT_CHARS_PER_PAGE)))

    def _compact_text_length(self, text: str) -> int:
        return len("".join((text or "").split()))

    def _is_reasonable_page_count_hint(self, page_count: int, *, compact_length: int) -> bool:
        if page_count <= 0:
            return False
        if compact_length <= 0:
            return True
        chars_per_page = compact_length / float(page_count)
        return (
            self.MIN_REASONABLE_COMPACT_CHARS_PER_PAGE
            <= chars_per_page
            <= self.MAX_REASONABLE_COMPACT_CHARS_PER_PAGE
        )

    def _read_app_page_count(self, file_path: str) -> int:
        try:
            with zipfile.ZipFile(file_path) as archive:
                xml_bytes = archive.read("docProps/app.xml")
        except Exception:
            return 0

        try:
            root = ET.fromstring(xml_bytes)
        except ET.ParseError:
            return 0

        for node in root.iter():
            if node.tag.endswith("Pages"):
                try:
                    value = int(str(node.text or "").strip())
                except (TypeError, ValueError):
                    return 0
                return max(0, value)
        return 0

    def _count_explicit_pages(self, file_path: str) -> int:
        try:
            with zipfile.ZipFile(file_path) as archive:
                root = ET.fromstring(archive.read("word/document.xml"))
        except Exception:
            return 0

        page_breaks = 0
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for paragraph in root.iterfind(f".//{namespace}p"):
            page_breaks += count_paragraph_page_breaks(paragraph)
            page_breaks += count_paragraph_section_page_breaks(paragraph)
        if page_breaks <= 0:
            return 0
        return page_breaks + 1

    def _build_pages_from_units(
        self,
        *,
        text: str,
        units: List[Any],
        page_count: int,
    ) -> List[Dict[str, Any]]:
        grouped: Dict[int, Dict[str, Any]] = {}
        last_page_no = 1

        for unit in units:
            unit_dict = (
                unit.to_dict(include_fragments=True)
                if hasattr(unit, "to_dict")
                else dict(unit)
            )
            if str(unit_dict.get("part_name") or "") != "word/document.xml":
                continue
            unit_text = str(unit_dict.get("text", "") or "").strip()
            estimated_page_no = unit_dict.get("estimated_page_no")
            if isinstance(estimated_page_no, int) and estimated_page_no > 0:
                page_no = estimated_page_no
            else:
                page_no = last_page_no
            last_page_no = max(last_page_no, page_no)
            page = grouped.setdefault(
                page_no,
                {
                    "page_number": page_no,
                    "texts": [],
                    "units": [],
                },
            )
            page["units"].append(unit_dict)
            if unit_text:
                page["texts"].append(unit_text)
            if self._coerce_int(unit_dict.get("start"), -1) >= 0:
                starts = page.setdefault("unit_starts", [])
                starts.append(self._coerce_int(unit_dict.get("start"), -1))
            if self._coerce_int(unit_dict.get("end"), -1) >= 0:
                ends = page.setdefault("unit_ends", [])
                ends.append(self._coerce_int(unit_dict.get("end"), -1))

        if not grouped:
            return []

        max_page_no = max(max(grouped), page_count, 1)
        ordered_pages: List[Dict[str, Any]] = []
        cursor = 0
        for page_no in range(1, max_page_no + 1):
            page = grouped.get(page_no) or {"page_number": page_no, "texts": [], "units": []}
            page_text = "\n".join(page.get("texts", [])).strip()
            unit_starts = [value for value in page.get("unit_starts", []) if isinstance(value, int) and value >= 0]
            unit_ends = [value for value in page.get("unit_ends", []) if isinstance(value, int) and value >= 0]
            if unit_starts and unit_ends:
                start = min(unit_starts)
                end = max(unit_ends)
                cursor = end
            else:
                start, end, cursor = self._locate_page_span(text=text, page_text=page_text, cursor=cursor)
            ordered_pages.append(
                {
                    "page_number": page_no,
                    "text": page_text,
                    "start": start,
                    "end": end,
                    "units": list(page.get("units", [])),
                }
            )
        return ordered_pages

    def _locate_page_span(self, *, text: str, page_text: str, cursor: int) -> tuple[int, int, int]:
        bounded_cursor = max(0, min(cursor, len(text)))
        if not page_text:
            return bounded_cursor, bounded_cursor, bounded_cursor

        start = text.find(page_text, bounded_cursor)
        if start < 0:
            start = text.find(page_text)
        if start < 0:
            start = bounded_cursor
            end = min(len(text), start + len(page_text))
            return start, end, end

        end = start + len(page_text)
        next_cursor = end
        while next_cursor < len(text) and text[next_cursor] == "\n":
            next_cursor += 1
        return start, end, next_cursor

    def _build_fallback_pages(self, *, text: str, page_count: int) -> List[Dict[str, Any]]:
        if not text:
            return [{"page_number": 1, "text": "", "start": 0, "end": 0, "units": []}]
        total_pages = max(1, page_count)
        if total_pages == 1:
            return [{"page_number": 1, "text": text, "start": 0, "end": len(text), "units": []}]

        chunk_size = max(1, len(text) // total_pages)
        pages: List[Dict[str, Any]] = []
        start = 0
        for page_no in range(1, total_pages + 1):
            if page_no == total_pages:
                end = len(text)
            else:
                end = min(len(text), start + chunk_size)
                newline_pos = text.rfind("\n", start, end)
                if newline_pos > start:
                    end = newline_pos
            page_text = text[start:end].strip("\n")
            actual_start = start
            if page_text:
                stripped_left = len(text[start:end]) - len(text[start:end].lstrip("\n"))
                actual_start = start + stripped_left
                actual_end = actual_start + len(page_text)
            else:
                actual_end = actual_start
            pages.append(
                {
                    "page_number": page_no,
                    "text": page_text,
                    "start": actual_start,
                    "end": actual_end,
                    "units": [],
                }
            )
            start = max(end, start + 1)
        return pages

    def _fill_missing_docx_page_numbers(
        self,
        *,
        units: List[DocxVisibleTextUnit],
        page_count: int,
    ) -> None:
        body_units = [unit for unit in units if unit.part_name == "word/document.xml"]
        total = max(1, len(body_units))
        for index, unit in enumerate(body_units):
            if isinstance(unit.estimated_page_no, int) and unit.estimated_page_no > 0:
                continue
            if page_count <= 1:
                unit.estimated_page_no = 1
                continue
            estimated = int((index / total) * page_count) + 1
            unit.estimated_page_no = max(1, min(page_count, estimated))

    def _coerce_int(self, value: Any, default: int = 0) -> int:
        try:
            return int(value)
        except (TypeError, ValueError):
            return default
