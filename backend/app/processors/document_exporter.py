"""Export anonymized content while preserving the source format when possible."""

from __future__ import annotations

import copy
import logging
import re
import shutil
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips
from PIL import Image, ImageDraw
import pypdfium2 as pdfium

from app.core.config import settings
from app.core.runtime_security import ensure_private_file
from app.engine.desensitization_engine import get_engine
from app.processors.docx_xml_utils import (
    apply_replacements_to_fragments,
    docx_contains_tracked_changes,
    normalize_docx_chinese_inline_spaces,
    replace_docx_text_by_ranges_with_report,
    replace_text_in_docx,
)
from app.services.default_numeric_masking import mask_default_numeric_text
from app.services.lowmem_entity_utils import sanitize_recognition_text

logger = logging.getLogger(__name__)

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TEXT_MIME_TYPE = "text/plain; charset=utf-8"
PDF_MIME_TYPE = "application/pdf"
DEFAULT_PDF_PAGE_WIDTH_PT = 595.0
DEFAULT_PDF_PAGE_HEIGHT_PT = 842.0
DEFAULT_DOCX_TABLE_WIDTH_DXA = 9360
DEFAULT_DOCX_TABLE_CELL_MARGINS_DXA = {
    "top": 120,
    "bottom": 120,
    "start": 120,
    "end": 120,
}
DEFAULT_DOCX_TABLE_INDENT_DXA = DEFAULT_DOCX_TABLE_CELL_MARGINS_DXA["start"]
DOCX_MAPPING_TABLE_HEADER_FILL = "E8EEF5"
DOCX_MAPPING_TABLE_TITLE = "替换对照"
DOCX_MAPPING_FONT_NAME = "PingFang SC"


class DocumentExporter:
    """Create downloadable output files for anonymized content."""

    EXPORT_TEXT_PRIORITY = {
        "PERSON": 1,
        "PERSON_NAME": 1,
        "ORGANIZATION": 2,
        "COMPANY_NAME": 2,
        "BANK_NAME": 3,
        "ACCOUNT_NAME": 4,
        "LOCATION": 5,
        "PROJECT": 6,
        "CONTRACT_NO": 7,
        "POSITION": 8,
        "CN_ID_CARD": 9,
        "CN_PHONE": 9,
        "LANDLINE_PHONE": 9,
        "CN_BANK_CARD": 9,
        "CN_CREDIT_CODE": 9,
        "EMAIL_ADDRESS": 9,
        "AMOUNT": 9,
    }
    EXPORT_ENTITY_TYPE_LABELS = {
        "PERSON": "人名",
        "PERSON_NAME": "人名",
        "ORGANIZATION": "组织机构",
        "COMPANY_NAME": "公司名称",
        "COURT": "法院",
        "ADDRESS": "地址",
        "LOCATION": "地址",
        "ALIAS": "简称",
        "POSITION": "职位",
        "CN_ID_CARD": "身份证号",
        "CN_PHONE": "手机号",
        "LANDLINE_PHONE": "座机号",
        "CN_BANK_CARD": "银行卡号",
        "CN_CREDIT_CODE": "统一社会信用代码",
        "EMAIL_ADDRESS": "邮箱",
        "AMOUNT": "金额",
        "PROJECT_CODE": "项目代号",
        "CONTRACT_NO": "合同编号",
        "CASE_NO": "案号",
        "PRODUCT_NAME": "产品名称",
        "SENSITIVE_TERM": "敏感术语",
        "PROJECT": "项目名称",
        "BANK_NAME": "开户行",
        "ACCOUNT_NAME": "户名",
    }

    def export(
        self,
        *,
        task_id: str,
        source_path: str,
        original_filename: str,
        source_text: str | None,
        source_metadata: Dict | None,
        source_structure: Dict | None,
        entities: List[Dict],
        anonymized_text: str,
        operator_config: Dict | None = None,
        coverage_first_final_export: Dict[str, Any] | None = None,
    ) -> Dict[str, str | bool | None]:
        suffix = Path(source_path).suffix.lower()
        coverage_first_final_export = (
            dict(coverage_first_final_export)
            if isinstance(coverage_first_final_export, dict)
            else {}
        )
        anonymized_text, text_inline_space_removed = self._normalize_export_text(anonymized_text or "")
        mapping_entities = self._build_mapping_export_entities(
            entities=entities,
            coverage_first_final_export=coverage_first_final_export,
        )
        replacement_entries = self._build_export_replacement_entries(entities, operator_config)
        final_rewrite_entries = self._build_final_rewrite_entries(coverage_first_final_export)
        if final_rewrite_entries:
            replacement_entries = final_rewrite_entries

        if suffix == ".docx":
            try:
                primary_result = self._export_docx(
                    task_id=task_id,
                    source_path=source_path,
                    original_filename=original_filename,
                    source_text=source_text,
                    replacement_entries=replacement_entries,
                    range_rewrite_entries=final_rewrite_entries,
                )
            except Exception as exc:
                logger.exception("DOCX export failed, falling back to TXT: %s", exc)
                primary_result = self._export_text(
                    task_id=task_id,
                    original_filename=original_filename,
                    anonymized_text=anonymized_text,
                    warning="DOCX export failed, so the result was saved as TXT.",
                    preserves_format=False,
                )
            primary_result = self._attach_text_inline_space_metadata(
                primary_result,
                removed_count=text_inline_space_removed,
            )
            return self._attach_mapping_export(
                primary_result=primary_result,
                task_id=task_id,
                original_filename=original_filename,
                entities=mapping_entities,
            )

        if suffix == ".txt":
            primary_result = self._export_text(
                task_id=task_id,
                original_filename=original_filename,
                anonymized_text=anonymized_text,
                warning=None,
                preserves_format=True,
            )
            primary_result = self._attach_text_inline_space_metadata(
                primary_result,
                removed_count=text_inline_space_removed,
            )
            return self._attach_mapping_export(
                primary_result=primary_result,
                task_id=task_id,
                original_filename=original_filename,
                entities=mapping_entities,
            )

        if suffix == ".pdf":
            primary_result = self._export_pdf_as_text(
                task_id=task_id,
                original_filename=original_filename,
                anonymized_text=anonymized_text,
                source_metadata=source_metadata,
            )
            primary_result = self._attach_text_inline_space_metadata(
                primary_result,
                removed_count=text_inline_space_removed,
            )
            return self._attach_mapping_export(
                primary_result=primary_result,
                task_id=task_id,
                original_filename=original_filename,
                entities=mapping_entities,
            )

        primary_result = self._export_text(
            task_id=task_id,
            original_filename=original_filename,
            anonymized_text=anonymized_text,
            warning="Format-preserving export is currently available for DOCX and TXT. PDF currently falls back to TXT.",
            preserves_format=False,
        )
        primary_result = self._attach_text_inline_space_metadata(
            primary_result,
            removed_count=text_inline_space_removed,
        )
        return self._attach_mapping_export(
            primary_result=primary_result,
            task_id=task_id,
            original_filename=original_filename,
            entities=mapping_entities,
        )

    def _normalize_export_text(self, value: str) -> tuple[str, int]:
        sanitized, index_map = sanitize_recognition_text(value or "")
        removed_count = max(0, len(value or "") - len(index_map))
        return mask_default_numeric_text(sanitized), removed_count

    def _attach_text_inline_space_metadata(
        self,
        result: Dict[str, str | bool | None],
        *,
        removed_count: int,
    ) -> Dict[str, str | bool | int | None]:
        updated = dict(result)
        updated["text_chinese_inline_space_normalized"] = removed_count > 0
        updated["text_chinese_inline_space_removed_count"] = int(removed_count)
        return updated

    def _attach_mapping_export(
        self,
        *,
        primary_result: Dict[str, str | bool | None],
        task_id: str,
        original_filename: str,
        entities: List[Dict[str, Any]],
    ) -> Dict[str, str | bool | None]:
        result = dict(primary_result)
        try:
            mapping_result = self._export_mapping_docx(
                task_id=task_id,
                original_filename=original_filename,
                entities=entities,
            )
        except Exception:
            logger.exception("Mapping directory export failed for %s", original_filename)
            warning = str(result.get("warning") or "").strip()
            mapping_warning = "Mapping directory export failed; please retry the export."
            result["warning"] = f"{warning} | {mapping_warning}" if warning else mapping_warning
            return result

        result.update(mapping_result)
        return result

    def _build_mapping_export_entities(
        self,
        *,
        entities: Iterable[Dict[str, Any]],
        coverage_first_final_export: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        """Build mapping rows from the same occurrence layer used for final rewrite."""

        if isinstance(coverage_first_final_export, dict):
            from_mapping_entities = [
                dict(item)
                for item in coverage_first_final_export.get("mapping_entities") or []
                if isinstance(item, dict)
            ]
            if from_mapping_entities:
                return from_mapping_entities
            from_directory_rows = self._mapping_entities_from_directory_rows(coverage_first_final_export)
            if from_directory_rows:
                return from_directory_rows
            from_rewrite_entries = self._mapping_entities_from_rewrite_entries(coverage_first_final_export)
            if from_rewrite_entries:
                return from_rewrite_entries
        return [dict(entity) for entity in entities or [] if isinstance(entity, dict)]

    def _mapping_entities_from_rewrite_entries(
        self,
        coverage_first_final_export: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        rewrite_entries = [
            item
            for item in coverage_first_final_export.get("rewrite_entries") or []
            if isinstance(item, dict)
        ]
        if not rewrite_entries:
            return []
        row_metadata = self._coverage_directory_row_metadata(coverage_first_final_export)
        mapping_entities: List[Dict[str, Any]] = []
        for input_order, entry in enumerate(rewrite_entries):
            status = str(entry.get("verification_status") or "").strip()
            if status == "blocked":
                continue
            source_text = str(entry.get("source_text") or "").strip()
            replacement = self._coerce_replacement_text(entry.get("replacement"))
            if not source_text or not replacement or source_text == replacement:
                continue
            subject_id = str(entry.get("subject_id") or "").strip()
            metadata = dict(row_metadata.get(subject_id) or {})
            mapping_entities.append(
                {
                    "type": str(entry.get("entity_type") or metadata.get("subject_type") or "").strip(),
                    "text": source_text,
                    "replacement": replacement,
                    "start": self._coerce_int(entry.get("start"), 10**12),
                    "end": self._coerce_int(entry.get("end"), 10**12),
                    "context_label": metadata.get("context_label") or "未命名字段",
                    "context_role": metadata.get("context_role") or "通用上下文",
                    "metadata": {
                        "coverage_first_subject_id": subject_id,
                        "coverage_first_mapping_source": "rewrite_entries",
                        "coverage_first_input_order": self._coerce_int(entry.get("input_order"), input_order),
                    },
                }
            )
        return mapping_entities

    def _mapping_entities_from_directory_rows(
        self,
        coverage_first_final_export: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        directory_rows = [
            item
            for item in coverage_first_final_export.get("directory_rows") or []
            if isinstance(item, dict)
        ]
        mapping_entities: List[Dict[str, Any]] = []
        for row in directory_rows:
            replacement = self._coerce_replacement_text(row.get("replacement"))
            if not replacement:
                continue
            subject_id = str(row.get("subject_id") or "").strip()
            for occurrence_index, occurrence in enumerate(row.get("occurrences") or [], start=1):
                if not isinstance(occurrence, dict):
                    continue
                source_text = str(occurrence.get("text") or "").strip()
                if not source_text or source_text == replacement:
                    continue
                mapping_entities.append(
                    {
                        "type": str(occurrence.get("type") or row.get("subject_type") or "").strip(),
                        "text": source_text,
                        "replacement": replacement,
                        "start": self._coerce_int(occurrence.get("start"), 10**12),
                        "end": self._coerce_int(occurrence.get("end"), 10**12),
                        "context_label": row.get("context_label") or "未命名字段",
                        "context_role": row.get("context_role") or "通用上下文",
                        "metadata": {
                            "coverage_first_subject_id": subject_id,
                            "coverage_first_mapping_source": "directory_rows",
                            "coverage_first_occurrence_index": occurrence_index,
                        },
                    }
                )
        return mapping_entities

    def _coverage_directory_row_metadata(
        self,
        coverage_first_final_export: Dict[str, Any],
    ) -> Dict[str, Dict[str, Any]]:
        rows = [
            item
            for item in coverage_first_final_export.get("directory_rows") or []
            if isinstance(item, dict)
        ]
        metadata: Dict[str, Dict[str, Any]] = {}
        for row in rows:
            subject_id = str(row.get("subject_id") or "").strip()
            if not subject_id:
                continue
            metadata[subject_id] = {
                "subject_type": row.get("subject_type"),
                "context_label": row.get("context_label"),
                "context_role": row.get("context_role"),
            }
        return metadata

    def _export_text(
        self,
        *,
        task_id: str,
        original_filename: str,
        anonymized_text: str,
        warning: str | None,
        preserves_format: bool,
    ) -> Dict[str, str | bool | None]:
        output_path = Path(settings.OUTPUT_DIR) / f"{task_id}_anonymized.txt"
        output_path.write_text(anonymized_text, encoding="utf-8")
        ensure_private_file(output_path)

        return {
            "output_path": str(output_path),
            "download_name": f"{Path(original_filename).stem}_anonymized.txt",
            "file_type": "txt",
            "media_type": TEXT_MIME_TYPE,
            "preserves_format": preserves_format,
            "warning": warning,
        }

    def _export_pdf_as_text(
        self,
        *,
        task_id: str,
        original_filename: str,
        anonymized_text: str,
        source_metadata: Dict | None,
    ) -> Dict[str, str | bool | None]:
        ocr_pages = self._extract_pdf_ocr_page_count(source_metadata)
        ocr_timeout_pages = self._extract_pdf_ocr_timeout_page_count(source_metadata)
        ocr_failed_pages = self._extract_pdf_ocr_failed_page_count(source_metadata)

        warning = "PDF 已提取为纯文本，并按既有文本脱敏策略导出为 TXT，便于后续继续交给 AI 审查。"
        if ocr_pages > 0:
            warning += f" 其中 {ocr_pages} 页来自 OCR 文本提取。"
        if ocr_timeout_pages > 0:
            warning += f" {ocr_timeout_pages} 页 OCR 超时，相关文本可能不完整，建议人工复核原 PDF。"
        elif ocr_failed_pages > 0:
            warning += f" {ocr_failed_pages} 页 OCR 未完成，相关文本可能缺失，建议人工复核原 PDF。"

        return self._export_text(
            task_id=task_id,
            original_filename=original_filename,
            anonymized_text=anonymized_text,
            warning=warning,
            preserves_format=False,
        )

    def _export_docx(
        self,
        *,
        task_id: str,
        source_path: str,
        original_filename: str,
        source_text: str | None,
        replacement_entries: List[Dict[str, Any]],
        range_rewrite_entries: List[Dict[str, Any]] | None = None,
    ) -> Dict[str, str | bool | None]:
        replacements = self._build_replacements_from_entries(replacement_entries)
        output_path = Path(settings.OUTPUT_DIR) / f"{task_id}_anonymized.docx"
        shutil.copy2(source_path, output_path)
        range_report: Dict[str, Any] = {
            "applied": False,
            "attempted_entry_count": 0,
            "normalized_entry_count": 0,
            "applied_entry_count": 0,
            "unapplied_entry_count": 0,
            "modified_part_count": 0,
            "modified_parts": [],
            "rejection_reason": "",
            "unapplied_ranges": [],
        }
        range_entries = list(range_rewrite_entries or [])
        if range_entries:
            try:
                range_report = replace_docx_text_by_ranges_with_report(
                    output_path,
                    range_entries,
                    source_text=source_text,
                )
            except Exception:
                logger.exception("DOCX final range rewrite failed for %s", output_path)
        precise_rewrite_complete = (
            bool(range_entries)
            and int(range_report.get("normalized_entry_count") or 0) > 0
            and int(range_report.get("unapplied_entry_count") or 0) == 0
            and not str(range_report.get("rejection_reason") or "").strip()
        )
        if precise_rewrite_complete:
            replacements = []
        if replacements:
            try:
                document = Document(output_path)
                self._replace_in_container(document, replacements)
                document.save(output_path)
            except Exception:
                logger.exception("DOCX visible-text replacement failed for %s", output_path)
            try:
                if docx_contains_tracked_changes(output_path):
                    replace_text_in_docx(output_path, replacements)
            except Exception:
                logger.exception("DOCX tracked-change XML replacement failed for %s", output_path)
        inline_space_report: Dict[str, Any] = {
            "applied": False,
            "removed_space_count": 0,
            "modified_part_count": 0,
            "modified_parts": [],
        }
        try:
            inline_space_report = normalize_docx_chinese_inline_spaces(output_path)
        except Exception:
            logger.exception("DOCX Chinese inline space normalization failed for %s", output_path)
        ensure_private_file(output_path)

        return {
            "output_path": str(output_path),
            "download_name": f"{Path(original_filename).stem}_anonymized.docx",
            "file_type": "docx",
            "media_type": DOCX_MIME_TYPE,
            "preserves_format": True,
            "warning": None,
            "docx_rewrite_method": "coverage_first_range_then_text" if range_entries else "text_replacement",
            "docx_precise_rewrite_complete": precise_rewrite_complete if range_entries else None,
            "docx_range_rewrite_required_count": int(range_report.get("normalized_entry_count") or 0),
            "docx_range_rewrite_applied_count": int(range_report.get("applied_entry_count") or 0),
            "docx_range_rewrite_unapplied_count": int(range_report.get("unapplied_entry_count") or 0),
            "docx_range_rewrite_preflight_blocked_count": int(range_report.get("attempted_entry_count") or 0)
            - int(range_report.get("normalized_entry_count") or 0),
            "docx_range_rewrite_rejection_reason": str(range_report.get("rejection_reason") or ""),
            "docx_range_rewrite_modified_part_count": int(range_report.get("modified_part_count") or 0),
            "docx_range_rewrite_modified_parts": list(range_report.get("modified_parts") or [])[:20],
            "docx_range_rewrite_unapplied_ranges": list(range_report.get("unapplied_ranges") or [])[:20],
            "docx_chinese_inline_space_normalized": bool(inline_space_report.get("applied")),
            "docx_chinese_inline_space_removed_count": int(inline_space_report.get("removed_space_count") or 0),
            "docx_chinese_inline_space_modified_part_count": int(inline_space_report.get("modified_part_count") or 0),
            "docx_chinese_inline_space_modified_parts": list(inline_space_report.get("modified_parts") or [])[:20],
        }

    def _replace_in_container(self, container, replacements: List[Tuple[str, str]]) -> None:
        for paragraph in container.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._replace_in_container(cell, replacements)

    def _replace_in_paragraph(self, paragraph, replacements: List[Tuple[str, str]]) -> None:
        if not paragraph.runs:
            return

        run_texts = [run.text for run in paragraph.runs]
        if not any(run_texts):
            return

        updated_texts = apply_replacements_to_fragments(run_texts, replacements)
        updated_texts = [mask_default_numeric_text(text) for text in updated_texts]
        if updated_texts == run_texts:
            return

        for run, updated_text in zip(paragraph.runs, updated_texts):
            run.text = updated_text

    def _export_pdf_as_docx(
        self,
        *,
        task_id: str,
        original_filename: str,
        source_text: str | None,
        source_metadata: Dict | None,
        source_structure: Dict | None,
        entities: List[Dict],
        operator_config: Dict | None,
    ) -> Dict[str, str | bool | None]:
        document = Document()
        self._remove_default_paragraph(document)

        replacements = self._build_replacements(entities, operator_config)
        pages = self._build_pdf_pages(
            source_text=source_text,
            source_metadata=source_metadata,
            source_structure=source_structure,
            replacements=replacements,
        )

        if not pages:
            pages = [
                {
                    "width": DEFAULT_PDF_PAGE_WIDTH_PT,
                    "height": DEFAULT_PDF_PAGE_HEIGHT_PT,
                    "blocks": self._build_blocks_from_text_fallback(
                        self._apply_text_replacements(source_text or "", replacements)
                    ),
                }
            ]

        for page_index, page in enumerate(pages):
            self._configure_pdf_section(document, page, page_index)
            self._append_structured_pdf_page(document, page)

        output_path = Path(settings.OUTPUT_DIR) / f"{task_id}_anonymized.docx"
        document.save(output_path)
        ensure_private_file(output_path)

        ocr_pages = 0
        ocr_timeout_pages = 0
        ocr_failed_pages = 0
        if isinstance(source_metadata, dict):
            try:
                ocr_pages = int(source_metadata.get("ocr_pages", 0) or 0)
            except (TypeError, ValueError):
                ocr_pages = 0
            try:
                ocr_timeout_pages = int(source_metadata.get("ocr_timeout_pages", 0) or 0)
            except (TypeError, ValueError):
                ocr_timeout_pages = 0
            try:
                ocr_failed_pages = int(source_metadata.get("ocr_failed_pages", 0) or 0)
            except (TypeError, ValueError):
                ocr_failed_pages = 0

        warning = (
            "PDF has been reconstructed into an editable DOCX with page-level layout recovery. "
            "Titles, line breaks, spacing, and simple tables are preserved as much as possible."
        )
        if ocr_pages > 0:
            warning += " Scanned pages still depend on OCR, so very complex layouts may differ from the original."
        if ocr_timeout_pages > 0:
            warning += f" OCR timed out on {ocr_timeout_pages} scanned page(s), so those pages may be partially missing and should be reviewed manually."
        elif ocr_failed_pages > 0:
            warning += f" OCR could not be completed on {ocr_failed_pages} scanned page(s), so those pages may need manual review."

        return {
            "output_path": str(output_path),
            "download_name": f"{Path(original_filename).stem}_anonymized.docx",
            "file_type": "docx",
            "media_type": DOCX_MIME_TYPE,
            "preserves_format": False,
            "warning": warning,
        }

    def _export_redacted_pdf(
        self,
        *,
        task_id: str,
        source_path: str,
        original_filename: str,
        source_text: str | None,
        source_structure: Dict | None,
        source_metadata: Dict | None,
        entities: List[Dict],
    ) -> Dict[str, str | bool | None]:
        redaction_plan = self._build_pdf_redaction_plan(
            source_text=source_text,
            source_structure=source_structure,
            entities=entities,
        )
        if entities and not any(redaction_plan.values()):
            raise ValueError("No redaction rectangles could be resolved for scanned PDF export.")

        pdf_document = pdfium.PdfDocument(source_path)
        render_scale = max(settings.PDF_OCR_RENDER_SCALE, 1.0)
        rendered_pages: List[Image.Image] = []

        try:
            for page_index in range(len(pdf_document)):
                page = pdf_document[page_index]
                bitmap = page.render(
                    scale=render_scale,
                    rev_byteorder=True,
                )
                image = bitmap.to_pil().convert("RGB")
                bitmap.close()
                page.close()

                rectangles = redaction_plan.get(page_index + 1, [])
                if rectangles:
                    self._draw_redaction_rectangles(image, rectangles)
                rendered_pages.append(image)
        finally:
            pdf_document.close()

        if not rendered_pages:
            raise ValueError("PDF contains no renderable pages.")

        output_path = Path(settings.OUTPUT_DIR) / f"{task_id}_anonymized.pdf"
        resolution = int(round(72 * render_scale))
        first_page = rendered_pages[0]
        remaining_pages = rendered_pages[1:]
        try:
            first_page.save(
                output_path,
                "PDF",
                save_all=True,
                append_images=remaining_pages,
                resolution=resolution,
            )
        finally:
            for image in rendered_pages:
                image.close()
        ensure_private_file(output_path)

        ocr_pages = self._extract_pdf_ocr_page_count(source_metadata)
        ocr_timeout_pages = self._extract_pdf_ocr_timeout_page_count(source_metadata)
        ocr_failed_pages = self._extract_pdf_ocr_failed_page_count(source_metadata)
        warning = (
            "扫描型 PDF 已按识别命中的关键字符位置直接加黑框导出。"
            " 如遇极小字号、倾斜扫描或 OCR 低质量页，建议人工复核。"
        )
        if ocr_pages > 1:
            warning += f" 当前共有 {ocr_pages} 页走 OCR 路线。"
        if ocr_timeout_pages > 0:
            warning += f" 其中 {ocr_timeout_pages} 页 OCR 超时，可能仍有未覆盖区域。"
        elif ocr_failed_pages > 0:
            warning += f" 其中 {ocr_failed_pages} 页 OCR 未完成，建议重点复核。"

        return {
            "output_path": str(output_path),
            "download_name": f"{Path(original_filename).stem}_anonymized.pdf",
            "file_type": "pdf",
            "media_type": PDF_MIME_TYPE,
            "preserves_format": True,
            "warning": warning,
        }

    def _build_pdf_redaction_plan(
        self,
        *,
        source_text: str | None,
        source_structure: Dict | None,
        entities: List[Dict],
    ) -> Dict[int, List[List[float]]]:
        if not isinstance(source_structure, dict):
            return {}

        raw_pages = source_structure.get("pages")
        if not isinstance(raw_pages, list):
            return {}

        page_offsets = self._build_pdf_page_offset_map(raw_pages)
        redaction_plan: Dict[int, List[List[float]]] = {}
        for page_index, raw_page in enumerate(raw_pages):
            if not isinstance(raw_page, dict):
                continue
            page_number = int(raw_page.get("page_number", page_index + 1) or (page_index + 1))
            page_text = str(raw_page.get("text", "") or "")
            offset_entry = page_offsets.get(page_index)
            if not page_text or not offset_entry:
                redaction_plan[page_number] = []
                continue

            line_map = self._build_page_line_map(page_text, raw_page.get("lines"))
            if not line_map["lines"]:
                redaction_plan[page_number] = []
                continue

            page_rectangles: List[List[float]] = []
            for entity in entities:
                if not isinstance(entity, dict):
                    continue
                if entity.get("replacement_method") == "preserve":
                    continue
                if entity.get("replacement") and entity.get("replacement") == entity.get("text"):
                    continue

                entity_start = int(entity.get("start", 0) or 0)
                entity_end = int(entity.get("end", 0) or 0)
                if entity_end <= offset_entry["start"] or entity_start >= offset_entry["end"]:
                    continue

                local_start = max(entity_start, offset_entry["start"]) - offset_entry["start"]
                local_end = min(entity_end, offset_entry["end"]) - offset_entry["start"]
                page_rectangles.extend(
                    self._resolve_entity_rectangles_for_page(
                        page_text=page_text,
                        line_map=line_map,
                        local_start=local_start,
                        local_end=local_end,
                    )
                )

            redaction_plan[page_number] = self._merge_relative_rectangles(page_rectangles)

        return redaction_plan

    def _build_pdf_page_offset_map(self, raw_pages: List[Dict[str, Any]]) -> Dict[int, Dict[str, int]]:
        non_empty_pages = [
            (index, str(page.get("text", "") or "").strip())
            for index, page in enumerate(raw_pages)
            if isinstance(page, dict) and str(page.get("text", "") or "").strip()
        ]
        offsets: Dict[int, Dict[str, int]] = {}
        cursor = 0
        for order, (page_index, page_text) in enumerate(non_empty_pages):
            start = cursor
            end = start + len(page_text)
            offsets[page_index] = {"start": start, "end": end}
            cursor = end
            if order < len(non_empty_pages) - 1:
                cursor += 2
        return offsets

    def _build_page_line_map(self, page_text: str, lines: Any) -> Dict[str, Any]:
        compact_page_text, raw_to_compact = self._build_compact_text_map(page_text)
        search_cursor = 0
        prepared_lines: List[Dict[str, Any]] = []
        if isinstance(lines, list):
            for item in lines:
                if not isinstance(item, dict):
                    continue
                text = str(item.get("text", "") or "").strip()
                bbox = self._coerce_relative_bbox(item.get("bbox"))
                compact_line = self._compact_for_redaction(text)
                if not text or bbox is None or not compact_line:
                    continue

                compact_start = compact_page_text.find(compact_line, search_cursor)
                if compact_start == -1:
                    compact_start = compact_page_text.find(compact_line)
                if compact_start == -1:
                    continue

                compact_end = compact_start + len(compact_line)
                prepared_lines.append(
                    {
                        "text": text,
                        "compact_text": compact_line,
                        "compact_start": compact_start,
                        "compact_end": compact_end,
                        "bbox": bbox,
                    }
                )
                search_cursor = compact_end

        return {
            "compact_page_text": compact_page_text,
            "raw_to_compact": raw_to_compact,
            "lines": prepared_lines,
        }

    def _build_compact_text_map(self, text: str) -> Tuple[str, List[int]]:
        compact_chars: List[str] = []
        raw_to_compact = [0]
        for char in text:
            if not char.isspace():
                compact_chars.append(char)
            raw_to_compact.append(len(compact_chars))
        return "".join(compact_chars), raw_to_compact

    def _resolve_entity_rectangles_for_page(
        self,
        *,
        page_text: str,
        line_map: Dict[str, Any],
        local_start: int,
        local_end: int,
    ) -> List[List[float]]:
        if local_end <= local_start:
            return []

        raw_to_compact = line_map.get("raw_to_compact") or []
        if not raw_to_compact:
            _, raw_to_compact = self._build_compact_text_map(page_text)
        local_start = max(0, min(local_start, len(page_text)))
        local_end = max(local_start, min(local_end, len(page_text)))
        compact_start = raw_to_compact[local_start]
        compact_end = raw_to_compact[local_end]
        if compact_end <= compact_start:
            return []

        rectangles: List[List[float]] = []
        for line in line_map.get("lines", []):
            overlap_start = max(compact_start, int(line.get("compact_start", 0) or 0))
            overlap_end = min(compact_end, int(line.get("compact_end", 0) or 0))
            if overlap_end <= overlap_start:
                continue

            line_local_start = overlap_start - int(line.get("compact_start", 0) or 0)
            line_local_end = overlap_end - int(line.get("compact_start", 0) or 0)
            rectangle = self._estimate_line_segment_rectangle(
                compact_text=str(line.get("compact_text", "") or ""),
                bbox=line.get("bbox"),
                start_index=line_local_start,
                end_index=line_local_end,
            )
            if rectangle is not None:
                rectangles.append(rectangle)

        return rectangles

    def _estimate_line_segment_rectangle(
        self,
        *,
        compact_text: str,
        bbox: Any,
        start_index: int,
        end_index: int,
    ) -> List[float] | None:
        normalized_bbox = self._coerce_relative_bbox(bbox)
        if normalized_bbox is None or not compact_text:
            return None

        start_index = max(0, min(start_index, len(compact_text)))
        end_index = max(start_index, min(end_index, len(compact_text)))
        if end_index <= start_index:
            return None

        weights = [self._estimate_character_weight(char) for char in compact_text]
        total_weight = sum(weights) or float(len(compact_text))
        cumulative = [0.0]
        for weight in weights:
            cumulative.append(cumulative[-1] + weight)

        left, top, right, bottom = normalized_bbox
        line_width = max(right - left, 0.001)
        line_height = max(bottom - top, 0.001)
        segment_left = left + line_width * (cumulative[start_index] / total_weight)
        segment_right = left + line_width * (cumulative[end_index] / total_weight)
        minimum_width = min(line_width, max(line_height * 0.5, 0.006))
        if segment_right - segment_left < minimum_width:
            center = (segment_left + segment_right) / 2
            half = minimum_width / 2
            segment_left = max(left, center - half)
            segment_right = min(right, center + half)

        pad_x = min(line_width * 0.04, max(line_height * 0.18, 0.002))
        pad_y = min(line_height * 0.24, 0.006)
        return [
            round(max(0.0, segment_left - pad_x), 5),
            round(max(0.0, top - pad_y), 5),
            round(min(1.0, segment_right + pad_x), 5),
            round(min(1.0, bottom + pad_y), 5),
        ]

    def _estimate_character_weight(self, char: str) -> float:
        if re.match(r"[\u4e00-\u9fff]", char):
            return 1.0
        if char.isdigit():
            return 0.62
        if re.match(r"[A-Za-z]", char):
            return 0.62
        if char in "（）()【】[]《》<>":
            return 0.55
        if char in "：:；;，,。.!！？?、/\\|_-":
            return 0.42
        return 0.58

    def _draw_redaction_rectangles(
        self,
        image: Image.Image,
        rectangles: List[List[float]],
    ) -> None:
        if not rectangles:
            return

        draw = ImageDraw.Draw(image)
        width, height = image.size
        for left, top, right, bottom in rectangles:
            x0 = max(0, min(int(left * width), width - 1))
            y0 = max(0, min(int(top * height), height - 1))
            x1 = max(x0 + 1, min(int(right * width), width))
            y1 = max(y0 + 1, min(int(bottom * height), height))
            draw.rectangle((x0, y0, x1, y1), fill="black")

    def _merge_relative_rectangles(self, rectangles: List[List[float]]) -> List[List[float]]:
        if not rectangles:
            return []

        ordered = sorted(rectangles, key=lambda item: (item[1], item[0], item[3], item[2]))
        merged: List[List[float]] = []
        for rectangle in ordered:
            if not merged:
                merged.append(rectangle[:])
                continue

            previous = merged[-1]
            if self._rectangles_can_merge(previous, rectangle):
                merged[-1] = [
                    min(previous[0], rectangle[0]),
                    min(previous[1], rectangle[1]),
                    max(previous[2], rectangle[2]),
                    max(previous[3], rectangle[3]),
                ]
                continue
            merged.append(rectangle[:])
        return merged

    def _rectangles_can_merge(self, left_rect: List[float], right_rect: List[float]) -> bool:
        left_height = max(left_rect[3] - left_rect[1], 0.001)
        right_height = max(right_rect[3] - right_rect[1], 0.001)
        same_band = (
            min(left_rect[3], right_rect[3]) - max(left_rect[1], right_rect[1])
        ) >= min(left_height, right_height) * 0.45
        horizontal_gap = right_rect[0] - left_rect[2]
        return same_band and horizontal_gap <= max(left_height, right_height) * 0.8

    def _compact_for_redaction(self, text: str) -> str:
        return re.sub(r"\s+", "", text or "")

    def _coerce_relative_bbox(self, bbox: Any) -> List[float] | None:
        if not isinstance(bbox, list) or len(bbox) != 4:
            return None
        try:
            numeric = [float(value) for value in bbox]
        except (TypeError, ValueError):
            return None
        left, top, right, bottom = [min(max(value, 0.0), 1.0) for value in numeric]
        if right <= left or bottom <= top:
            return None
        return [left, top, right, bottom]

    def _extract_pdf_ocr_page_count(self, source_metadata: Dict | None) -> int:
        if not isinstance(source_metadata, dict):
            return 0
        try:
            return int(source_metadata.get("ocr_pages", 0) or 0)
        except (TypeError, ValueError):
            return 0

    def _extract_pdf_ocr_timeout_page_count(self, source_metadata: Dict | None) -> int:
        if not isinstance(source_metadata, dict):
            return 0
        try:
            return int(source_metadata.get("ocr_timeout_pages", 0) or 0)
        except (TypeError, ValueError):
            return 0

    def _extract_pdf_ocr_failed_page_count(self, source_metadata: Dict | None) -> int:
        if not isinstance(source_metadata, dict):
            return 0
        try:
            return int(source_metadata.get("ocr_failed_pages", 0) or 0)
        except (TypeError, ValueError):
            return 0

    def _build_export_replacement_entries(
        self,
        entities: Iterable[Dict],
        operator_config: Dict | None,
    ) -> List[Dict[str, Any]]:
        engine = get_engine()
        config = copy.deepcopy(engine.pipeline_manager._get_default_operator_config())
        if operator_config:
            config.update(copy.deepcopy(operator_config))

        replacement_map: Dict[str, Dict[str, Any]] = {}
        indexed_entities = [
            (index, entity)
            for index, entity in enumerate(entities)
            if isinstance(entity, dict) and entity.get("text")
        ]
        sorted_entities = sorted(
            indexed_entities,
            key=lambda item: len(str(item[1]["text"])),
            reverse=True,
        )

        for input_order, entity in sorted_entities:
            entity_type = str(entity.get("type") or "").strip()
            source_text = str(entity["text"]).strip()
            if not entity_type or not source_text:
                continue
            if entity.get("replacement_method") == "preserve":
                continue
            if (
                entity_type in engine.pipeline_manager.contextual_desensitizer.PRESERVE_TYPES
                and not entity.get("replacement")
                and not (operator_config and entity_type in operator_config)
            ):
                continue

            replacement = self._coerce_replacement_text(entity.get("replacement"))
            if not replacement:
                entity_rule = config.get(entity_type, config.get("default", {}))
                operator_name = entity_rule.get("operator", "mask")
                params = entity_rule.get("params", {})
                operator = engine.operator_registry.get_operator(operator_name)
                if operator is None:
                    operator = engine.operator_registry.get_operator("mask")

                replacement = self._coerce_replacement_text(
                    operator.operate(source_text, 0, len(source_text), entity_type, **params)
                )
            if not replacement or replacement == source_text:
                continue

            candidate = {
                "source_text": source_text,
                "replacement": replacement,
                "entity_type": entity_type,
                "sort_start": self._coerce_int(entity.get("start"), 10**12),
                "input_order": input_order,
            }
            existing = replacement_map.get(source_text)
            if existing is not None:
                if existing["replacement"] != replacement:
                    preferred = self._prefer_export_entity(
                        existing["entity_type"],
                        entity_type,
                    )
                    if preferred != existing["entity_type"]:
                        replacement_map[source_text] = candidate
                    logger.warning(
                        "Conflicting replacements detected for '%s'; using %s instead of %s.",
                        source_text,
                        replacement_map[source_text]["entity_type"],
                        existing["entity_type"],
                    )
                else:
                    existing["sort_start"] = min(existing["sort_start"], candidate["sort_start"])
                    existing["input_order"] = min(existing["input_order"], candidate["input_order"])
                continue

            replacement_map[source_text] = candidate

        return list(replacement_map.values())

    def _build_final_rewrite_entries(
        self,
        coverage_first_final_export: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        if not isinstance(coverage_first_final_export, dict) or not coverage_first_final_export.get("enabled"):
            return []
        entries: List[Dict[str, Any]] = []
        for input_order, raw_entry in enumerate(coverage_first_final_export.get("rewrite_entries") or []):
            if not isinstance(raw_entry, dict):
                continue
            if str(raw_entry.get("verification_status") or "").strip() == "blocked":
                continue
            source_text = str(raw_entry.get("source_text") or "").strip()
            replacement = self._coerce_replacement_text(raw_entry.get("replacement"))
            if not source_text or not replacement or source_text == replacement:
                continue
            entry = dict(raw_entry)
            entry["source_text"] = source_text
            entry["replacement"] = replacement
            entry["entity_type"] = str(raw_entry.get("entity_type") or "").strip()
            entry["sort_start"] = self._coerce_int(raw_entry.get("start"), 10**12)
            entry["input_order"] = self._coerce_int(raw_entry.get("input_order"), input_order)
            entries.append(entry)
        return sorted(
            entries,
            key=lambda item: (
                self._coerce_int(item.get("sort_start"), 10**12),
                -len(str(item.get("source_text") or "")),
                self._coerce_int(item.get("input_order"), 10**12),
            ),
        )

    def _build_replacements(
        self,
        entities: Iterable[Dict],
        operator_config: Dict | None,
    ) -> List[Tuple[str, str]]:
        replacement_entries = self._build_export_replacement_entries(entities, operator_config)
        return self._build_replacements_from_entries(replacement_entries)

    def _build_replacements_from_entries(
        self,
        replacement_entries: Iterable[Dict[str, Any]],
    ) -> List[Tuple[str, str]]:
        ordered_entries = sorted(
            (
                entry
                for entry in replacement_entries
                if entry.get("source_text") and entry.get("replacement")
            ),
            key=lambda item: len(str(item["source_text"])),
            reverse=True,
        )
        return [
            (str(entry["source_text"]), str(entry["replacement"]))
            for entry in ordered_entries
        ]

    def _export_mapping_docx(
        self,
        *,
        task_id: str,
        original_filename: str,
        entities: List[Dict[str, Any]],
    ) -> Dict[str, str | bool | None]:
        output_path = Path(settings.OUTPUT_DIR) / f"{task_id}_mapping.docx"
        document = Document()
        self._remove_default_paragraph(document)
        self._append_mapping_directory_content(document, entities)
        document.save(output_path)
        ensure_private_file(output_path)

        return {
            "mapping_output_path": str(output_path),
            "mapping_download_name": f"{Path(original_filename).stem}_mapping.docx",
            "mapping_file_type": "docx",
            "mapping_media_type": DOCX_MIME_TYPE,
        }

    def _coerce_replacement_text(self, value: Any) -> str:
        if value is None:
            return ""

        replacement = str(value).strip()
        if replacement.lower() == "none":
            return ""
        return replacement

    def _prefer_export_entity(self, existing_type: str, candidate_type: str) -> str:
        existing_priority = self.EXPORT_TEXT_PRIORITY.get(existing_type, 99)
        candidate_priority = self.EXPORT_TEXT_PRIORITY.get(candidate_type, 99)
        if candidate_priority < existing_priority:
            return candidate_type
        return existing_type

    def _append_mapping_directory_content(
        self,
        document: Document,
        entities: List[Dict[str, Any]],
    ) -> None:
        mapping_rows = self._build_frontend_mapping_rows(entities)

        title = document.add_paragraph()
        title.paragraph_format.space_before = Pt(0)
        title.paragraph_format.space_after = Pt(8)
        title.paragraph_format.line_spacing = 1.0
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.add_run(DOCX_MAPPING_TABLE_TITLE)
        title_run.bold = True
        title_run.font.size = Pt(14)
        self._set_mapping_run_font(title_run)

        if not mapping_rows:
            note = document.add_paragraph()
            note.paragraph_format.space_before = Pt(0)
            note.paragraph_format.space_after = Pt(0)
            note.paragraph_format.line_spacing = 1.15
            note_run = note.add_run("当前没有可展示的替换对照")
            note_run.font.size = Pt(10.5)
            self._set_mapping_run_font(note_run)
            return

        table = document.add_table(rows=1, cols=4)
        if self._docx_style_exists(document, "Table Grid"):
            table.style = "Table Grid"

        header_row = table.rows[0]
        self._populate_mapping_table_row(
            header_row,
            values=["类型", "原文", "替换后", "上下文"],
            is_header=True,
        )

        for row_data in mapping_rows:
            row = table.add_row()
            self._populate_mapping_table_row(
                row,
                values=[
                    row_data["type_label"],
                    row_data["source_text"],
                    row_data["replacement"],
                    row_data["context_display"],
                ],
                is_header=False,
            )

        table_width_dxa = self._section_content_width_dxa(document.sections[-1])
        column_widths_dxa = self._column_widths_from_weights(
            [1.25, 2.65, 2.85, 2.25],
            total_width_dxa=table_width_dxa,
        )
        self._apply_exact_table_geometry(
            table,
            column_widths_dxa,
            table_width_dxa=table_width_dxa,
            indent_dxa=DEFAULT_DOCX_TABLE_INDENT_DXA,
            cell_margins_dxa=DEFAULT_DOCX_TABLE_CELL_MARGINS_DXA,
        )
        self._set_table_borders(table)
        self._set_table_header_repeat(header_row)

    def _docx_style_exists(self, document: Document, style_name: str) -> bool:
        try:
            document.styles[style_name]
        except KeyError:
            return False
        return True

    def _build_frontend_mapping_rows(
        self,
        entities: Iterable[Dict[str, Any]],
    ) -> List[Dict[str, str]]:
        ordered_entities = sorted(
            (
                entity
                for entity in entities
                if isinstance(entity, dict)
            ),
            key=lambda item: self._coerce_int(item.get("start"), 10**12),
        )
        unique_rows: Dict[str, Dict[str, str]] = {}
        rows: List[Dict[str, str]] = []
        for entity in ordered_entities:
            source_text = str(entity.get("text") or "").strip()
            replacement = self._coerce_replacement_text(entity.get("replacement"))
            if not source_text or not replacement:
                continue
            if replacement == source_text:
                continue
            entity_type = str(entity.get("type") or "").strip()
            unique_key = "::".join((entity_type, source_text, replacement))
            if unique_key in unique_rows:
                continue

            context_label = str(entity.get("context_label") or "").strip() or "未命名字段"
            context_role = str(entity.get("context_role") or "").strip() or "通用上下文"
            row_payload = {
                "type_label": self.EXPORT_ENTITY_TYPE_LABELS.get(entity_type, entity_type or "未分类"),
                "source_text": source_text,
                "replacement": replacement,
                "context_display": f"{context_label}\n{context_role}",
            }
            unique_rows[unique_key] = row_payload
            rows.append(row_payload)
        return rows

    def _populate_mapping_table_row(self, row, values: List[str], *, is_header: bool) -> None:
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if is_header or index == 0
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            text_value = str(value or "").strip() or "-"
            lines = text_value.splitlines() or ["-"]
            for line_index, line in enumerate(lines):
                run = paragraph.add_run(line or "-")
                run.bold = is_header
                run.font.size = Pt(10.5)
                self._set_mapping_run_font(run)
                if line_index < len(lines) - 1:
                    run.add_break()
            if is_header:
                self._set_cell_shading(cell, DOCX_MAPPING_TABLE_HEADER_FILL)

    def _set_mapping_run_font(self, run) -> None:
        run.font.name = DOCX_MAPPING_FONT_NAME
        r_pr = run._element.get_or_add_rPr()
        r_fonts = self._ensure_oxml_child(r_pr, "w:rFonts")
        for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(attribute), DOCX_MAPPING_FONT_NAME)

    def _set_table_header_repeat(self, row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        header = tr_pr.find(qn("w:tblHeader"))
        if header is None:
            header = OxmlElement("w:tblHeader")
            tr_pr.append(header)
        header.set(qn("w:val"), "true")

    def _set_table_borders(self, table) -> None:
        tbl_pr = table._tbl.tblPr
        borders = self._ensure_oxml_child(tbl_pr, "w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = self._ensure_oxml_child(borders, f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "6")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "C9D2DE")

    def _set_cell_shading(self, cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = self._ensure_oxml_child(tc_pr, "w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), fill)

    def _apply_exact_table_geometry(
        self,
        table,
        column_widths_dxa: List[int],
        *,
        table_width_dxa: int | None = None,
        indent_dxa: int | None = None,
        cell_margins_dxa: Dict[str, int] | None = None,
    ) -> None:
        widths = [int(width) for width in column_widths_dxa]
        if not widths or any(width <= 0 for width in widths):
            raise ValueError("column_widths_dxa must contain only positive widths")

        width_total = int(table_width_dxa if table_width_dxa is not None else sum(widths))
        if sum(widths) != width_total:
            raise ValueError(
                f"column widths must sum to table_width_dxa: sum={sum(widths)} width={width_total}"
            )

        cell_margins = dict(DEFAULT_DOCX_TABLE_CELL_MARGINS_DXA)
        if cell_margins_dxa:
            cell_margins.update({key: int(value) for key, value in cell_margins_dxa.items()})
        resolved_indent_dxa = (
            int(cell_margins.get("start", DEFAULT_DOCX_TABLE_INDENT_DXA))
            if indent_dxa is None
            else int(indent_dxa)
        )

        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        self._set_oxml_width(tbl_pr, "w:tblW", width_total)

        table_indent = self._ensure_oxml_child(tbl_pr, "w:tblInd")
        table_indent.set(qn("w:type"), "dxa")
        table_indent.set(qn("w:w"), str(resolved_indent_dxa))

        layout = self._ensure_oxml_child(tbl_pr, "w:tblLayout")
        layout.set(qn("w:type"), "fixed")

        self._replace_table_grid(table, widths)

        for col_index, width in enumerate(widths):
            table.columns[col_index].width = Twips(width)

        for row in table.rows:
            for col_index, cell in enumerate(row.cells):
                width = widths[col_index]
                cell.width = Twips(width)
                tc_pr = cell._tc.get_or_add_tcPr()
                self._set_oxml_width(tc_pr, "w:tcW", width)
                self._set_table_cell_margins(cell, cell_margins)

    def _replace_table_grid(self, table, column_widths_dxa: List[int]) -> None:
        tbl = table._tbl
        grid = tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in column_widths_dxa:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(int(width)))
            grid.append(grid_col)

    def _set_table_cell_margins(self, cell, margins_dxa: Dict[str, int]) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = self._ensure_oxml_child(tc_pr, "w:tcMar")
        for side in ("top", "bottom", "start", "end"):
            margin = self._ensure_oxml_child(tc_mar, f"w:{side}")
            margin.set(qn("w:w"), str(int(margins_dxa[side])))
            margin.set(qn("w:type"), "dxa")

    def _set_oxml_width(self, parent, tag: str, width_dxa: int) -> None:
        width = self._ensure_oxml_child(parent, tag)
        width.set(qn("w:type"), "dxa")
        width.set(qn("w:w"), str(int(width_dxa)))

    def _ensure_oxml_child(self, parent, tag: str):
        child = parent.find(qn(tag))
        if child is None:
            child = OxmlElement(tag)
            parent.append(child)
        return child

    def _section_content_width_dxa(self, section) -> int:
        width_dxa = (
            self._length_to_dxa(section.page_width)
            - self._length_to_dxa(section.left_margin)
            - self._length_to_dxa(section.right_margin)
        )
        return max(3600, width_dxa or DEFAULT_DOCX_TABLE_WIDTH_DXA)

    def _length_to_dxa(self, length) -> int:
        try:
            return int(round(length.twips))
        except Exception:
            return int(round(float(length)))

    def _column_widths_from_weights(
        self,
        weights: List[float],
        *,
        total_width_dxa: int,
    ) -> List[int]:
        if not weights or any(weight <= 0 for weight in weights):
            raise ValueError("weights must be positive")

        total_weight = float(sum(weights))
        widths = [int(round(total_width_dxa * (weight / total_weight))) for weight in weights]
        widths[-1] += total_width_dxa - sum(widths)
        if any(width <= 0 for width in widths):
            raise ValueError(f"invalid computed widths: {widths}")
        return widths

    def _coerce_int(self, value: Any, default: int) -> int:
        try:
            return int(value)
        except (TypeError, ValueError):
            return default

    def _build_pdf_pages(
        self,
        *,
        source_text: str | None,
        source_metadata: Dict | None,
        source_structure: Dict | None,
        replacements: List[Tuple[str, str]],
    ) -> List[Dict[str, Any]]:
        pages: List[Dict[str, Any]] = []
        if isinstance(source_structure, dict):
            raw_pages = source_structure.get("pages")
            if isinstance(raw_pages, list):
                for raw_page in raw_pages:
                    if not isinstance(raw_page, dict):
                        continue
                    page_text = str(raw_page.get("text", "") or "")
                    page_blocks = raw_page.get("blocks")
                    blocks = self._apply_replacements_to_blocks(page_blocks, replacements)
                    if not blocks:
                        blocks = self._build_blocks_from_text_fallback(
                            self._apply_text_replacements(page_text, replacements)
                        )

                    width = self._coerce_float(
                        raw_page.get("width"),
                        DEFAULT_PDF_PAGE_WIDTH_PT,
                    )
                    height = self._coerce_float(
                        raw_page.get("height"),
                        DEFAULT_PDF_PAGE_HEIGHT_PT,
                    )

                    pages.append(
                        {
                            "width": width,
                            "height": height,
                            "blocks": blocks,
                        }
                    )

        if pages:
            return pages

        fallback_text = self._apply_text_replacements(source_text or "", replacements).strip()
        if not fallback_text:
            return []

        page_count = 1
        if isinstance(source_metadata, dict):
            try:
                page_count = max(1, int(source_metadata.get("pages", 1) or 1))
            except (TypeError, ValueError):
                page_count = 1

        if page_count <= 1:
            return [
                {
                    "width": DEFAULT_PDF_PAGE_WIDTH_PT,
                    "height": DEFAULT_PDF_PAGE_HEIGHT_PT,
                    "blocks": self._build_blocks_from_text_fallback(fallback_text),
                }
            ]

        chunks = [chunk.strip() for chunk in fallback_text.split("\n\n") if chunk.strip()]
        return [
            {
                "width": DEFAULT_PDF_PAGE_WIDTH_PT,
                "height": DEFAULT_PDF_PAGE_HEIGHT_PT,
                "blocks": self._build_blocks_from_text_fallback(chunk),
            }
            for chunk in chunks
        ]

    def _apply_replacements_to_blocks(
        self,
        blocks: Any,
        replacements: List[Tuple[str, str]],
    ) -> List[Dict[str, Any]]:
        if not isinstance(blocks, list):
            return []

        normalized_blocks: List[Dict[str, Any]] = []
        for raw_block in blocks:
            if not isinstance(raw_block, dict):
                continue

            block_type = str(raw_block.get("type", "line")).strip().lower()
            if block_type == "table":
                rows = raw_block.get("rows")
                if not isinstance(rows, list):
                    continue
                normalized_rows: List[List[str]] = []
                for row in rows:
                    if not isinstance(row, list):
                        continue
                    normalized_row = [
                        self._apply_text_replacements(str(cell), replacements).strip()
                        for cell in row
                    ]
                    normalized_rows.append(normalized_row)
                normalized_blocks.append(
                    {
                        "type": "table",
                        "rows": normalized_rows,
                        "align": "left",
                    }
                )
                continue

            if block_type == "spacer":
                normalized_blocks.append(
                    {
                        "type": "spacer",
                        "count": max(
                            1,
                            int(raw_block.get("count", raw_block.get("blank_before", 1)) or 1),
                        ),
                    }
                )
                continue

            text = self._apply_text_replacements(str(raw_block.get("text", "")), replacements).strip()
            if not text:
                continue

            align = str(raw_block.get("align", "left")).strip().lower()
            if align not in {"left", "center", "right"}:
                align = "left"

            normalized_blocks.append(
                {
                    "type": block_type if block_type in {"title", "paragraph", "line"} else "line",
                    "text": text,
                    "align": align,
                    "indent_pt": self._coerce_float(raw_block.get("indent_pt"), 0.0),
                    "indent": int(raw_block.get("indent", 0) or 0),
                    "space_before_pt": self._coerce_float(raw_block.get("space_before_pt"), 0.0),
                    "blank_before": max(0, int(raw_block.get("blank_before", 0) or 0)),
                    "font_size_hint": self._coerce_float(raw_block.get("font_size_hint"), 0.0),
                }
            )

        return self._collapse_spacer_blocks(normalized_blocks)

    def _build_blocks_from_text_fallback(self, text: str) -> List[Dict[str, Any]]:
        if not text.strip():
            return []

        blocks: List[Dict[str, Any]] = []
        pending_rows: List[List[str]] = []
        for index, raw_line in enumerate(text.splitlines()):
            stripped = raw_line.strip()
            if not stripped:
                if pending_rows:
                    blocks.append({"type": "table", "rows": pending_rows})
                    pending_rows = []
                blocks.append({"type": "spacer", "count": 1})
                continue

            cells = self._split_fallback_cells(raw_line)
            if len(cells) >= 2:
                pending_rows.append(cells)
                continue

            if pending_rows:
                blocks.append({"type": "table", "rows": pending_rows})
                pending_rows = []

            block_type = "title" if index == 0 and self._looks_like_center_title(stripped) else "line"
            blocks.append(
                {
                    "type": block_type,
                    "text": stripped,
                    "align": "center" if block_type == "title" else "left",
                    "indent_pt": 0.0,
                    "space_before_pt": 0.0,
                }
            )

        if pending_rows:
            blocks.append({"type": "table", "rows": pending_rows})

        return self._collapse_spacer_blocks(blocks)

    def _split_fallback_cells(self, line: str) -> List[str]:
        if "\t" in line:
            parts = [part.strip() for part in line.split("\t")]
        else:
            parts = [part.strip() for part in re.split(r"\s{3,}", line)]
        return [part for part in parts if part]

    def _collapse_spacer_blocks(self, blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        collapsed: List[Dict[str, Any]] = []
        for block in blocks:
            if str(block.get("type", "")).lower() == "spacer":
                count = max(1, int(block.get("count", 1) or 1))
                if collapsed and str(collapsed[-1].get("type", "")).lower() == "spacer":
                    collapsed[-1]["count"] = min(
                        3,
                        int(collapsed[-1].get("count", 1) or 1) + count,
                    )
                else:
                    collapsed.append({"type": "spacer", "count": min(3, count)})
                continue

            collapsed.append(block)

        return collapsed

    def _configure_pdf_section(self, document: Document, page: Dict[str, Any], page_index: int) -> None:
        if page_index == 0:
            section = document.sections[0]
        else:
            section = document.add_section(WD_SECTION.NEW_PAGE)

        section.page_width = Pt(self._coerce_float(page.get("width"), DEFAULT_PDF_PAGE_WIDTH_PT))
        section.page_height = Pt(self._coerce_float(page.get("height"), DEFAULT_PDF_PAGE_HEIGHT_PT))
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)
        section.top_margin = Pt(32)
        section.bottom_margin = Pt(32)

    def _append_structured_pdf_page(self, document: Document, page: Dict[str, Any]) -> None:
        blocks = page.get("blocks") if isinstance(page, dict) else None
        if not isinstance(blocks, list) or not blocks:
            document.add_paragraph("")
            return

        for index, block in enumerate(blocks):
            block_type = str(block.get("type", "line")).strip().lower()
            if block_type == "spacer":
                count = max(1, int(block.get("count", 1) or 1))
                for _ in range(count):
                    document.add_paragraph("")
                continue

            if block_type == "table":
                self._append_table_block(document, block)
                continue

            self._append_text_block(document, block, index)

    def _append_text_block(
        self,
        document: Document,
        block: Dict[str, Any],
        block_index: int,
    ) -> None:
        text = str(block.get("text", "")).strip()
        if not text:
            return

        blank_before = max(0, int(block.get("blank_before", 0) or 0))
        for _ in range(blank_before):
            document.add_paragraph("")

        paragraph = document.add_paragraph()
        paragraph.alignment = self._resolve_alignment(block.get("align"))
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(
            min(18.0, self._coerce_float(block.get("space_before_pt"), 0.0))
        )
        paragraph_format.line_spacing = 1.0

        indent_pt = self._coerce_float(block.get("indent_pt"), 0.0)
        if indent_pt <= 0:
            indent_level = max(0, int(block.get("indent", 0) or 0))
            indent_pt = indent_level * 18.0
        if indent_pt > 0:
            paragraph_format.left_indent = Pt(min(indent_pt, 144.0))

        font_size = self._resolve_font_size(block)
        lines = text.splitlines()
        for line_index, line in enumerate(lines):
            run = paragraph.add_run(line)
            run.font.size = Pt(font_size)
            if str(block.get("type", "")).lower() == "title":
                run.bold = True
            if line_index < len(lines) - 1:
                paragraph.add_run().add_break(WD_BREAK.LINE)

        if block_index == 0 and str(block.get("type", "")).lower() == "title":
            paragraph_format.space_after = Pt(6)

    def _append_table_block(self, document: Document, block: Dict[str, Any]) -> None:
        rows = block.get("rows")
        if not isinstance(rows, list) or not rows:
            return

        max_cols = max(len(row) for row in rows if isinstance(row, list)) if rows else 0
        if max_cols <= 0:
            return

        table = document.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"
        table.autofit = True

        for row_index, row in enumerate(rows):
            if not isinstance(row, list):
                continue
            for col_index in range(max_cols):
                text = str(row[col_index]).strip() if col_index < len(row) else ""
                cell = table.cell(row_index, col_index)
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                run = paragraph.add_run(text)
                run.font.size = Pt(10.5)

    def _resolve_alignment(self, align: Any) -> WD_ALIGN_PARAGRAPH:
        value = str(align or "left").strip().lower()
        if value == "center":
            return WD_ALIGN_PARAGRAPH.CENTER
        if value == "right":
            return WD_ALIGN_PARAGRAPH.RIGHT
        return WD_ALIGN_PARAGRAPH.LEFT

    def _resolve_font_size(self, block: Dict[str, Any]) -> float:
        block_type = str(block.get("type", "line")).strip().lower()
        font_size_hint = self._coerce_float(block.get("font_size_hint"), 0.0)
        if font_size_hint > 0:
            font_size_hint = max(9.0, min(font_size_hint, 16.0))

        if block_type == "title":
            return max(14.0, font_size_hint or 14.0)
        if block_type == "paragraph":
            return font_size_hint or 11.0
        return font_size_hint or 10.5

    def _looks_like_center_title(self, text: str) -> bool:
        if len(text) > 40:
            return False
        return not any(token in text for token in ["：", ":", "。", "；", ";", "，", ","])

    def _remove_default_paragraph(self, document: Document) -> None:
        if len(document.paragraphs) != 1:
            return
        paragraph = document.paragraphs[0]
        if paragraph.text.strip():
            return
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    def _apply_text_replacements(
        self,
        text: str,
        replacements: List[Tuple[str, str]],
    ) -> str:
        updated = text
        for original, replacement in replacements:
            updated = updated.replace(original, replacement)
        return mask_default_numeric_text(updated)

    def _coerce_float(self, value: Any, default: float) -> float:
        try:
            numeric = float(value)
        except (TypeError, ValueError):
            return default
        if numeric <= 0:
            return default
        return numeric
