"""Local high-quality PDF normalization before desensitization.

This service is intentionally worker-oriented. It may import OpenCV, PDF renderers,
or call local OCR runtimes, but it should only be used from a one-shot worker so
the API process does not retain heavyweight OCR state.
"""

from __future__ import annotations

import base64
import difflib
import gc
import hashlib
import json
import logging
import os
import re
import shlex
import shutil
import subprocess
import sys
import tempfile
import time
import uuid
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from statistics import mean
from typing import Any, Dict, Iterable, List, Optional

import pdfplumber
import pypdfium2 as pdfium
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt
from PIL import Image, ImageOps

from app.core.config import SOURCE_BACKEND_ROOT, settings
from app.core.runtime_security import ensure_private_directory, ensure_private_file
from app.services.ocr_model_profiles import paddleocr_vl_profile, ppstructurev3_profile
from app.services.pdf_normalization_service import PdfNormalizationService, RapidOCROnNXService

logger = logging.getLogger(__name__)
LOCAL_PDF_FRONTLINE_PIPELINE_VERSION = "quality_first_adaptive_v1"


def _coerce_bool_setting(value: Any, *, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return default
    normalized = str(value).strip().lower()
    if not normalized:
        return default
    if normalized in {"1", "true", "yes", "on", "force", "enabled"}:
        return True
    if normalized in {"0", "false", "no", "off", "disabled", "none"}:
        return False
    return default


def _build_wps_region_prompt(*, field_type: str = "", text_hint: str = "") -> str:
    field_hint = f"字段类型：{field_type}\n" if field_type else ""
    candidate_hint = f"WPS 候选原文：{text_hint[:700]}\n" if text_hint else ""
    return (
        "你是 WPS PDF 转 DOCX 忠实度审查的局部 OCR 复核员。截图是从原 PDF 裁剪出的局部区域，WPS 候选原文只是待核验对象，不是真相。\n"
        "工作步骤：\n"
        "1. 先忽略 WPS 候选原文，只抄录截图中确定可见的文字，写入 visible_text。\n"
        "2. 再比较 visible_text 与 WPS 候选原文是否互相支持。\n"
        "3. 最后按下面定义给 verdict。\n"
        "抄录规则：不要补全、不要纠错、不要改写、不要总结；看不清就少写，不要猜；不要输出 Markdown 或多余字段。\n"
        "支持规则：空格、换行、全半角、普通标点、冒号样式差异不算错误；繁简体或 OCR 常见近形字只有在关键字段中改变含义时才算冲突。\n"
        "verdict 定义：\n"
        "- match：截图中文字能支持 WPS 候选原文，只有排版/标点/空格差异。\n"
        "- partial：截图只覆盖候选原文的一部分，未发现关键字段直接冲突。\n"
        "- mismatch：截图中可见关键字段与候选原文直接冲突，且不是裁剪缺失或看不清。\n"
        "- unreadable：截图过糊、过小、遮挡、空白或只能看到无意义碎片。\n"
        "candidate_supported：只有 match 时为 true；partial、mismatch、unreadable 都为 false。\n"
        "关键字段：金额、日期、证件号、统一社会信用代码、银行账号、案号、合同编号、公司名、人名、地址。关键字段不同才把对应值写入 mismatch_values。\n"
        "看不清：candidate_supported=false，verdict=unreadable 或 partial，visible_text 只写确定字符，mismatch_values 为空。\n"
        f"{field_hint}{candidate_hint}"
        "candidate_supported 必须是 JSON 布尔值 true 或 false，不能输出字符串。\n"
        "必须只输出 JSON，字段固定为："
        "{\"visible_text\":\"截图中确定可见文字\","
        "\"candidate_supported\":false,"
        "\"verdict\":\"match|partial|mismatch|unreadable\","
        "\"mismatch_values\":[\"只列直接冲突的关键字段值\"],"
        "\"confidence\":0.0,"
        "\"reason\":\"20字内中文原因\"}"
    )


def _auto_bool_setting(value: Any, *, auto_default: bool = False, default: bool = False) -> bool:
    if isinstance(value, str) and value.strip().lower() == "auto":
        return auto_default
    return _coerce_bool_setting(value, default=default)


class LocalHighQualityPdfNormalizationError(RuntimeError):
    """Stable-code PDF frontline error for UI and task metadata."""

    def __init__(self, code: str, message: str, *, metadata: Optional[Dict[str, Any]] = None) -> None:
        super().__init__(f"{code}: {message}")
        self.code = code
        self.metadata = metadata or {}


@dataclass
class LocalPdfPageResult:
    page_number: int
    width: float
    height: float
    text: str
    source: str
    quality: str
    lines: List[Dict[str, Any]]
    blocks: List[Dict[str, Any]]
    confidence: float = 0.0
    warnings: Optional[List[str]] = None
    engine: str = "glm_ocr_ollama"
    used_vl_fallback: bool = False

    def to_public_dict(self) -> Dict[str, Any]:
        return {
            "page_number": self.page_number,
            "width": self.width,
            "height": self.height,
            "text": self.text,
            "source": self.source,
            "ocr_quality": self.quality,
            "ocr_confidence": self.confidence,
            "engine": self.engine,
            "used_vl_fallback": self.used_vl_fallback,
            "warnings": sorted(set(self.warnings or [])),
            "lines": self.lines,
            "blocks": self.blocks,
        }

    @classmethod
    def from_public_dict(cls, payload: Dict[str, Any]) -> "LocalPdfPageResult":
        return cls(
            page_number=int(payload.get("page_number") or 0),
            width=float(payload.get("width") or 595.0),
            height=float(payload.get("height") or 842.0),
            text=str(payload.get("text") or ""),
            source=str(payload.get("source") or "glm_ocr"),
            quality=str(payload.get("ocr_quality") or payload.get("quality") or "failed"),
            confidence=float(payload.get("ocr_confidence") or payload.get("confidence") or 0.0),
            lines=list(payload.get("lines") or []),
            blocks=list(payload.get("blocks") or []),
            warnings=list(payload.get("warnings") or []),
            engine=str(payload.get("engine") or "glm_ocr_ollama"),
            used_vl_fallback=bool(payload.get("used_vl_fallback")),
        )


class GlmOcrOllamaClient:
    """Small Ollama adapter for the local GLM-OCR model."""

    def __init__(self, *, base_url: str, model: str, timeout: int, max_tokens: int) -> None:
        self.base_url = str(base_url or "").rstrip("/")
        self.model = str(model or "").strip()
        self.timeout = max(1, int(timeout or 120))
        self.max_tokens = max(256, min(int(max_tokens or 8192), 8192))
        self.available = bool(self.base_url and self.model and self._model_installed())

    def _model_installed(self) -> bool:
        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=2)
            if response.status_code != 200:
                return False
            payload = response.json()
        except Exception:
            return False
        names = {
            str(item.get("name") or item.get("model") or "").strip()
            for item in payload.get("models", [])
            if isinstance(item, dict)
        }
        return self.model in names

    def recognize(self, image_bytes: bytes, *, page_number: int, total_pages: int) -> Dict[str, Any]:
        return self.recognize_text_official(image_bytes, page_number=page_number, total_pages=total_pages)

    def recognize_text_official(self, image_bytes: bytes, *, page_number: int, total_pages: int) -> Dict[str, Any]:
        return self._recognize_with_prompt(
            image_bytes,
            prompt="Text Recognition:",
            page_number=page_number,
            total_pages=total_pages,
            max_tokens=self.max_tokens,
            task="text",
        )

    def recognize_table_official(self, image_bytes: bytes, *, page_number: int, total_pages: int) -> Dict[str, Any]:
        return self._recognize_with_prompt(
            image_bytes,
            prompt="Table Recognition:",
            page_number=page_number,
            total_pages=total_pages,
            max_tokens=max(self.max_tokens, 8192),
            task="table",
        )

    def recognize_formula_official(self, image_bytes: bytes, *, page_number: int, total_pages: int) -> Dict[str, Any]:
        return self._recognize_with_prompt(
            image_bytes,
            prompt="Formula Recognition:",
            page_number=page_number,
            total_pages=total_pages,
            max_tokens=max(self.max_tokens, 8192),
            task="formula",
        )

    def recognize_region_text_strict(
        self,
        image_bytes: bytes,
        *,
        page_number: int,
        total_pages: int,
        field_type: str = "",
        text_hint: str = "",
    ) -> Dict[str, Any]:
        prompt = _build_wps_region_prompt(field_type=field_type, text_hint=text_hint)
        raw = self._raw_generate(image_bytes, prompt=prompt, max_tokens=min(self.max_tokens, 1024))
        parsed = self._extract_json(raw)
        if isinstance(parsed, dict):
            visible_text = str(parsed.get("visible_text") or parsed.get("text") or "").strip()
            if len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", visible_text)) < 1 or self._looks_like_region_prompt_leakage(visible_text):
                visible_text = ""
            confidence = 0.0
            try:
                confidence = float(parsed.get("confidence") or 0.0)
            except Exception:
                confidence = 0.0
            verdict = str(parsed.get("verdict") or "").strip().lower()
            if verdict not in {"match", "partial", "mismatch", "unreadable"}:
                verdict = "partial" if visible_text else "unreadable"
            return {
                "text": visible_text,
                "quality": "medium" if visible_text else "failed",
                "layout": "region_text",
                "blocks": [],
                "lines": [{"text": visible_text, "confidence": max(0.0, min(confidence or 0.82, 1.0)), "source": "glm_ocr_region_strict"}] if visible_text else [],
                "warnings": [] if visible_text else ["glm_ocr_empty_response"],
                "task": "region_text",
                "raw_response": raw[:2000],
                "confidence": max(0.0, min(confidence or 0.0, 1.0)),
                "is_readable": _coerce_bool_setting(parsed.get("is_readable"), default=bool(visible_text)) and bool(visible_text),
                "candidate_supported": _coerce_bool_setting(parsed.get("candidate_supported"), default=False),
                "verdict": verdict,
                "mismatch_values": [str(item).strip() for item in parsed.get("mismatch_values") or [] if str(item).strip()][:5],
                "reason": str(parsed.get("reason") or ""),
            }
        parsed_text = self._parse_official_response(raw, task="region_text")
        parsed_text.setdefault("raw_response", raw[:2000])
        if self._looks_like_region_prompt_leakage(str(parsed_text.get("text") or "")):
            return {
                "text": "",
                "quality": "failed",
                "layout": "region_text",
                "blocks": [],
                "lines": [],
                "warnings": ["glm_ocr_prompt_leakage"],
                "task": "region_text",
                "raw_response": raw[:2000],
                "is_readable": False,
                "reason": "prompt_leakage_or_repeated_json",
            }
        return parsed_text

    def _looks_like_region_prompt_leakage(self, text: str) -> bool:
        value = str(text or "")
        if not value.strip():
            return False
        bad_markers = ("必须只输出 JSON", "visible_text", "is_readable", "confidence", "不要输出 Markdown")
        if any(marker in value for marker in bad_markers):
            return True
        if value.count("{") >= 2 or value.count("}") >= 2:
            return True
        return False

    def _recognize_with_prompt(
        self,
        image_bytes: bytes,
        *,
        prompt: str,
        page_number: int,
        total_pages: int,
        max_tokens: int,
        task: str,
    ) -> Dict[str, Any]:
        if not image_bytes:
            return {"text": "", "quality": "failed", "warnings": ["empty_image_input"], "lines": [], "blocks": []}
        raw_text = self._raw_generate(image_bytes, prompt=prompt, max_tokens=max_tokens)
        parsed = self._parse_official_response(raw_text, task=task)
        parsed.setdefault("raw_response", raw_text[:2000])
        return parsed

    def recognize_custom_prompt(
        self,
        image_bytes: bytes,
        *,
        prompt: str,
        page_number: int,
        total_pages: int,
        max_tokens: Optional[int] = None,
        task: str = "text",
    ) -> Dict[str, Any]:
        return self._recognize_with_prompt(
            image_bytes,
            prompt=prompt,
            page_number=page_number,
            total_pages=total_pages,
            max_tokens=max_tokens or self.max_tokens,
            task=task,
        )

    def _raw_generate(self, image_bytes: bytes, *, prompt: str, max_tokens: int) -> str:
        payload = {
            "model": self.model,
            "prompt": prompt,
            "images": [base64.b64encode(image_bytes).decode("ascii")],
            "stream": False,
            "keep_alive": "0s",
            "options": {
                "temperature": 0.0,
                "top_p": 0.00001,
                "top_k": 1,
                "repeat_penalty": 1.1,
                "num_predict": max(256, min(int(max_tokens or self.max_tokens), 8192)),
            },
        }
        response = requests.post(f"{self.base_url}/api/generate", json=payload, timeout=self.timeout)
        if response.status_code != 200:
            detail = str(getattr(response, "text", "") or "").strip().replace("\n", " ")[:500]
            raise RuntimeError(f"glm_ocr_http_{response.status_code}: {detail}")
        return str((response.json() or {}).get("response") or "").strip()

    def _prompt(self, *, page_number: int, total_pages: int) -> str:
        return "Text Recognition:"

    def _parse_response(self, raw_text: str) -> Dict[str, Any]:
        parsed = self._extract_json(raw_text)
        if isinstance(parsed, dict):
            return parsed
        return self._parse_official_response(raw_text, task="text")

    def _parse_official_response(self, raw_text: str, *, task: str) -> Dict[str, Any]:
        cleaned = self._strip_markdown_fences(raw_text)
        lines = [
            {"text": line.strip(), "confidence": 0.9, "source": f"glm_ocr_official_{task}"}
            for line in cleaned.splitlines()
            if line.strip()
        ]
        readable_chars = len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", cleaned))
        abnormal_chars = len(re.findall(r"[�□■◆◇●○]", cleaned))
        abnormal_ratio = abnormal_chars / max(len(cleaned.strip()), 1)
        if not cleaned.strip():
            quality = "failed"
        elif readable_chars >= 80 and abnormal_ratio <= 0.04:
            quality = "high"
        elif readable_chars >= 20 and abnormal_ratio <= 0.08:
            quality = "medium"
        else:
            quality = "low"
        return {
            "text": cleaned,
            "quality": quality,
            "layout": "plain_text",
            "blocks": [],
            "lines": lines,
            "warnings": [] if cleaned.strip() else ["glm_ocr_empty_response"],
            "task": task,
        }

    @staticmethod
    def _strip_markdown_fences(text: str) -> str:
        stripped = str(text or "").strip()
        stripped = re.sub(r"^```(?:json|markdown|md)?\s*", "", stripped, flags=re.I)
        stripped = re.sub(r"\s*```$", "", stripped)
        return stripped.strip()

    def _extract_json(self, text: str) -> Optional[Dict[str, Any]]:
        candidate = self._strip_markdown_fences(text)
        for value in (candidate, self._balanced_json(candidate)):
            if not value:
                continue
            try:
                parsed = json.loads(value)
            except Exception:
                continue
            if isinstance(parsed, dict):
                return parsed
        return None

    @staticmethod
    def _balanced_json(text: str) -> str:
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1 or end <= start:
            return ""
        return text[start : end + 1]


class PaddleOcrVlFallbackClient(GlmOcrOllamaClient):
    """Optional fallback adapter.

    v1 only executes this fallback when the configured model is an installed
    Ollama tag. HF/local model execution is deliberately treated as missing so
    runtime never downloads weights implicitly.
    """

    def __init__(
        self,
        *,
        base_url: str,
        model: str,
        timeout: int,
        max_tokens: int,
        api_mode: str = "ollama_generate",
        model_dir: Optional[str] = None,
        python_path: Optional[str] = None,
        command: Optional[str] = None,
    ) -> None:
        self.api_mode = str(api_mode or "ollama_generate").strip().lower()
        self.command = str(command or "").strip()
        self.local_model_dir = self._resolve_local_model_dir(model_dir)
        self.python_path = self._resolve_python_path(python_path)
        if self.api_mode == "paddleocr_python":
            self.base_url = str(base_url or "").rstrip("/")
            self.model = str(model or "").strip()
            self.timeout = max(1, int(timeout or 120))
            self.max_tokens = max(256, min(int(max_tokens or 8192), 8192))
            self.model_configured_as_local_path = False
            self.available = bool(self.local_model_dir and self.local_model_dir.exists() and self.python_path and Path(self.python_path).exists())
            return
        if self.command:
            self.base_url = str(base_url or "").rstrip("/")
            self.model = str(model or "").strip()
            self.timeout = max(1, int(timeout or 120))
            self.max_tokens = max(256, min(int(max_tokens or 8192), 8192))
            self.model_configured_as_local_path = False
            self.available = self._command_available()
            return
        self.model_configured_as_local_path = "/" in str(model or "") and not Path(str(model)).exists()
        super().__init__(base_url=base_url, model=model, timeout=timeout, max_tokens=max_tokens)
        if self.model_configured_as_local_path:
            self.available = False

    def recognize(
        self,
        image_bytes: bytes,
        *,
        page_number: int,
        total_pages: int,
        use_doc_unwarping: Optional[bool] = None,
    ) -> Dict[str, Any]:
        if self.api_mode == "paddleocr_python":
            return self._recognize_with_local_paddleocr_vl(
                image_bytes,
                page_number=page_number,
                total_pages=total_pages,
                use_doc_unwarping=use_doc_unwarping,
            )
        if not self.command:
            return super().recognize(image_bytes, page_number=page_number, total_pages=total_pages)
        image_path = self._write_temp_image(image_bytes)
        try:
            formatted = self.command.format(
                image=str(image_path),
                page=page_number,
                total=total_pages,
                model=self.model,
            )
            result = subprocess.run(
                shlex.split(formatted),
                capture_output=True,
                text=True,
                timeout=self.timeout,
                check=False,
            )
            if result.returncode != 0:
                raise RuntimeError(f"paddleocr_vl_command_failed:{result.stderr[-800:]}")
            raw_text = str(result.stdout or "").strip()
            parsed = self._parse_response(raw_text)
            parsed.setdefault("raw_response", raw_text[:2000])
            return parsed
        finally:
            try:
                image_path.unlink(missing_ok=True)
            except Exception:
                pass

    def _recognize_with_local_paddleocr_vl(
        self,
        image_bytes: bytes,
        *,
        page_number: int,
        total_pages: int,
        use_doc_unwarping: Optional[bool] = None,
    ) -> Dict[str, Any]:
        if not self.local_model_dir:
            raise RuntimeError("paddleocr_vl_model_dir_missing")
        image_path = self._write_temp_image(image_bytes)
        input_path = self._temp_json_file(image_path.parent, "paddleocr-vl-in-")
        output_path = self._temp_json_file(image_path.parent, "paddleocr-vl-out-")
        try:
            payload = {
                **paddleocr_vl_profile(
                    mode="photo_warped" if bool(use_doc_unwarping) else "hard_document",
                    server_url=str(getattr(settings, "LOCAL_PDF_VL_REC_SERVER_URL", "http://localhost:8111/") or ""),
                    threads=max(1, int(getattr(settings, "LOCAL_PDF_VL_THREADS", 2) or 2)),
                ),
                "image_path": str(image_path),
                "model_dir": str(self.local_model_dir),
                "page_number": page_number,
                "total_pages": total_pages,
                "vl_rec_backend": str(getattr(settings, "LOCAL_PDF_VL_REC_BACKEND", "auto") or "auto"),
                "vl_rec_server_url": str(getattr(settings, "LOCAL_PDF_VL_REC_SERVER_URL", "http://localhost:8111/") or ""),
                "vl_rec_api_model_name": str(getattr(settings, "LOCAL_PDF_VL_MODEL", "") or self.model),
                "use_layout_detection": bool(getattr(settings, "LOCAL_PDF_VL_LAYOUT_DETECTION", True)),
                "use_doc_orientation_classify": bool(getattr(settings, "LOCAL_PDF_VL_DOC_ORIENTATION", True)),
                "use_doc_unwarping": bool(use_doc_unwarping)
                if use_doc_unwarping is not None
                else _auto_bool_setting(getattr(settings, "LOCAL_PDF_VL_DOC_UNWARPING", False), auto_default=False),
                "use_ocr_for_image_block": bool(getattr(settings, "LOCAL_PDF_VL_OCR_FOR_IMAGE_BLOCK", True)),
                "use_seal_recognition": bool(getattr(settings, "LOCAL_PDF_VL_SEAL_RECOGNITION", False)),
                "use_chart_recognition": bool(getattr(settings, "LOCAL_PDF_VL_CHART_RECOGNITION", False)),
                "min_pixels": int(getattr(settings, "LOCAL_PDF_VL_MIN_PIXELS", 112896) or 112896),
                "max_pixels": int(getattr(settings, "LOCAL_PDF_VL_MAX_PIXELS", 1605632) or 1605632),
                "max_new_tokens": min(int(self.max_tokens or 8192), 8192),
                "threads": max(1, int(getattr(settings, "LOCAL_PDF_VL_THREADS", 2) or 2)),
            }
            input_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            ensure_private_file(input_path)
            env = os.environ.copy()
            threads = str(payload["threads"])
            env.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
            env["OMP_NUM_THREADS"] = threads
            env["OPENBLAS_NUM_THREADS"] = threads
            env["MKL_NUM_THREADS"] = threads
            env["VECLIB_MAXIMUM_THREADS"] = threads
            existing_pythonpath = str(env.get("PYTHONPATH") or "").strip()
            env["PYTHONPATH"] = (
                str(SOURCE_BACKEND_ROOT)
                if not existing_pythonpath
                else f"{SOURCE_BACKEND_ROOT}{os.pathsep}{existing_pythonpath}"
            )
            result = subprocess.run(
                [
                    self.python_path,
                    "-m",
                    "app.workers.local_paddleocr_vl_sidecar",
                    str(input_path),
                    str(output_path),
                ],
                cwd=str(SOURCE_BACKEND_ROOT),
                env=env,
                capture_output=True,
                text=True,
                timeout=self.timeout,
                check=False,
            )
            output_text = output_path.read_text(encoding="utf-8") if output_path.exists() else ""
            parsed = json.loads(output_text) if output_text.strip() else {}
            if result.returncode != 0 or not parsed.get("ok", True):
                raise RuntimeError(str(parsed.get("error") or result.stderr or result.stdout or "paddleocr_vl_failed")[-1200:])
            parsed.setdefault("quality", "medium")
            parsed.setdefault("lines", [])
            return parsed
        finally:
            for path in (image_path, input_path, output_path):
                try:
                    path.unlink(missing_ok=True)
                except Exception:
                    pass

    def _command_available(self) -> bool:
        try:
            probe = self.command.format(image="/tmp/probe.jpg", page=1, total=1, model=self.model)
            parts = shlex.split(probe)
        except Exception:
            return False
        if not parts:
            return False
        executable = parts[0]
        return Path(executable).exists() or bool(shutil.which(executable))

    @staticmethod
    def _resolve_local_model_dir(configured: Optional[str]) -> Optional[Path]:
        candidates = [
            str(configured or "").strip(),
            str(Path.home() / ".paddlex" / "official_models" / "PaddleOCR-VL-1.5"),
        ]
        for candidate in candidates:
            if not candidate:
                continue
            path = Path(candidate).expanduser()
            if path.exists():
                return path.resolve()
        return None

    @staticmethod
    def _resolve_python_path(configured: Optional[str]) -> str:
        candidate = str(configured or sys.executable).strip()
        if not candidate:
            return sys.executable
        path = Path(candidate).expanduser()
        if path.exists():
            return str(path if path.is_absolute() else path.absolute())
        return candidate

    @staticmethod
    def _temp_json_file(directory: Path, prefix: str) -> Path:
        handle = tempfile.NamedTemporaryFile(
            mode="w",
            encoding="utf-8",
            prefix=prefix,
            suffix=".json",
            dir=str(directory),
            delete=False,
        )
        handle.close()
        return Path(handle.name)

    @staticmethod
    def _write_temp_image(image_bytes: bytes) -> Path:
        handle = tempfile.NamedTemporaryFile(prefix="paddleocr-vl-page-", suffix=".jpg", delete=False)
        try:
            handle.write(image_bytes)
            return Path(handle.name)
        finally:
            handle.close()

    def _prompt(self, *, page_number: int, total_pages: int) -> str:
        return (
            "ocr: Carefully transcribe this difficult PDF page. "
            f"Page {page_number}/{total_pages}. Return strict JSON: "
            '{"text":"...", "quality":"high|medium|low|failed", '
            '"lines":[{"text":"...", "bbox":[0,0,1000,1000], "confidence":0.0}], "warnings":["..."]}. '
            "Do not summarize or infer invisible content."
        )


class LocalHighQualityPdfNormalizationService:
    """Normalize local_high_quality PDF inputs into a DOCX source for desensitization."""

    def __init__(self, *, progress_path: Optional[str] = None) -> None:
        self.lowmem_pdf_service = PdfNormalizationService()
        self._rapid_layout_ocr: Optional[RapidOCROnNXService] = None
        self._progress_path = Path(progress_path).expanduser() if progress_path else None

    def normalize_pdf(self, file_path: str, *, output_dir: Optional[str] = None) -> Dict[str, Any]:
        source_path = Path(file_path).expanduser().resolve()
        if source_path.suffix.lower() != ".pdf":
            raise ValueError(f"LocalHighQualityPdfNormalizationService only accepts PDFs: {source_path}")

        output_root = Path(output_dir).expanduser().resolve() if output_dir else self._default_output_dir()
        ensure_private_directory(output_root)
        normalized_path = output_root / f"{source_path.stem}_{uuid.uuid4().hex[:10]}_local_hq.docx"

        classification = self.lowmem_pdf_service.classify_pdf(str(source_path))
        warnings: List[str] = []
        if classification.get("pdf_kind") == "native_text":
            return self._normalize_native_text_pdf(
                source_path=source_path,
                normalized_path=normalized_path,
                classification=classification,
                warnings=warnings,
            )
        return self._normalize_scanned_or_mixed_pdf(
            source_path=source_path,
            normalized_path=normalized_path,
            classification=classification,
            warnings=warnings,
        )

    def _normalize_native_text_pdf(
        self,
        *,
        source_path: Path,
        normalized_path: Path,
        classification: Dict[str, Any],
        warnings: List[str],
    ) -> Dict[str, Any]:
        try:
            self.lowmem_pdf_service._convert_text_pdf_with_pdf2docx(str(source_path), normalized_path)
            docx_recovery_engine = "pdf2docx"
        except Exception as exc:
            logger.warning("local_high_quality pdf2docx failed; using native rebuild: %s", exc)
            warnings.append("pdf2docx_failed_fallback_used")
            self.lowmem_pdf_service._rebuild_native_docx(classification.get("pages") or [], normalized_path)
            docx_recovery_engine = "native_rebuild"

        metadata = self._build_metadata(
            source_path=str(source_path),
            normalized_path=str(normalized_path),
            classification=classification,
            warnings=warnings,
            docx_recovery_engine=docx_recovery_engine,
            quality_counts={"high": 0, "medium": 0, "low": 0, "failed": 0},
            vl_fallback_pages=[],
            alignment_summary={"aligned_blocks": 0, "low_alignment_pages": []},
        )
        metadata["pdf_normalization_engine"] = docx_recovery_engine
        metadata["ocr_quality_gate"] = "not_applicable"
        ensure_private_file(normalized_path)
        return {
            "normalized_file_path": str(normalized_path),
            "metadata": metadata,
            "structure": {"pages": self.lowmem_pdf_service._public_pages(classification.get("pages") or [])},
        }

    def _normalize_scanned_or_mixed_pdf(
        self,
        *,
        source_path: Path,
        normalized_path: Path,
        classification: Dict[str, Any],
        warnings: List[str],
    ) -> Dict[str, Any]:
        glm_client = GlmOcrOllamaClient(
            base_url=settings.OLLAMA_BASE_URL,
            model=settings.LOCAL_PDF_GLM_MODEL,
            timeout=settings.PDF_OCR_TIMEOUT,
            max_tokens=settings.LOCAL_PDF_GLM_MAX_TOKENS,
        )
        if not glm_client.available:
            metadata = self._build_metadata(
                source_path=str(source_path),
                normalized_path=str(normalized_path),
                classification=classification,
                warnings=[*warnings, "glm_ocr_model_missing"],
                docx_recovery_engine="not_started",
                quality_counts={"high": 0, "medium": 0, "low": 0, "failed": 0},
                vl_fallback_pages=[],
                alignment_summary={"aligned_blocks": 0, "low_alignment_pages": []},
            )
            raise LocalHighQualityPdfNormalizationError(
                "glm_ocr_model_missing",
                f"GLM-OCR model is not available in Ollama: {settings.LOCAL_PDF_GLM_MODEL}",
                metadata=metadata,
            )

        recovery_probe = self._run_ppstructure_recovery_sidecar(source_path, normalized_path.with_suffix(".ppstructure.docx"))
        docx_recovery_engine = "ppstructure_recovery_to_doc" if recovery_probe.get("available") else "layout_ir_docx_renderer"
        recovery_warning = str(recovery_probe.get("warning") or "").strip()
        if not recovery_probe.get("available") and recovery_warning:
            warnings.append(recovery_warning)
        elif recovery_probe.get("warning"):
            warnings.append(str(recovery_probe["warning"]))

        vl_client = None
        if settings.LOCAL_PDF_VL_FALLBACK_ENABLED:
            vl_client = PaddleOcrVlFallbackClient(
                base_url=settings.OLLAMA_BASE_URL,
                model=settings.LOCAL_PDF_VL_MODEL,
                timeout=max(int(settings.LOCAL_PDF_VL_TIMEOUT or 0), int(settings.PDF_OCR_TIMEOUT or 0)),
                max_tokens=min(int(settings.LOCAL_PDF_GLM_MAX_TOKENS or 8192), 8192),
                api_mode=settings.LOCAL_PDF_VL_API_MODE,
                model_dir=settings.LOCAL_PDF_VL_MODEL_DIR,
                python_path=self._resolve_ocr_python(),
                command=settings.LOCAL_PDF_VL_COMMAND
                if str(settings.LOCAL_PDF_VL_API_MODE or "").strip().lower() == "command"
                else None,
            )

        pages = classification.get("pages") or []
        total_pages = len(pages)
        classified_by_page = {
            int(page.get("page_number") or index + 1): page for index, page in enumerate(pages)
        }
        page_index_by_number = {
            int(page.get("page_number") or index + 1): index for index, page in enumerate(pages)
        }
        image_page_count = sum(1 for page in pages if page.get("kind") != "native_text")
        vl_fallback_budget = self._vl_fallback_budget(total_pages=total_pages, image_page_count=image_page_count)
        checkpoint_dir, checkpoint_id = self._prepare_frontline_checkpoint(source_path, classification)
        checkpoint_resume_pages: List[int] = []

        result_pages: List[LocalPdfPageResult] = []
        quality_counts = {"high": 0, "medium": 0, "low": 0, "failed": 0}
        vl_fallback_pages: List[int] = []
        vl_fallback_attempted_pages: List[int] = []
        vl_fallback_failed_pages: List[int] = []
        vl_fallback_rejected_pages: List[int] = []
        ocr_candidate_scores: List[Dict[str, Any]] = []
        high_res_retry_pages: List[int] = []
        high_res_retry_used_pages: List[int] = []
        frontline_page_timings: List[Dict[str, Any]] = []
        rendered_temp_context = tempfile.TemporaryDirectory(prefix="local-pdf-frontline-pages-")
        rendered_image_root = Path(rendered_temp_context.name)
        ensure_private_directory(rendered_image_root)
        rendered_image_paths: Dict[int, Path] = {}

        pdf_document = pdfium.PdfDocument(str(source_path))
        try:
            for page_index, classified_page in enumerate(pages):
                page_number = int(classified_page.get("page_number") or page_index + 1)
                page_start = time.perf_counter()
                checkpoint_payload = self._load_checkpoint_page(checkpoint_dir, page_number)
                if checkpoint_payload:
                    page_result = LocalPdfPageResult.from_public_dict(checkpoint_payload.get("result") or {})
                    result_pages.append(page_result)
                    checkpoint_resume_pages.append(page_number)
                    frontline_page_timings.append(
                        {
                            "page": page_number,
                            "seconds": 0.0,
                            "stage": "checkpoint_resume",
                            "quality": page_result.quality,
                        }
                    )
                    warnings.append("checkpoint_resume_used")
                    continue

                if classified_page.get("kind") == "native_text":
                    page_result = self._native_page_result(classified_page)
                    result_pages.append(page_result)
                    metrics = self._page_quality_metrics(page_result, classified_page)
                    timing = {
                        "page": page_number,
                        "seconds": round(time.perf_counter() - page_start, 4),
                        "stage": "native_text",
                        "quality": page_result.quality,
                    }
                    frontline_page_timings.append(timing)
                    self._write_checkpoint_page(checkpoint_dir, page_result, metrics, timing)
                    continue

                try:
                    logger.info("Local high-quality PDF anchor OCR page %s/%s started.", page_number, total_pages)
                    self._emit_progress(
                        stage="anchor_preflight",
                        current=page_number,
                        total=total_pages,
                        page=page_number,
                        message=f"正在快速识别第 {page_number}/{total_pages} 页...",
                    )
                    image = self._render_pdf_page_image(
                        pdf_document,
                        page_index,
                        classified_page,
                        dpi=self._glm_render_dpi(),
                    )
                    rendered_image_paths[page_number] = self._save_frontline_page_image(
                        image,
                        rendered_image_root,
                        page_number,
                    )
                    page_result = self.build_anchor_page_result(classified_page, image, warnings=warnings)
                    result_pages.append(page_result)
                    metrics = self._page_quality_metrics(page_result, classified_page)
                    frontline_page_timings.append(
                        {
                            "page": page_number,
                            "seconds": round(time.perf_counter() - page_start, 4),
                            "stage": "anchor_preflight",
                            "engine": page_result.engine,
                            "quality": page_result.quality,
                            "readable_chars": metrics["readable_chars"],
                        }
                    )
                    logger.info(
                        "Local high-quality PDF anchor OCR page %s/%s finished: quality=%s chars=%s.",
                        page_number,
                        total_pages,
                        page_result.quality,
                        self.lowmem_pdf_service.count_readable_chars(page_result.text),
                    )
                except Exception as exc:
                    logger.warning("Anchor OCR failed on local high quality PDF page %s: %s", page_number, exc, exc_info=True)
                    warnings.append("ppocrv5_anchor_failed")
                    page_result = LocalPdfPageResult(
                        page_number=page_number,
                        width=float(classified_page.get("width") or 595.0),
                        height=float(classified_page.get("height") or 842.0),
                        text="",
                        source="ocr_failed",
                        quality="failed",
                        confidence=0.0,
                        lines=[],
                        blocks=[],
                        warnings=["ppocrv5_anchor_failed"],
                        engine="ppocrv5_anchor",
                    )
                    result_pages.append(page_result)
                    frontline_page_timings.append(
                        {
                            "page": page_number,
                            "seconds": round(time.perf_counter() - page_start, 4),
                            "stage": "anchor_failed",
                            "quality": "failed",
                        }
                    )
                finally:
                    if "image" in locals():
                        del image
                    gc.collect()
        finally:
            pdf_document.close()

        self._emit_progress(
            stage="layout_analysis",
            current=total_pages,
            total=total_pages,
            message="正在分析 PDF 版面结构...",
        )
        self._apply_ppstructurev3_batch_layout(
            result_pages=result_pages,
            rendered_image_paths=rendered_image_paths,
            warnings=warnings,
        )
        glm_page_numbers: List[int] = []
        glm_tiers: Dict[int, str] = {}
        anchor_only_pages: List[int] = []
        glm_page_pages: List[int] = []
        glm_block_pages: List[int] = []
        quality_gate_rejected_pages: List[int] = []
        for page_result in list(sorted(result_pages, key=lambda item: item.page_number)):
            if page_result.source == "native" or page_result.page_number in checkpoint_resume_pages:
                continue
            classified_page = classified_by_page.get(page_result.page_number) or {}
            tier = self.classify_page_ocr_tier(page_result, classified_page)
            if tier in {"glm_page", "glm_block"}:
                glm_page_numbers.append(page_result.page_number)
                glm_tiers[page_result.page_number] = tier
                continue
            final_anchor = self._mark_anchor_output_used(page_result)
            self._replace_result_page(result_pages, page_result.page_number, final_anchor)
            anchor_only_pages.append(page_result.page_number)
            metrics = self._page_quality_metrics(final_anchor, classified_page)
            timing = {
                "page": page_result.page_number,
                "seconds": None,
                "stage": "anchor_final",
                "engine": final_anchor.engine,
                "quality": final_anchor.quality,
                "readable_chars": metrics["readable_chars"],
            }
            frontline_page_timings.append(timing)
            self._write_checkpoint_page(checkpoint_dir, final_anchor, metrics, timing)

        if glm_page_numbers:
            pdf_document = pdfium.PdfDocument(str(source_path))
            try:
                for page_number in glm_page_numbers:
                    page_index = page_index_by_number.get(page_number)
                    classified_page = classified_by_page.get(page_number) or {}
                    image_path = rendered_image_paths.get(page_number)
                    if page_index is None or image_path is None:
                        continue
                    page_start = time.perf_counter()
                    current = self._page_result_by_number(result_pages, page_number)
                    tier = glm_tiers.get(page_number, "glm_page")
                    try:
                        logger.info("Local high-quality PDF %s OCR page %s/%s started.", tier, page_number, total_pages)
                        self._emit_progress(
                            stage=tier,
                            current=page_number,
                            total=total_pages,
                            page=page_number,
                            message=f"正在精识别第 {page_number}/{total_pages} 页...",
                        )
                        with Image.open(image_path) as image:
                            image.load()
                            page_result = self.run_glm_for_page_or_blocks(
                                glm_client=glm_client,
                                image=image,
                                current=current,
                                classified_page=classified_page,
                                page_number=page_number,
                                total_pages=total_pages,
                                tier=tier,
                            )
                            if self._should_high_res_retry(page_result):
                                high_res_retry_pages.append(page_number)
                                retry_result = self._run_high_res_glm_retry(
                                    pdf_document=pdf_document,
                                    page_index=page_index,
                                    classified_page=classified_page,
                                    glm_client=glm_client,
                                    page_number=page_number,
                                    total_pages=total_pages,
                                )
                                if retry_result and self._is_better_page_result(page_result, retry_result):
                                    page_result = retry_result
                                    high_res_retry_used_pages.append(page_number)
                            if current is not None:
                                page_result = self._merge_text_with_existing_layout(
                                    page_result,
                                    current,
                                    warning="glm_text_existing_layout",
                                    engine_suffix="layout_reuse",
                                )
                            if not self._page_has_layout_signal(page_result):
                                page_result = self._try_rapidocr_layout_assist(page_result, image)
                        page_result, rejected = self.compare_anchor_glm_quality(
                            anchor=current,
                            glm_result=page_result,
                            classified_page=classified_page,
                        )
                        if rejected:
                            quality_gate_rejected_pages.append(page_number)
                        self._replace_result_page(result_pages, page_number, page_result)
                        if tier == "glm_block":
                            glm_block_pages.append(page_number)
                        else:
                            glm_page_pages.append(page_number)
                        metrics = self._page_quality_metrics(page_result, classified_page)
                        timing = {
                            "page": page_number,
                            "seconds": round(time.perf_counter() - page_start, 4),
                            "stage": tier,
                            "engine": page_result.engine,
                            "quality": page_result.quality,
                            "readable_chars": metrics["readable_chars"],
                            "quality_gate_rejected": rejected,
                        }
                        frontline_page_timings.append(timing)
                        self._write_checkpoint_page(checkpoint_dir, page_result, metrics, timing)
                        logger.info(
                            "Local high-quality PDF GLM OCR page %s/%s finished: engine=%s quality=%s chars=%s.",
                            page_number,
                            total_pages,
                            page_result.engine,
                            page_result.quality,
                            self.lowmem_pdf_service.count_readable_chars(page_result.text),
                        )
                    except Exception as exc:
                        logger.warning("GLM-OCR failed on local high quality PDF page %s: %s", page_number, exc, exc_info=True)
                        warnings.append("glm_ocr_failed_pages")
                        page_result = self._anchor_fallback_after_glm_failure(current, classified_page)
                        self._replace_result_page(result_pages, page_number, page_result)
                        metrics = self._page_quality_metrics(page_result, classified_page)
                        timing = {
                            "page": page_number,
                            "seconds": round(time.perf_counter() - page_start, 4),
                            "stage": "glm_failed_anchor_fallback" if page_result.text.strip() else "glm_failed",
                            "quality": page_result.quality,
                        }
                        frontline_page_timings.append(timing)
                        self._write_checkpoint_page(checkpoint_dir, page_result, metrics, timing)
                    finally:
                        gc.collect()
            finally:
                pdf_document.close()

        page_quality_metrics = self._collect_page_quality_metrics(result_pages, classified_by_page)
        vl_candidates = self._rank_vl_fallback_candidates(result_pages, classified_by_page, page_quality_metrics)
        vl_fallback_candidate_pages = [int(candidate["page"]) for candidate in vl_candidates]
        selected_vl_candidates = vl_candidates[:vl_fallback_budget]
        if vl_candidates and len(vl_candidates) > vl_fallback_budget:
            warnings.append("paddleocr_vl_budget_exhausted")

        if selected_vl_candidates:
            logger.info(
                "Local high-quality PDF selected %s/%s PaddleOCR-VL fallback pages: %s.",
                len(selected_vl_candidates),
                len(vl_candidates),
                [candidate["page"] for candidate in selected_vl_candidates],
            )
            pdf_document = pdfium.PdfDocument(str(source_path))
            try:
                for candidate_info in selected_vl_candidates:
                    page_number = int(candidate_info["page"])
                    page_index = page_index_by_number.get(page_number)
                    classified_page = classified_by_page.get(page_number) or {}
                    if page_index is None:
                        continue
                    vl_start = time.perf_counter()
                    vl_fallback_attempted_pages.append(page_number)
                    self._emit_progress(
                        stage="vl_fallback",
                        current=len(vl_fallback_attempted_pages),
                        total=max(1, len(selected_vl_candidates)),
                        page=page_number,
                        message=f"正在兜底复核第 {page_number} 页...",
                    )
                    try:
                        image = self._render_pdf_page_image(
                            pdf_document,
                            page_index,
                            classified_page,
                            dpi=max(self._glm_render_dpi(), self._glm_high_res_dpi()),
                        )
                        try:
                            image_bytes = self._enhanced_image_bytes(image)
                            vl_page_result = self._try_vl_fallback(
                                vl_client=vl_client,
                                image_bytes=image_bytes,
                                classified_page=classified_page,
                                page_number=page_number,
                                total_pages=total_pages,
                                warnings=warnings,
                            )
                        finally:
                            if "image_bytes" in locals():
                                del image_bytes
                    finally:
                        if "image" in locals():
                            del image
                        gc.collect()
                    if vl_page_result:
                        current = self._page_result_by_number(result_pages, page_number)
                        candidate = self._merge_candidate_text_with_existing_layout(vl_page_result, current) if current else vl_page_result
                        current_score = self._page_result_score(current) if current else 0.0
                        candidate_score = self._page_result_score(candidate)
                        ocr_candidate_scores.append(
                            {
                                "page": page_number,
                                "current_engine": current.engine if current else None,
                                "candidate_engine": candidate.engine,
                                "current_score": round(current_score, 4),
                                "candidate_score": round(candidate_score, 4),
                                "current_quality": current.quality if current else None,
                                "candidate_quality": candidate.quality,
                                "current_chars": self.lowmem_pdf_service.count_readable_chars(current.text) if current else 0,
                                "candidate_chars": self.lowmem_pdf_service.count_readable_chars(candidate.text),
                                "selection_reason": candidate_info.get("reason"),
                            }
                        )
                        if current is None or self._is_better_page_result(current, candidate):
                            candidate.used_vl_fallback = True
                            self._replace_result_page(result_pages, page_number, candidate)
                            vl_fallback_pages.append(page_number)
                            warnings.append("paddleocr_vl_fallback_used")
                            metrics = self._page_quality_metrics(candidate, classified_page)
                            vl_timing = {
                                "page": page_number,
                                "seconds": round(time.perf_counter() - vl_start, 4),
                                "stage": "paddleocr_vl_fallback",
                                "quality": candidate.quality,
                                "adopted": True,
                            }
                            frontline_page_timings.append(vl_timing)
                            self._write_checkpoint_page(
                                checkpoint_dir,
                                candidate,
                                metrics,
                                vl_timing,
                            )
                        else:
                            vl_fallback_rejected_pages.append(page_number)
                            frontline_page_timings.append(
                                {
                                    "page": page_number,
                                    "seconds": round(time.perf_counter() - vl_start, 4),
                                    "stage": "paddleocr_vl_fallback",
                                    "quality": candidate.quality,
                                    "adopted": False,
                                }
                            )
                    else:
                        vl_fallback_failed_pages.append(page_number)
                        frontline_page_timings.append(
                            {
                                "page": page_number,
                                "seconds": round(time.perf_counter() - vl_start, 4),
                                "stage": "paddleocr_vl_fallback_failed",
                                "quality": "failed",
                                "adopted": False,
                            }
                        )
            finally:
                pdf_document.close()

        page_quality_metrics = self._collect_page_quality_metrics(result_pages, classified_by_page)
        failed_pages = [
            page.page_number for page in result_pages if page.quality == "failed" or not page.text.strip()
        ]
        low_quality_pages = [
            page.page_number for page in result_pages if page.quality in {"low", "failed"}
        ]
        low_alignment_pages = [
            page.page_number for page in result_pages if not self._page_has_layout_signal(page)
        ]
        for page_result in result_pages:
            quality_counts[page_result.quality] = int(quality_counts.get(page_result.quality, 0)) + 1

        if failed_pages:
            rendered_temp_context.cleanup()
            warnings.append("ocr_quality_gate_failed")
            metadata = self._build_metadata(
                source_path=str(source_path),
                normalized_path=str(normalized_path),
                classification=classification,
                warnings=warnings,
                docx_recovery_engine=docx_recovery_engine,
                quality_counts=quality_counts,
                vl_fallback_pages=vl_fallback_pages,
                vl_fallback_attempted_pages=vl_fallback_attempted_pages,
                vl_fallback_failed_pages=vl_fallback_failed_pages,
                vl_fallback_rejected_pages=vl_fallback_rejected_pages,
                ocr_candidate_scores=ocr_candidate_scores,
                alignment_summary={"aligned_blocks": 0, "low_alignment_pages": low_alignment_pages},
                page_quality_metrics=page_quality_metrics,
                vl_fallback_candidate_pages=vl_fallback_candidate_pages,
                vl_fallback_budget=vl_fallback_budget,
                checkpoint_id=checkpoint_id,
                checkpoint_resume_pages=checkpoint_resume_pages,
                frontline_page_timings=frontline_page_timings,
                anchor_only_pages=anchor_only_pages,
                glm_page_pages=glm_page_pages,
                glm_block_pages=glm_block_pages,
                quality_gate_rejected_pages=quality_gate_rejected_pages,
            )
            metadata["ocr_failed_pages"] = sorted(set(failed_pages))
            metadata["ocr_high_res_retry_pages"] = sorted(set(high_res_retry_pages))
            metadata["ocr_high_res_retry_used_pages"] = sorted(set(high_res_retry_used_pages))
            raise LocalHighQualityPdfNormalizationError(
                "ocr_quality_gate_failed",
                f"Local high-quality PDF OCR failed on pages: {sorted(set(failed_pages))}",
                metadata=metadata,
            )

        if low_quality_pages:
            warnings.append("ocr_quality_gate_failed")
            warnings.append("ocr_low_quality_pages")
        if low_alignment_pages:
            warnings.append("layout_alignment_low")

        final_docx_recovery_engine = "layout_ir_docx_renderer"
        self._emit_progress(
            stage="docx_render",
            current=total_pages,
            total=total_pages,
            message="正在生成可编辑 DOCX...",
        )
        if self._try_write_recovered_docx_with_final_text(recovery_probe, result_pages, normalized_path):
            final_docx_recovery_engine = "ppstructure_recovery_to_doc_text_rewrite"
            warnings.append("ppstructure_docx_recovery_used")
        else:
            if not recovery_probe.get("disabled"):
                warnings.append("docx_renderer_fallback_used")
            self._render_layout_ir_docx(result_pages, normalized_path)
        ensure_private_file(normalized_path)

        metadata = self._build_metadata(
            source_path=str(source_path),
            normalized_path=str(normalized_path),
            classification=classification,
            warnings=warnings,
            docx_recovery_engine=final_docx_recovery_engine,
            quality_counts=quality_counts,
            vl_fallback_pages=vl_fallback_pages,
            vl_fallback_attempted_pages=vl_fallback_attempted_pages,
            vl_fallback_failed_pages=vl_fallback_failed_pages,
            vl_fallback_rejected_pages=vl_fallback_rejected_pages,
            ocr_candidate_scores=ocr_candidate_scores,
            alignment_summary={
                "aligned_blocks": sum(
                    len(page.blocks)
                    for page in result_pages
                    if page.blocks and self._page_has_layout_signal(page)
                ),
                "low_alignment_pages": sorted(set(low_alignment_pages)),
            },
            page_quality_metrics=page_quality_metrics,
            vl_fallback_candidate_pages=vl_fallback_candidate_pages,
            vl_fallback_budget=vl_fallback_budget,
            checkpoint_id=checkpoint_id,
            checkpoint_resume_pages=checkpoint_resume_pages,
            frontline_page_timings=frontline_page_timings,
            anchor_only_pages=anchor_only_pages,
            glm_page_pages=glm_page_pages,
            glm_block_pages=glm_block_pages,
            quality_gate_rejected_pages=quality_gate_rejected_pages,
        )
        metadata.update(
            {
                "ocr_pages": sum(1 for page in result_pages if page.source in {"glm_ocr", "paddleocr_vl", "ppocr_anchor"}),
                "ocr_failed_pages": [],
                "ocr_low_quality_pages": sorted(set(low_quality_pages)),
                "ocr_quality_gate": "failed" if low_quality_pages else "passed",
                "ocr_high_res_retry_pages": sorted(set(high_res_retry_pages)),
                "ocr_high_res_retry_used_pages": sorted(set(high_res_retry_used_pages)),
                "ppstructure_recovery_available": bool(recovery_probe.get("available")),
                "ppstructure_recovery_disabled": bool(recovery_probe.get("disabled")),
                "ppstructure_recovery_output_path": str(recovery_probe.get("output_docx_path") or "") or None,
            }
        )
        self._cleanup_checkpoint(checkpoint_dir)
        rendered_temp_context.cleanup()
        return {
            "normalized_file_path": str(normalized_path),
            "metadata": metadata,
            "structure": {"pages": [page.to_public_dict() for page in result_pages]},
        }

    def _prepare_frontline_checkpoint(
        self,
        source_path: Path,
        classification: Dict[str, Any],
    ) -> tuple[Optional[Path], Optional[str]]:
        if not bool(getattr(settings, "LOCAL_PDF_PAGE_CHECKPOINT_ENABLED", True)):
            return None, None
        try:
            checkpoint_id = self._checkpoint_id(source_path, classification)
            root = Path(settings.RUNTIME_ROOT) / "local_pdf_frontline_checkpoints"
            ensure_private_directory(root)
            checkpoint_dir = root / checkpoint_id
            ensure_private_directory(checkpoint_dir)
            manifest = {
                "checkpoint_id": checkpoint_id,
                "source_path": str(source_path),
                "settings_hash": self._frontline_settings_hash(),
                "pipeline_version": LOCAL_PDF_FRONTLINE_PIPELINE_VERSION,
                "profile": str(getattr(settings, "LOCAL_PDF_FRONTLINE_PROFILE", "")),
                "created_at": time.time(),
                "ttl_hours": int(getattr(settings, "LOCAL_PDF_CHECKPOINT_TTL_HOURS", 24) or 24),
            }
            self._atomic_write_json(checkpoint_dir / "manifest.json", manifest)
            self._cleanup_expired_checkpoints(root)
            return checkpoint_dir, checkpoint_id
        except Exception:
            logger.debug("Local PDF frontline checkpoint init failed.", exc_info=True)
            return None, None

    def _checkpoint_id(self, source_path: Path, classification: Dict[str, Any]) -> str:
        source_hash = hashlib.sha256()
        with source_path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                source_hash.update(chunk)
        pages = [
            {
                "page_number": page.get("page_number"),
                "kind": page.get("kind"),
                "width": page.get("width"),
                "height": page.get("height"),
            }
            for page in classification.get("pages") or []
        ]
        payload = {
            "source_sha256": source_hash.hexdigest(),
            "settings_hash": self._frontline_settings_hash(),
            "pages": pages,
        }
        digest = hashlib.sha256(json.dumps(payload, sort_keys=True, ensure_ascii=False).encode("utf-8")).hexdigest()
        return digest[:32]

    def _frontline_settings_hash(self) -> str:
        keys = [
            "LOCAL_PDF_FRONTLINE_PROFILE",
            "LOCAL_PDF_GLM_MODEL",
            "LOCAL_PDF_GLM_POLICY",
            "LOCAL_PDF_GLM_SCOPE",
            "LOCAL_PDF_GLM_RENDER_DPI",
            "LOCAL_PDF_GLM_HIGH_RES_DPI",
            "LOCAL_PDF_PPSTRUCTUREV3_POLICY",
            "LOCAL_PDF_PPOCRV5_POLICY",
            "LOCAL_PDF_FAST_ANCHOR_MIN_CONFIDENCE",
            "LOCAL_PDF_FAST_ANCHOR_MIN_CHARS",
            "LOCAL_PDF_FAST_ANCHOR_MAX_LOW_CONF_RATIO",
            "LOCAL_PDF_FAST_ANCHOR_MAX_ABNORMAL_RATIO",
            "LOCAL_PDF_VL_FALLBACK_POLICY",
            "LOCAL_PDF_VL_MAX_PAGES",
            "LOCAL_PDF_VL_MAX_PAGE_RATIO",
            "LOCAL_PDF_DOCX_RECOVERY_ENGINE",
        ]
        payload = {key: getattr(settings, key, None) for key in keys}
        payload["pipeline_version"] = LOCAL_PDF_FRONTLINE_PIPELINE_VERSION
        return hashlib.sha256(json.dumps(payload, sort_keys=True, ensure_ascii=False).encode("utf-8")).hexdigest()

    def _load_checkpoint_page(self, checkpoint_dir: Optional[Path], page_number: int) -> Optional[Dict[str, Any]]:
        if checkpoint_dir is None:
            return None
        page_path = checkpoint_dir / f"page_{int(page_number):04d}.json"
        if not page_path.exists():
            return None
        try:
            payload = json.loads(page_path.read_text(encoding="utf-8"))
        except Exception:
            return None
        if not isinstance(payload, dict) or not isinstance(payload.get("result"), dict):
            return None
        return payload

    def _write_checkpoint_page(
        self,
        checkpoint_dir: Optional[Path],
        page_result: LocalPdfPageResult,
        metrics: Dict[str, Any],
        timing: Dict[str, Any],
    ) -> None:
        if checkpoint_dir is None:
            return
        try:
            payload = {
                "page": page_result.page_number,
                "updated_at": time.time(),
                "result": page_result.to_public_dict(),
                "metrics": metrics,
                "timing": timing,
            }
            self._atomic_write_json(checkpoint_dir / f"page_{page_result.page_number:04d}.json", payload)
        except Exception:
            logger.debug("Local PDF frontline checkpoint page write failed.", exc_info=True)

    def _cleanup_checkpoint(self, checkpoint_dir: Optional[Path]) -> None:
        if checkpoint_dir is None:
            return
        try:
            shutil.rmtree(checkpoint_dir, ignore_errors=True)
        except Exception:
            logger.debug("Local PDF frontline checkpoint cleanup failed: %s", checkpoint_dir, exc_info=True)

    def _cleanup_expired_checkpoints(self, root: Path) -> None:
        ttl_seconds = max(1, int(getattr(settings, "LOCAL_PDF_CHECKPOINT_TTL_HOURS", 24) or 24)) * 3600
        cutoff = time.time() - ttl_seconds
        try:
            for child in root.iterdir():
                if not child.is_dir():
                    continue
                manifest = child / "manifest.json"
                try:
                    created_at = float((json.loads(manifest.read_text(encoding="utf-8")) or {}).get("created_at") or 0)
                except Exception:
                    created_at = child.stat().st_mtime
                if created_at < cutoff:
                    shutil.rmtree(child, ignore_errors=True)
        except Exception:
            logger.debug("Local PDF frontline expired checkpoint cleanup failed.", exc_info=True)

    def _atomic_write_json(self, path: Path, payload: Dict[str, Any]) -> None:
        ensure_private_directory(path.parent)
        temp_path = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
        temp_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
        ensure_private_file(temp_path)
        temp_path.replace(path)
        ensure_private_file(path)

    def _emit_progress(
        self,
        *,
        stage: str,
        current: int = 0,
        total: int = 0,
        message: str = "",
        page: Optional[int] = None,
    ) -> None:
        if self._progress_path is None:
            return
        payload: Dict[str, Any] = {
            "stage": stage,
            "current": max(0, int(current or 0)),
            "total": max(0, int(total or 0)),
            "message": message,
            "updated_at": time.time(),
        }
        if page is not None:
            payload["page"] = int(page)
        try:
            self._atomic_write_json(self._progress_path, payload)
        except Exception:
            logger.debug("Local PDF frontline progress write failed.", exc_info=True)

    def _enhanced_image_bytes(self, image: Image.Image) -> bytes:
        enhanced = self._enhance_for_ocr(image)
        try:
            return self._image_to_jpeg_bytes(enhanced)
        finally:
            del enhanced

    def _collect_page_quality_metrics(
        self,
        result_pages: List[LocalPdfPageResult],
        classified_by_page: Dict[int, Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        return [
            self._page_quality_metrics(page, classified_by_page.get(page.page_number) or {})
            for page in sorted(result_pages, key=lambda item: item.page_number)
        ]

    def _page_quality_metrics(
        self,
        page_result: LocalPdfPageResult,
        classified_page: Dict[str, Any],
    ) -> Dict[str, Any]:
        readable_chars = self.lowmem_pdf_service.count_readable_chars(page_result.text)
        anchor_texts = [
            str(line.get("layout_text") or line.get("anchor_text") or "")
            for line in page_result.lines
            if str(line.get("layout_text") or line.get("anchor_text") or "").strip()
            and self._safe_float(line.get("confidence"), default=0.0) >= 0.7
        ]
        anchor_chars = sum(self.lowmem_pdf_service.count_readable_chars(text) for text in anchor_texts)
        anchor_recall = self._anchor_recall(page_result.text, anchor_texts) if anchor_texts else None
        char_count_ratio = (
            round(readable_chars / max(anchor_chars, 1), 4) if anchor_chars else None
        )
        duplicate_ratio = self._duplicate_ratio(page_result.text)
        abnormal_ratio = self._abnormal_char_ratio(page_result.text)
        layout_signal = self._page_has_layout_signal(page_result)
        score = self._page_result_score(page_result)
        quality_failed = page_result.quality == "failed" or not str(page_result.text or "").strip()
        quality_low = page_result.quality == "low"
        confidence_failed = (
            classified_page.get("kind") != "native_text"
            and page_result.quality != "high"
            and float(page_result.confidence or 0.0) < 0.72
        )
        anchor_failed = anchor_recall is not None and anchor_recall < 0.85
        chars_failed = char_count_ratio is not None and char_count_ratio < 0.75
        abnormal_failed = abnormal_ratio >= 0.08
        duplicate_failed = duplicate_ratio >= 0.18
        layout_failed = not layout_signal and classified_page.get("kind") != "native_text"
        warning_set = set(page_result.warnings or [])
        non_json_response = "glm_ocr_non_json_response" in warning_set
        unreadable_marker = "[UNREADABLE]" in str(page_result.text or "").upper()
        format_warning_requires_vl = non_json_response and any(
            [
                quality_failed,
                quality_low,
                confidence_failed,
                anchor_failed,
                chars_failed,
                abnormal_failed,
                duplicate_failed,
                layout_failed,
                unreadable_marker,
            ]
        )
        requires_vl = any(
            [
                quality_failed,
                quality_low,
                confidence_failed,
                anchor_failed,
                chars_failed,
                abnormal_failed,
                duplicate_failed,
                layout_failed,
                format_warning_requires_vl,
                unreadable_marker,
            ]
        )
        reason_parts: List[str] = []
        if quality_failed:
            reason_parts.append("failed")
        elif quality_low:
            reason_parts.append("low_quality")
        if anchor_failed:
            reason_parts.append("low_anchor_recall")
        if confidence_failed:
            reason_parts.append("low_confidence")
        if chars_failed:
            reason_parts.append("low_char_ratio")
        if abnormal_failed:
            reason_parts.append("abnormal_chars")
        if duplicate_failed:
            reason_parts.append("duplicate_text")
        if layout_failed:
            reason_parts.append("layout_missing")
        metrics = {
            "page": page_result.page_number,
            "quality": page_result.quality,
            "engine": page_result.engine,
            "source": page_result.source,
            "readable_chars": readable_chars,
            "confidence": round(float(page_result.confidence or 0.0), 4),
            "anchor_recall": round(anchor_recall, 4) if anchor_recall is not None else None,
            "char_count_ratio": char_count_ratio,
            "abnormal_char_ratio": round(abnormal_ratio, 4),
            "duplicate_ratio": round(duplicate_ratio, 4),
            "layout_coverage": 1.0 if layout_signal else 0.0,
            "score": round(score, 4),
            "requires_vl": bool(requires_vl),
            "reason": ",".join(reason_parts) if reason_parts else "passed",
        }
        if page_result.source == "ppocr_anchor":
            anchor_score = self.score_anchor_page(page_result, classified_page)
            metrics.update(
                {
                    "ocr_tier": self.classify_page_ocr_tier(page_result, classified_page),
                    "anchor_low_confidence_line_ratio": anchor_score["low_confidence_line_ratio"],
                    "anchor_low_confidence_line_count": anchor_score["low_confidence_line_count"],
                    "anchor_complex_layout": anchor_score["complex_layout"],
                    "anchor_bbox_order_unstable": anchor_score["bbox_order_unstable"],
                    "anchor_gate_reasons": list(anchor_score["reasons"]),
                }
            )
        return metrics

    def _anchor_recall(self, text: str, anchor_texts: List[str]) -> float:
        if not anchor_texts:
            return 1.0
        compact_text = self._compact_layout_match_text(text)
        if not compact_text:
            return 0.0
        matched = 0
        for anchor in anchor_texts:
            compact_anchor = self._compact_layout_match_text(anchor)
            if not compact_anchor:
                continue
            if compact_anchor in compact_text:
                matched += 1
                continue
            ratio = difflib.SequenceMatcher(None, compact_anchor, compact_text).quick_ratio()
            if ratio >= 0.72:
                matched += 1
        return matched / max(len(anchor_texts), 1)

    def _duplicate_ratio(self, text: str) -> float:
        lines = [
            self._compact_layout_match_text(line)
            for line in str(text or "").splitlines()
            if self._compact_layout_match_text(line)
        ]
        if len(lines) < 4:
            return 0.0
        duplicate_count = len(lines) - len(set(lines))
        return max(0.0, min(1.0, duplicate_count / max(len(lines), 1)))

    def _rank_vl_fallback_candidates(
        self,
        result_pages: List[LocalPdfPageResult],
        classified_by_page: Dict[int, Dict[str, Any]],
        page_quality_metrics: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        if not settings.LOCAL_PDF_VL_FALLBACK_ENABLED:
            return []
        policy = str(getattr(settings, "LOCAL_PDF_VL_FALLBACK_POLICY", "quality_gate") or "quality_gate").strip().lower()
        if policy in {"off", "disabled", "none", "false"}:
            return []
        metrics_by_page = {int(item.get("page") or 0): item for item in page_quality_metrics}
        candidates: List[Dict[str, Any]] = []
        for page_result in result_pages:
            classified_page = classified_by_page.get(page_result.page_number) or {}
            if classified_page.get("kind") == "native_text":
                continue
            metrics = metrics_by_page.get(page_result.page_number) or self._page_quality_metrics(page_result, classified_page)
            if policy in {"extreme_vl", "extreme", "all_pages"}:
                requires_vl = True
            elif policy == "always":
                # "always" now means every OCR page enters the candidate pool, but the budget still applies.
                requires_vl = True
            else:
                requires_vl = bool(metrics.get("requires_vl")) or self._should_use_vl_fallback(page_result, classified_page)
            if not requires_vl:
                continue
            severity = self._vl_candidate_severity(metrics)
            candidates.append(
                {
                    "page": page_result.page_number,
                    "severity": round(severity, 4),
                    "quality": page_result.quality,
                    "score": metrics.get("score"),
                    "reason": metrics.get("reason") or "quality_gate",
                }
            )
        return sorted(candidates, key=lambda item: (-float(item.get("severity") or 0.0), int(item.get("page") or 0)))

    def _vl_candidate_severity(self, metrics: Dict[str, Any]) -> float:
        severity = 0.0
        quality = str(metrics.get("quality") or "").lower()
        if quality == "failed":
            severity += 100.0
        elif quality == "low":
            severity += 70.0
        elif quality == "medium":
            severity += 30.0
        confidence = self._safe_float(metrics.get("confidence"), default=0.0)
        severity += max(0.0, 0.86 - confidence) * 40.0
        anchor_recall = metrics.get("anchor_recall")
        if isinstance(anchor_recall, (int, float)):
            severity += max(0.0, 0.85 - float(anchor_recall)) * 60.0
        char_ratio = metrics.get("char_count_ratio")
        if isinstance(char_ratio, (int, float)):
            severity += max(0.0, 0.75 - float(char_ratio)) * 50.0
        severity += self._safe_float(metrics.get("abnormal_char_ratio"), default=0.0) * 80.0
        severity += self._safe_float(metrics.get("duplicate_ratio"), default=0.0) * 45.0
        if not metrics.get("layout_coverage"):
            severity += 10.0
        return severity

    def _vl_fallback_budget(self, *, total_pages: int, image_page_count: int) -> int:
        if not settings.LOCAL_PDF_VL_FALLBACK_ENABLED:
            return 0
        policy = str(getattr(settings, "LOCAL_PDF_VL_FALLBACK_POLICY", "quality_gate") or "quality_gate").strip().lower()
        if policy in {"off", "disabled", "none", "false"}:
            return 0
        if policy in {"extreme_vl", "extreme", "all_pages"}:
            return max(0, int(image_page_count or total_pages or 0))
        max_pages = max(0, int(getattr(settings, "LOCAL_PDF_VL_MAX_PAGES", 2) or 0))
        ratio = max(0.0, min(1.0, float(getattr(settings, "LOCAL_PDF_VL_MAX_PAGE_RATIO", 0.2) or 0.0)))
        if max_pages <= 0 or ratio <= 0.0 or image_page_count <= 0:
            return 0
        ratio_budget = max(1, int(image_page_count * ratio + 0.999999))
        return max(0, min(max_pages, ratio_budget, image_page_count))

    @staticmethod
    def _page_result_by_number(result_pages: List[LocalPdfPageResult], page_number: int) -> Optional[LocalPdfPageResult]:
        for page in result_pages:
            if page.page_number == page_number:
                return page
        return None

    @staticmethod
    def _replace_result_page(
        result_pages: List[LocalPdfPageResult],
        page_number: int,
        replacement: LocalPdfPageResult,
    ) -> None:
        for index, page in enumerate(result_pages):
            if page.page_number == page_number:
                result_pages[index] = replacement
                return
        result_pages.append(replacement)

    def _native_page_result(self, page: Dict[str, Any]) -> LocalPdfPageResult:
        text = self.lowmem_pdf_service.normalize_text(str(page.get("text") or ""))
        lines = list(page.get("lines") or [])
        return LocalPdfPageResult(
            page_number=int(page.get("page_number") or 0),
            width=float(page.get("width") or 595.0),
            height=float(page.get("height") or 842.0),
            text=text,
            source="native",
            quality="high",
            confidence=1.0,
            lines=lines,
            blocks=list(page.get("blocks") or self._blocks_from_lines(lines, text)),
            engine="native_text",
        )

    def _page_result_from_ocr_payload(
        self,
        payload: Dict[str, Any],
        *,
        classified_page: Dict[str, Any],
        source: str,
        engine: str,
    ) -> LocalPdfPageResult:
        text = self.lowmem_pdf_service.normalize_text(str(payload.get("text") or ""))
        blocks = self._normalize_blocks(payload.get("blocks"))
        lines = self._normalize_lines(payload.get("lines"))
        if not text and blocks:
            text = self._blocks_to_text(blocks)
        if not text and lines:
            text = "\n".join(str(line.get("text") or "") for line in lines if str(line.get("text") or "").strip())
        if not lines:
            lines = self._lines_from_text(text)
        if not blocks:
            blocks = self._blocks_from_lines(lines, text)
        confidence = self._avg_confidence(lines)
        quality = self._normalize_quality(str(payload.get("quality") or ""), text=text, confidence=confidence)
        warnings = [str(item).strip() for item in payload.get("warnings") or [] if str(item).strip()]
        return LocalPdfPageResult(
            page_number=int(classified_page.get("page_number") or 0),
            width=float(classified_page.get("width") or 595.0),
            height=float(classified_page.get("height") or 842.0),
            text=text,
            source=source,
            quality=quality,
            confidence=round(confidence, 4),
            lines=lines,
            blocks=blocks,
            warnings=warnings,
            engine=engine,
        )

    def _try_vl_fallback(
        self,
        *,
        vl_client: Optional[PaddleOcrVlFallbackClient],
        image_bytes: bytes,
        classified_page: Dict[str, Any],
        page_number: int,
        total_pages: int,
        warnings: List[str],
        use_doc_unwarping: Optional[bool] = None,
    ) -> Optional[LocalPdfPageResult]:
        if vl_client is None or not vl_client.available:
            warnings.append("paddleocr_vl_missing")
            return None
        try:
            try:
                raw = vl_client.recognize(
                    image_bytes,
                    page_number=page_number,
                    total_pages=total_pages,
                    use_doc_unwarping=use_doc_unwarping,
                )
            except TypeError as exc:
                if "use_doc_unwarping" not in str(exc):
                    raise
                raw = vl_client.recognize(image_bytes, page_number=page_number, total_pages=total_pages)
            return self._page_result_from_ocr_payload(
                raw,
                classified_page=classified_page,
                source="paddleocr_vl",
                engine="paddleocr_vl_1_5",
            )
        except Exception as exc:
            logger.warning("PaddleOCR-VL fallback failed on page %s: %s", page_number, exc, exc_info=True)
            warnings.append("paddleocr_vl_fallback_failed")
            return None

    def build_anchor_page_result(
        self,
        classified_page: Dict[str, Any],
        image: Image.Image,
        *,
        warnings: List[str],
    ) -> LocalPdfPageResult:
        page_number = int(classified_page.get("page_number") or 0)
        if not bool(getattr(settings, "LOCAL_PDF_PPOCRV5_ENABLED", True)):
            warnings.append("ppocrv5_anchor_disabled")
            return LocalPdfPageResult(
                page_number=page_number,
                width=float(classified_page.get("width") or 595.0),
                height=float(classified_page.get("height") or 842.0),
                text="",
                source="ocr_failed",
                quality="failed",
                confidence=0.0,
                lines=[],
                blocks=[],
                warnings=["ppocrv5_anchor_disabled"],
                engine="ppocrv5_anchor",
            )
        try:
            lines = self._detect_rapidocr_layout_lines(image)
        except Exception as exc:
            logger.debug("PP-OCRv5 anchor OCR unavailable: %s", exc, exc_info=True)
            warnings.append("ppocrv5_anchor_failed")
            lines = []
        normalized_lines: List[Dict[str, Any]] = []
        for line in self._normalize_lines(lines):
            text = self.lowmem_pdf_service.normalize_text(str(line.get("text") or ""))
            if not text:
                continue
            normalized_lines.append(
                {
                    **line,
                    "text": text,
                    "anchor_text": text,
                    "layout_text": text,
                    "source": "ppocrv5_anchor",
                }
            )
        text = "\n".join(str(line.get("text") or "") for line in normalized_lines if str(line.get("text") or "").strip())
        confidence = round(self._avg_confidence(normalized_lines), 4)
        readable_chars = self.lowmem_pdf_service.count_readable_chars(text)
        low_confidence_lines = [
            line
            for line in normalized_lines
            if self._safe_float(line.get("confidence"), default=0.0) < 0.72
        ]
        low_confidence_ratio = len(low_confidence_lines) / max(len(normalized_lines), 1)
        abnormal_ratio = self._abnormal_char_ratio(text)
        min_chars = max(20, int(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MIN_CHARS", 120) or 120))
        min_confidence = float(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MIN_CONFIDENCE", 0.93) or 0.93)
        max_low_ratio = float(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MAX_LOW_CONF_RATIO", 0.10) or 0.10)
        max_abnormal_ratio = float(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MAX_ABNORMAL_RATIO", 0.025) or 0.025)
        if not text.strip():
            quality = "failed"
        elif (
            confidence >= min_confidence
            and readable_chars >= min_chars
            and low_confidence_ratio <= max_low_ratio
            and abnormal_ratio <= max_abnormal_ratio
            and "[UNREADABLE]" not in text.upper()
        ):
            quality = "high"
        elif confidence >= 0.72 or readable_chars >= min_chars:
            quality = "medium"
        else:
            quality = "low"
        page_warnings = ["ppocrv5_anchor_used"] if normalized_lines else ["ppocrv5_anchor_empty"]
        if quality != "high":
            page_warnings.append("ppocrv5_anchor_low_quality")
        return LocalPdfPageResult(
            page_number=page_number,
            width=float(classified_page.get("width") or 595.0),
            height=float(classified_page.get("height") or 842.0),
            text=text,
            source="ppocr_anchor",
            quality=quality,
            confidence=confidence,
            lines=normalized_lines,
            blocks=self._blocks_from_lines(normalized_lines, text),
            warnings=page_warnings,
            engine="ppocrv5_anchor_rapidocr",
        )

    def _fast_anchor_page_result(
        self,
        classified_page: Dict[str, Any],
        image: Image.Image,
        *,
        warnings: List[str],
    ) -> LocalPdfPageResult:
        return self.build_anchor_page_result(classified_page, image, warnings=warnings)

    def score_anchor_page(self, page_result: LocalPdfPageResult, classified_page: Dict[str, Any]) -> Dict[str, Any]:
        readable_chars = self.lowmem_pdf_service.count_readable_chars(page_result.text)
        line_count = len([line for line in page_result.lines if str(line.get("text") or "").strip()])
        low_confidence_count = sum(
            1
            for line in page_result.lines
            if str(line.get("text") or "").strip()
            and self._safe_float(line.get("confidence"), default=0.0) < 0.72
        )
        low_confidence_ratio = low_confidence_count / max(line_count, 1)
        abnormal_ratio = self._abnormal_char_ratio(page_result.text)
        complex_layout = self._page_has_complex_layout(page_result, classified_page)
        bbox_order_unstable = self._bbox_order_unstable(page_result.lines)
        has_unreadable = "[UNREADABLE]" in str(page_result.text or "").upper()
        min_confidence = float(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MIN_CONFIDENCE", 0.93) or 0.93)
        min_chars = max(20, int(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MIN_CHARS", 120) or 120))
        max_low_ratio = float(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MAX_LOW_CONF_RATIO", 0.10) or 0.10)
        max_abnormal_ratio = float(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MAX_ABNORMAL_RATIO", 0.025) or 0.025)
        reasons: List[str] = []
        if page_result.quality != "high":
            reasons.append("anchor_not_high")
        if page_result.confidence < min_confidence:
            reasons.append("anchor_confidence_low")
        if readable_chars < min_chars:
            reasons.append("anchor_chars_low")
        if low_confidence_ratio > max_low_ratio:
            reasons.append("anchor_low_conf_ratio_high")
        if abnormal_ratio > max_abnormal_ratio:
            reasons.append("anchor_abnormal_ratio_high")
        if complex_layout:
            reasons.append("complex_layout")
        if bbox_order_unstable:
            reasons.append("bbox_order_unstable")
        if has_unreadable:
            reasons.append("unreadable_marker")
        can_skip_glm = not reasons
        return {
            "confidence": round(float(page_result.confidence or 0.0), 4),
            "readable_chars": readable_chars,
            "line_count": line_count,
            "low_confidence_line_count": low_confidence_count,
            "low_confidence_line_ratio": round(low_confidence_ratio, 4),
            "abnormal_char_ratio": round(abnormal_ratio, 4),
            "complex_layout": complex_layout,
            "bbox_order_unstable": bbox_order_unstable,
            "has_unreadable_marker": has_unreadable,
            "can_skip_glm": can_skip_glm,
            "reasons": reasons,
        }

    def classify_page_ocr_tier(self, page_result: LocalPdfPageResult, classified_page: Dict[str, Any]) -> str:
        policy = str(getattr(settings, "LOCAL_PDF_GLM_POLICY", "quality_gate") or "quality_gate").strip().lower()
        if page_result.source == "native":
            return "native_text"
        if policy in {"off", "disabled", "none", "false"}:
            return "fast_verified" if page_result.text.strip() else "failed"
        if policy == "always":
            return "glm_page"
        anchor_score = self.score_anchor_page(page_result, classified_page)
        if anchor_score.get("can_skip_glm"):
            return "fast_verified"
        if not page_result.text.strip() or page_result.quality in {"failed", "low"}:
            return "glm_page"
        scope = str(getattr(settings, "LOCAL_PDF_GLM_SCOPE", "page_or_block") or "page_or_block").strip().lower()
        if scope in {"block", "page_or_block", "adaptive"} and self.select_glm_blocks(page_result, classified_page):
            return "glm_block"
        return "glm_page"

    def _should_run_glm_main(self, page_result: LocalPdfPageResult, classified_page: Dict[str, Any]) -> bool:
        return self.classify_page_ocr_tier(page_result, classified_page) in {"glm_page", "glm_block"}

    def _bbox_order_unstable(self, lines: List[Dict[str, Any]]) -> bool:
        previous_top = -1.0
        inversions = 0
        usable = 0
        for line in lines:
            bbox = self._normalize_bbox(line.get("bbox"))
            text = str(line.get("text") or "").strip()
            if not text:
                continue
            usable += 1
            top = float(bbox[1])
            if previous_top >= 0 and top + 0.01 < previous_top:
                inversions += 1
            previous_top = max(previous_top, top)
        return usable >= 4 and inversions / max(usable - 1, 1) > 0.10

    def select_glm_blocks(self, page_result: LocalPdfPageResult, classified_page: Dict[str, Any]) -> List[Dict[str, Any]]:
        scope = str(getattr(settings, "LOCAL_PDF_GLM_SCOPE", "page_or_block") or "page_or_block").strip().lower()
        if scope not in {"block", "page_or_block", "adaptive"}:
            return []
        if not page_result.lines:
            return []
        complex_layout = self._page_has_complex_layout(page_result, classified_page)
        selected: List[Dict[str, Any]] = []
        for index, line in enumerate(page_result.lines):
            text = str(line.get("text") or "").strip()
            if not text:
                continue
            confidence = self._safe_float(line.get("confidence"), default=0.0)
            bbox = self._normalize_bbox(line.get("bbox"))
            needs_block = confidence < 0.86 or "[UNREADABLE]" in text.upper()
            if complex_layout and confidence < 0.94:
                needs_block = True
            if needs_block:
                selected.append(
                    {
                        "index": index,
                        "bbox": bbox,
                        "reason": "low_confidence_line" if confidence < 0.86 else "complex_layout_line",
                        "text": text,
                    }
                )
        return selected[:6]

    def run_glm_for_page_or_blocks(
        self,
        *,
        glm_client: GlmOcrOllamaClient,
        image: Image.Image,
        current: Optional[LocalPdfPageResult],
        classified_page: Dict[str, Any],
        page_number: int,
        total_pages: int,
        tier: str,
    ) -> LocalPdfPageResult:
        if tier != "glm_block" or current is None:
            return self._run_glm_for_full_page(
                glm_client=glm_client,
                image=image,
                classified_page=classified_page,
                page_number=page_number,
                total_pages=total_pages,
            )

        selected_blocks = self.select_glm_blocks(current, classified_page)
        if not selected_blocks:
            return self._run_glm_for_full_page(
                glm_client=glm_client,
                image=image,
                classified_page=classified_page,
                page_number=page_number,
                total_pages=total_pages,
            )

        block_results: List[Dict[str, Any]] = []
        for block in selected_blocks:
            crop = self._crop_image_by_bbox(image, block.get("bbox"))
            try:
                image_bytes = self._enhanced_image_bytes(crop)
                raw = glm_client.recognize(image_bytes, page_number=page_number, total_pages=total_pages)
            finally:
                if "image_bytes" in locals():
                    del image_bytes
                del crop
            block_page = self._page_result_from_ocr_payload(
                raw,
                classified_page=classified_page,
                source="glm_ocr",
                engine="glm_ocr_ollama+block",
            )
            block_results.append({"selection": block, "result": block_page})

        return self.merge_best_text_into_layout_ir(current, block_results)

    def _run_glm_for_full_page(
        self,
        *,
        glm_client: GlmOcrOllamaClient,
        image: Image.Image,
        classified_page: Dict[str, Any],
        page_number: int,
        total_pages: int,
    ) -> LocalPdfPageResult:
        try:
            image_bytes = self._enhanced_image_bytes(image)
            raw_glm = glm_client.recognize(image_bytes, page_number=page_number, total_pages=total_pages)
        finally:
            if "image_bytes" in locals():
                del image_bytes
        return self._page_result_from_ocr_payload(
            raw_glm,
            classified_page=classified_page,
            source="glm_ocr",
            engine="glm_ocr_ollama",
        )

    def _crop_image_by_bbox(self, image: Image.Image, bbox_value: Any, *, margin: float = 0.025) -> Image.Image:
        bbox = self._normalize_bbox(bbox_value)
        width, height = image.size
        left = max(0, int((bbox[0] - margin) * width))
        top = max(0, int((bbox[1] - margin) * height))
        right = min(width, int((bbox[2] + margin) * width))
        bottom = min(height, int((bbox[3] + margin) * height))
        if right <= left or bottom <= top:
            return image.copy()
        return image.crop((left, top, right, bottom))

    def merge_best_text_into_layout_ir(
        self,
        anchor_page: LocalPdfPageResult,
        block_results: List[Dict[str, Any]],
    ) -> LocalPdfPageResult:
        if not block_results:
            return anchor_page
        lines = [dict(line) for line in anchor_page.lines]
        warnings = [*(anchor_page.warnings or []), "glm_block_used"]
        adopted = 0
        for item in block_results:
            selection = dict(item.get("selection") or {})
            block_result = item.get("result")
            if not isinstance(block_result, LocalPdfPageResult):
                continue
            text = self.lowmem_pdf_service.normalize_text(str(block_result.text or ""))
            line_index = int(selection.get("index") if selection.get("index") is not None else -1)
            if not text or line_index < 0 or line_index >= len(lines):
                continue
            current_text = str(lines[line_index].get("text") or "")
            current_confidence = self._safe_float(lines[line_index].get("confidence"), default=0.0)
            if (
                block_result.quality in {"high", "medium"}
                and self._abnormal_char_ratio(text) <= self._abnormal_char_ratio(current_text)
                and (block_result.confidence >= current_confidence or block_result.quality == "high")
            ):
                lines[line_index] = {
                    **lines[line_index],
                    "text": text.replace("\n", " "),
                    "confidence": round(max(block_result.confidence, current_confidence), 4),
                    "source": "glm_block_text_anchor_layout",
                    "anchor_text": current_text,
                    "retry_confidence_delta": round(block_result.confidence - current_confidence, 4),
                }
                adopted += 1
        if adopted <= 0:
            warnings.append("glm_block_rejected")
            return anchor_page
        text = "\n".join(str(line.get("text") or "") for line in lines if str(line.get("text") or "").strip())
        confidence = round(self._avg_confidence(lines), 4)
        quality = self._normalize_quality("high" if confidence >= 0.9 else "medium", text=text, confidence=confidence)
        warnings.append("glm_block_text_merged")
        return LocalPdfPageResult(
            page_number=anchor_page.page_number,
            width=anchor_page.width,
            height=anchor_page.height,
            text=text,
            source="glm_ocr",
            quality=quality,
            confidence=confidence,
            lines=lines,
            blocks=self._blocks_from_lines(lines, text),
            warnings=sorted(set(warnings)),
            engine="glm_ocr_ollama+block+anchor_layout",
            used_vl_fallback=anchor_page.used_vl_fallback,
        )

    def compare_anchor_glm_quality(
        self,
        *,
        anchor: Optional[LocalPdfPageResult],
        glm_result: LocalPdfPageResult,
        classified_page: Dict[str, Any],
    ) -> tuple[LocalPdfPageResult, bool]:
        if anchor is None or not anchor.text.strip():
            return glm_result, False
        glm_metrics = self._page_quality_metrics(glm_result, classified_page)
        reject_reasons: List[str] = []
        if glm_result.quality == "failed" or not glm_result.text.strip():
            reject_reasons.append("glm_failed")
        if isinstance(glm_metrics.get("anchor_recall"), (int, float)) and float(glm_metrics["anchor_recall"]) < 0.85:
            reject_reasons.append("low_anchor_recall")
        if isinstance(glm_metrics.get("char_count_ratio"), (int, float)) and float(glm_metrics["char_count_ratio"]) < 0.75:
            reject_reasons.append("low_char_count_ratio")
        if float(glm_metrics.get("abnormal_char_ratio") or 0.0) >= 0.08:
            reject_reasons.append("abnormal_chars")
        if float(glm_metrics.get("duplicate_ratio") or 0.0) >= 0.18:
            reject_reasons.append("duplicate_text")
        if "[UNREADABLE]" in str(glm_result.text or "").upper():
            reject_reasons.append("unreadable_marker")

        if reject_reasons or (
            anchor.quality == "high"
            and not self._is_better_page_result(anchor, glm_result)
            and self._page_result_score(anchor) >= self._page_result_score(glm_result)
        ):
            warnings = [*(anchor.warnings or []), "glm_quality_gate_rejected"]
            if reject_reasons:
                warnings.extend(f"glm_reject_{reason}" for reason in reject_reasons)
            return LocalPdfPageResult(
                page_number=anchor.page_number,
                width=anchor.width,
                height=anchor.height,
                text=anchor.text,
                source=anchor.source,
                quality=anchor.quality,
                confidence=anchor.confidence,
                lines=anchor.lines,
                blocks=anchor.blocks,
                warnings=sorted(set(warnings)),
                engine=anchor.engine,
                used_vl_fallback=anchor.used_vl_fallback,
            ), True
        return glm_result, False

    def _should_run_glm_main_legacy(self, page_result: LocalPdfPageResult, classified_page: Dict[str, Any]) -> bool:
        policy = str(getattr(settings, "LOCAL_PDF_GLM_POLICY", "quality_gate") or "quality_gate").strip().lower()
        if policy in {"off", "disabled", "none", "false"}:
            return False
        if policy == "always":
            return True
        if page_result.source == "native":
            return False
        if page_result.quality != "high":
            return True
        min_confidence = float(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MIN_CONFIDENCE", 0.90) or 0.90)
        min_chars = max(20, int(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MIN_CHARS", 80) or 80))
        if page_result.confidence < min_confidence:
            return True
        if self.lowmem_pdf_service.count_readable_chars(page_result.text) < min_chars:
            return True
        if not self._page_has_layout_signal(page_result):
            return True
        if self._abnormal_char_ratio(page_result.text) >= 0.04:
            return True
        if "[UNREADABLE]" in str(page_result.text or "").upper():
            return True
        if self._page_has_complex_layout(page_result, classified_page):
            return True
        return False

    def _page_has_complex_layout(self, page_result: LocalPdfPageResult, classified_page: Dict[str, Any]) -> bool:
        block_types = {str(block.get("type") or "").lower() for block in page_result.blocks}
        if {"table", "formula", "chart"} & block_types:
            return True
        if str(classified_page.get("kind") or "").lower() == "mixed_pdf":
            return True
        lines = [line for line in page_result.lines if str(line.get("text") or "").strip()]
        if len(lines) < 2:
            return False
        left_values = [self._safe_float((line.get("bbox") or [0, 0, 0, 0])[0], default=0.0) for line in lines]
        right_values = [self._safe_float((line.get("bbox") or [0, 0, 0, 0])[2], default=0.0) for line in lines]
        left_spread = max(left_values) - min(left_values)
        right_spread = max(right_values) - min(right_values)
        # A broad spread on both edges usually means multi-column or floating blocks.
        return left_spread > 0.22 and right_spread > 0.22

    def _mark_anchor_output_used(self, page_result: LocalPdfPageResult) -> LocalPdfPageResult:
        warnings = [*(page_result.warnings or []), "ppocr_anchor_output_used"]
        return LocalPdfPageResult(
            page_number=page_result.page_number,
            width=page_result.width,
            height=page_result.height,
            text=page_result.text,
            source=page_result.source,
            quality=page_result.quality,
            confidence=page_result.confidence,
            lines=page_result.lines,
            blocks=page_result.blocks,
            warnings=sorted(set(warnings)),
            engine=page_result.engine,
            used_vl_fallback=page_result.used_vl_fallback,
        )

    def _anchor_fallback_after_glm_failure(
        self,
        current: Optional[LocalPdfPageResult],
        classified_page: Dict[str, Any],
    ) -> LocalPdfPageResult:
        if current is not None and current.text.strip():
            warnings = [*(current.warnings or []), "glm_ocr_failed_anchor_fallback"]
            return LocalPdfPageResult(
                page_number=current.page_number,
                width=current.width,
                height=current.height,
                text=current.text,
                source=current.source,
                quality=current.quality,
                confidence=current.confidence,
                lines=current.lines,
                blocks=current.blocks,
                warnings=sorted(set(warnings)),
                engine=current.engine,
                used_vl_fallback=current.used_vl_fallback,
            )
        return LocalPdfPageResult(
            page_number=int(classified_page.get("page_number") or 0),
            width=float(classified_page.get("width") or 595.0),
            height=float(classified_page.get("height") or 842.0),
            text="",
            source="ocr_failed",
            quality="failed",
            confidence=0.0,
            lines=[],
            blocks=[],
            warnings=["glm_ocr_failed"],
        )

    def _merge_text_with_existing_layout(
        self,
        candidate: LocalPdfPageResult,
        layout_page: LocalPdfPageResult,
        *,
        warning: str,
        engine_suffix: str,
    ) -> LocalPdfPageResult:
        if not candidate.text.strip() or not self._page_has_layout_signal(layout_page):
            return candidate
        merged = self._merge_glm_text_with_layout_blocks(
            candidate,
            layout_page.blocks or self._blocks_from_lines(layout_page.lines, layout_page.text),
        )
        if merged is candidate and layout_page.lines:
            merged = self._merge_glm_text_with_layout_lines(candidate, layout_page.lines)
        if merged is candidate:
            return candidate
        warnings = [*(merged.warnings or []), warning]
        confidence = max(merged.confidence, candidate.confidence, layout_page.confidence)
        quality = merged.quality
        if (
            quality in {"failed", "low", "medium"}
            and confidence >= float(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MIN_CONFIDENCE", 0.90) or 0.90)
            and self.lowmem_pdf_service.count_readable_chars(merged.text) >= int(getattr(settings, "LOCAL_PDF_FAST_ANCHOR_MIN_CHARS", 80) or 80)
            and self._abnormal_char_ratio(merged.text) < 0.04
        ):
            quality = "high"
        return LocalPdfPageResult(
            page_number=merged.page_number,
            width=merged.width,
            height=merged.height,
            text=merged.text,
            source=merged.source,
            quality=quality,
            confidence=round(confidence, 4),
            lines=merged.lines,
            blocks=merged.blocks,
            warnings=sorted(set(warnings)),
            engine=f"{candidate.engine}+{engine_suffix}",
            used_vl_fallback=candidate.used_vl_fallback,
        )

    def _try_rapidocr_layout_assist(self, page_result: LocalPdfPageResult, image: Image.Image) -> LocalPdfPageResult:
        warnings = set(page_result.warnings or [])
        if self._page_has_layout_signal(page_result) and "glm_ocr_non_json_response" not in warnings:
            return page_result
        try:
            layout_lines = self._detect_rapidocr_layout_lines(image)
        except Exception as exc:
            logger.debug("RapidOCR layout assist unavailable: %s", exc, exc_info=True)
            return page_result
        if not layout_lines:
            return page_result
        return self._merge_glm_text_with_layout_lines(page_result, layout_lines)

    def _detect_rapidocr_layout_lines(self, image: Image.Image) -> List[Dict[str, Any]]:
        if self._rapid_layout_ocr is None:
            self._rapid_layout_ocr = RapidOCROnNXService(
                enable_angle_cls=False,
                max_side_len=max(int(settings.LOCAL_PDF_RENDER_DPI or 200) * 12, int(settings.PDF_NORMALIZE_MAX_IMAGE_EDGE)),
            )
        import numpy as np

        # Layout coordinates must stay relative to the original rendered page.
        # The OCR image enhancer may crop borders, which is useful for text
        # recognition but would shift bboxes and degrade DOCX reconstruction.
        array = np.array(ImageOps.exif_transpose(image.convert("RGB")))
        result = self._rapid_layout_ocr.recognize(array, use_cls=False)
        return self._normalize_lines(result.get("lines"))

    def _merge_glm_text_with_layout_lines(
        self,
        page_result: LocalPdfPageResult,
        layout_lines: List[Dict[str, Any]],
    ) -> LocalPdfPageResult:
        text_lines = [
            self.lowmem_pdf_service.normalize_text(line)
            for line in str(page_result.text or "").splitlines()
            if self.lowmem_pdf_service.normalize_text(line)
        ]
        if not text_lines or not layout_lines:
            return page_result

        matched_layout = self._match_text_lines_to_layout(text_lines, layout_lines)
        matched_count = sum(1 for item in matched_layout if item is not None)
        coverage = matched_count / max(len(text_lines), 1)
        if coverage < 0.45:
            return page_result

        new_lines: List[Dict[str, Any]] = []
        for index, text in enumerate(text_lines):
            layout_line = matched_layout[index]
            fallback_top = min(0.96, 0.07 + index * 0.035)
            if layout_line is None:
                bbox = [0.08, round(fallback_top, 5), 0.92, round(min(0.99, fallback_top + 0.025), 5)]
                confidence = 0.0
                layout_text = ""
            else:
                bbox = layout_line.get("bbox") or [0.08, round(fallback_top, 5), 0.92, round(min(0.99, fallback_top + 0.025), 5)]
                confidence = self._safe_float(layout_line.get("confidence"), default=0.0)
                layout_text = str(layout_line.get("text") or "")
            new_lines.append(
                {
                    "text": text,
                    "bbox": bbox,
                    "confidence": round(confidence, 4),
                    "source": "glm_text_rapidocr_layout",
                    "layout_text": layout_text,
                }
            )

        avg_confidence = self._avg_confidence(new_lines)
        new_quality = page_result.quality
        readable_chars = self.lowmem_pdf_service.count_readable_chars(page_result.text)
        min_readable_for_page = max(12, min(40, len(text_lines) * 8))
        if new_quality in {"failed", "low"} and coverage >= 0.6 and readable_chars >= min_readable_for_page:
            new_quality = "medium"
        if new_quality == "medium" and avg_confidence >= 0.86 and coverage >= 0.9:
            new_quality = "high"

        new_warnings = list(page_result.warnings or [])
        new_warnings.append("rapidocr_layout_assist_used")
        if coverage < 0.75:
            new_warnings.append("rapidocr_layout_assist_partial")

        return LocalPdfPageResult(
            page_number=page_result.page_number,
            width=page_result.width,
            height=page_result.height,
            text=page_result.text,
            source=page_result.source,
            quality=new_quality,
            confidence=round(avg_confidence, 4),
            lines=new_lines,
            blocks=self._blocks_from_lines(new_lines, page_result.text),
            warnings=sorted(set(new_warnings)),
            engine=page_result.engine,
            used_vl_fallback=page_result.used_vl_fallback,
        )

    def _match_text_lines_to_layout(
        self,
        text_lines: List[str],
        layout_lines: List[Dict[str, Any]],
    ) -> List[Optional[Dict[str, Any]]]:
        if not text_lines:
            return []
        if not layout_lines:
            return [None for _ in text_lines]

        count_delta = abs(len(text_lines) - len(layout_lines))
        if count_delta <= max(2, int(max(len(text_lines), len(layout_lines)) * 0.25)):
            return [layout_lines[index] if index < len(layout_lines) else None for index in range(len(text_lines))]

        used: set[int] = set()
        matched: List[Optional[Dict[str, Any]]] = []
        for text in text_lines:
            compact_text = self._compact_layout_match_text(text)
            best_index: Optional[int] = None
            best_score = 0.0
            for index, layout_line in enumerate(layout_lines):
                if index in used:
                    continue
                layout_text = self._compact_layout_match_text(str(layout_line.get("text") or ""))
                if not compact_text or not layout_text:
                    continue
                score = difflib.SequenceMatcher(None, compact_text, layout_text).ratio()
                if compact_text in layout_text or layout_text in compact_text:
                    score = max(score, 0.82)
                if score > best_score:
                    best_index = index
                    best_score = score
            if best_index is not None and best_score >= 0.42:
                used.add(best_index)
                matched.append(layout_lines[best_index])
            else:
                matched.append(None)
        return matched

    @staticmethod
    def _compact_layout_match_text(text: str) -> str:
        return re.sub(r"\s+", "", str(text or "")).strip()

    def _save_frontline_page_image(self, image: Image.Image, directory: Path, page_number: int) -> Path:
        ensure_private_directory(directory)
        image_path = directory / f"page_{int(page_number):04d}.jpg"
        image.convert("RGB").save(image_path, format="JPEG", quality=92, optimize=True)
        ensure_private_file(image_path)
        return image_path

    def _apply_ppstructurev3_batch_layout(
        self,
        *,
        result_pages: List[LocalPdfPageResult],
        rendered_image_paths: Dict[int, Path],
        warnings: List[str],
        use_doc_unwarping: Optional[bool] = None,
    ) -> None:
        if not result_pages or not rendered_image_paths:
            return
        if use_doc_unwarping is None:
            batch_probe = self._run_ppstructurev3_layout_batch_sidecar(rendered_image_paths)
        else:
            batch_probe = self._run_ppstructurev3_layout_batch_sidecar(
                rendered_image_paths,
                use_doc_unwarping=use_doc_unwarping,
            )
        page_payloads = batch_probe.get("pages") if isinstance(batch_probe, dict) else None
        if not isinstance(page_payloads, list):
            warning = str(batch_probe.get("warning") or "") if isinstance(batch_probe, dict) else ""
            if warning:
                warnings.append(warning)
            return
        probes_by_page = {
            int(item.get("page") or 0): item
            for item in page_payloads
            if isinstance(item, dict) and int(item.get("page") or 0) > 0
        }
        for index, page_result in enumerate(list(result_pages)):
            if page_result.source == "native":
                continue
            image_path = rendered_image_paths.get(page_result.page_number)
            probe = probes_by_page.get(page_result.page_number)
            if not image_path or not probe:
                continue
            if not probe.get("available"):
                warning = str(probe.get("warning") or "").strip()
                if warning:
                    warnings.append(warning)
                continue
            try:
                with Image.open(image_path) as image:
                    merged = self._apply_ppstructurev3_probe_to_page(
                        page_result,
                        probe,
                        image_width=image.width,
                        image_height=image.height,
                        warnings=warnings,
                    )
            except Exception:
                logger.debug("Failed to apply PP-StructureV3 batch layout on page %s.", page_result.page_number, exc_info=True)
                continue
            if merged is not page_result:
                result_pages[index] = merged

    def _apply_ppstructurev3_probe_to_page(
        self,
        page_result: LocalPdfPageResult,
        probe: Dict[str, Any],
        *,
        image_width: int,
        image_height: int,
        warnings: List[str],
    ) -> LocalPdfPageResult:
        blocks = self._blocks_from_ppstructurev3_payload(
            probe.get("raw"),
            image_width=image_width,
            image_height=image_height,
        )
        if not blocks:
            warnings.append("ppstructurev3_layout_low")
            return page_result
        merged = self._merge_glm_text_with_layout_blocks(page_result, blocks)
        if merged is page_result:
            warnings.append("ppstructurev3_layout_low")
            return page_result
        merged.engine = f"{page_result.engine}+ppstructurev3"
        page_warnings = list(merged.warnings or [])
        page_warnings.append("ppstructurev3_layout_used")
        merged.warnings = sorted(set(page_warnings))
        warnings.append("ppstructurev3_layout_used")
        return merged

    def _try_ppstructurev3_layout(
        self,
        page_result: LocalPdfPageResult,
        image: Image.Image,
        *,
        warnings: List[str],
    ) -> LocalPdfPageResult:
        if not settings.LOCAL_PDF_PPSTRUCTUREV3_ENABLED:
            return page_result
        policy = str(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_POLICY", "always") or "always").strip().lower()
        if (
            policy != "always"
            and self._page_has_layout_signal(page_result)
            and "glm_ocr_non_json_response" not in set(page_result.warnings or [])
        ):
            return page_result
        probe = self._run_ppstructurev3_layout_sidecar(image)
        if not probe.get("available"):
            warning = str(probe.get("warning") or "ppstructurev3_layout_missing")
            if warning:
                warnings.append(warning)
            return page_result
        return self._apply_ppstructurev3_probe_to_page(
            page_result,
            probe,
            image_width=image.width,
            image_height=image.height,
            warnings=warnings,
        )

    def _run_ppstructurev3_layout_batch_sidecar(
        self,
        image_paths: Dict[int, Path],
        *,
        use_doc_unwarping: Optional[bool] = None,
    ) -> Dict[str, Any]:
        if not settings.LOCAL_PDF_PPSTRUCTUREV3_ENABLED:
            return {"available": False, "warning": "ppstructurev3_layout_missing", "error": "disabled", "pages": []}
        policy = str(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_POLICY", "always") or "always").strip().lower()
        if policy in {"off", "disabled", "none", "false"}:
            return {"available": False, "warning": "ppstructurev3_layout_missing", "error": "disabled", "pages": []}
        mode = str(settings.LOCAL_PDF_PPSTRUCTUREV3_API_MODE or "auto").strip().lower()
        command = str(settings.LOCAL_PDF_PPSTRUCTUREV3_COMMAND or "").strip()
        config_path = str(settings.LOCAL_PDF_PPSTRUCTUREV3_CONFIG or "").strip()
        if not command and not config_path:
            return {
                "available": False,
                "warning": "ppstructurev3_layout_missing",
                "error": "ppstructurev3_config_or_command_missing",
                "pages": [],
            }
        python_path = self._resolve_ocr_python()
        if not python_path:
            return {"available": False, "warning": "ppstructurev3_layout_missing", "error": "ocr_python_missing", "pages": []}

        first_path = next(iter(image_paths.values()), None)
        base_dir = first_path.parent if first_path else Path(tempfile.gettempdir())
        input_path = self._temp_json_file(base_dir, "ppstructurev3-batch-in-")
        output_path = self._temp_json_file(base_dir, "ppstructurev3-batch-out-")
        try:
            payload = {
                **ppstructurev3_profile(
                    mode="photo" if bool(use_doc_unwarping) else "document_table",
                    threads=max(1, int(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_THREADS", 2) or 2)),
                ),
                "mode": mode,
                "command": command,
                "config_path": config_path,
                "timeout": int(settings.LOCAL_PDF_PPSTRUCTUREV3_TIMEOUT),
                "use_doc_orientation_classify": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_DOC_ORIENTATION", True)),
                "use_doc_unwarping": bool(use_doc_unwarping)
                if use_doc_unwarping is not None
                else _auto_bool_setting(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_DOC_UNWARPING", False), auto_default=False),
                "use_textline_orientation": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_TEXTLINE_ORIENTATION", True)),
                "use_table_recognition": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_TABLE_RECOGNITION", True)),
                "use_region_detection": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_REGION_DETECTION", True)),
                "use_seal_recognition": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_SEAL_RECOGNITION", False)),
                "use_formula_recognition": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_FORMULA_RECOGNITION", False)),
                "use_chart_recognition": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_CHART_RECOGNITION", False)),
                "threads": max(1, int(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_THREADS", 2) or 2)),
                "pages": [
                    {"page": int(page_number), "image_path": str(path)}
                    for page_number, path in sorted(image_paths.items())
                ],
            }
            input_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            ensure_private_file(input_path)
            env = os.environ.copy()
            threads = str(payload["threads"])
            env.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
            env["OMP_NUM_THREADS"] = threads
            env["OPENBLAS_NUM_THREADS"] = threads
            env["MKL_NUM_THREADS"] = threads
            env["VECLIB_MAXIMUM_THREADS"] = threads
            existing_pythonpath = str(env.get("PYTHONPATH") or "").strip()
            env["PYTHONPATH"] = (
                str(SOURCE_BACKEND_ROOT)
                if not existing_pythonpath
                else f"{SOURCE_BACKEND_ROOT}{os.pathsep}{existing_pythonpath}"
            )
            result = subprocess.run(
                [
                    python_path,
                    "-m",
                    "app.workers.local_ppstructurev3_sidecar",
                    str(input_path),
                    str(output_path),
                ],
                cwd=str(SOURCE_BACKEND_ROOT),
                env=env,
                capture_output=True,
                text=True,
                timeout=max(30, int(settings.LOCAL_PDF_PPSTRUCTUREV3_TIMEOUT)),
                check=False,
            )
            output_text = output_path.read_text(encoding="utf-8") if output_path.exists() else ""
            if output_text.strip():
                try:
                    parsed = json.loads(output_text)
                except Exception:
                    parsed = None
                if isinstance(parsed, dict):
                    return parsed
            return {
                "available": False,
                "warning": "ppstructurev3_layout_failed",
                "error": (result.stderr or result.stdout or output_text)[-1000:],
                "pages": [],
            }
        except Exception as exc:
            return {"available": False, "warning": "ppstructurev3_layout_failed", "error": str(exc), "pages": []}
        finally:
            for path in (input_path, output_path):
                try:
                    path.unlink(missing_ok=True)
                except Exception:
                    pass

    def _run_ppstructurev3_layout_sidecar(
        self,
        image: Image.Image,
        *,
        use_doc_unwarping: Optional[bool] = None,
    ) -> Dict[str, Any]:
        mode = str(settings.LOCAL_PDF_PPSTRUCTUREV3_API_MODE or "auto").strip().lower()
        if mode == "off":
            return {"available": False, "warning": "ppstructurev3_layout_missing", "error": "disabled"}
        command = str(settings.LOCAL_PDF_PPSTRUCTUREV3_COMMAND or "").strip()
        config_path = str(settings.LOCAL_PDF_PPSTRUCTUREV3_CONFIG or "").strip()
        if not command and not config_path:
            return {
                "available": False,
                "warning": "ppstructurev3_layout_missing",
                "error": "ppstructurev3_config_or_command_missing",
            }
        python_path = self._resolve_ocr_python()
        if not python_path:
            return {"available": False, "warning": "ppstructurev3_layout_missing", "error": "ocr_python_missing"}

        image_path = self._temp_image_file()
        input_path = self._temp_json_file(image_path.parent, "ppstructurev3-in-")
        output_path = self._temp_json_file(image_path.parent, "ppstructurev3-out-")
        try:
            image.convert("RGB").save(image_path, format="JPEG", quality=92, optimize=True)
            ensure_private_file(image_path)
            payload = {
                **ppstructurev3_profile(
                    mode="photo" if bool(use_doc_unwarping) else "document_table",
                    threads=max(1, int(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_THREADS", 2) or 2)),
                ),
                "image_path": str(image_path),
                "mode": mode,
                "command": command,
                "config_path": config_path,
                "timeout": int(settings.LOCAL_PDF_PPSTRUCTUREV3_TIMEOUT),
                "use_doc_orientation_classify": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_DOC_ORIENTATION", True)),
                "use_doc_unwarping": bool(use_doc_unwarping)
                if use_doc_unwarping is not None
                else _auto_bool_setting(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_DOC_UNWARPING", False), auto_default=False),
                "use_textline_orientation": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_TEXTLINE_ORIENTATION", True)),
                "use_table_recognition": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_TABLE_RECOGNITION", True)),
                "use_region_detection": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_REGION_DETECTION", True)),
                "use_seal_recognition": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_SEAL_RECOGNITION", False)),
                "use_formula_recognition": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_FORMULA_RECOGNITION", False)),
                "use_chart_recognition": bool(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_CHART_RECOGNITION", False)),
                "threads": max(1, int(getattr(settings, "LOCAL_PDF_PPSTRUCTUREV3_THREADS", 2) or 2)),
            }
            input_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            ensure_private_file(input_path)
            env = os.environ.copy()
            threads = str(payload["threads"])
            env.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
            env["OMP_NUM_THREADS"] = threads
            env["OPENBLAS_NUM_THREADS"] = threads
            env["MKL_NUM_THREADS"] = threads
            env["VECLIB_MAXIMUM_THREADS"] = threads
            existing_pythonpath = str(env.get("PYTHONPATH") or "").strip()
            env["PYTHONPATH"] = (
                str(SOURCE_BACKEND_ROOT)
                if not existing_pythonpath
                else f"{SOURCE_BACKEND_ROOT}{os.pathsep}{existing_pythonpath}"
            )
            result = subprocess.run(
                [
                    python_path,
                    "-m",
                    "app.workers.local_ppstructurev3_sidecar",
                    str(input_path),
                    str(output_path),
                ],
                cwd=str(SOURCE_BACKEND_ROOT),
                env=env,
                capture_output=True,
                text=True,
                timeout=max(30, int(settings.LOCAL_PDF_PPSTRUCTUREV3_TIMEOUT)),
                check=False,
            )
            output_text = output_path.read_text(encoding="utf-8") if output_path.exists() else ""
            if output_text.strip():
                try:
                    parsed = json.loads(output_text)
                except Exception:
                    parsed = None
                if isinstance(parsed, dict):
                    return parsed
            return {
                "available": False,
                "warning": "ppstructurev3_layout_failed",
                "error": (result.stderr or result.stdout or output_text)[-1000:],
            }
        except Exception as exc:
            return {"available": False, "warning": "ppstructurev3_layout_failed", "error": str(exc)}
        finally:
            for path in (image_path, input_path, output_path):
                try:
                    path.unlink(missing_ok=True)
                except Exception:
                    pass

    @staticmethod
    def _temp_image_file() -> Path:
        handle = tempfile.NamedTemporaryFile(
            mode="wb",
            prefix="ppstructurev3-page-",
            suffix=".jpg",
            delete=False,
        )
        handle.close()
        return Path(handle.name)

    def _blocks_from_ppstructurev3_payload(
        self,
        payload: Any,
        *,
        image_width: int,
        image_height: int,
    ) -> List[Dict[str, Any]]:
        blocks: List[Dict[str, Any]] = []

        def collect(node: Any) -> None:
            if isinstance(node, dict):
                direct = self._ppstructure_node_to_block(node, image_width=image_width, image_height=image_height)
                if direct is not None:
                    blocks.append(direct)
                for value in node.values():
                    collect(value)
            elif isinstance(node, list):
                for value in node:
                    collect(value)

        collect(payload)
        deduped: List[Dict[str, Any]] = []
        seen: set[tuple[str, str, tuple[float, ...]]] = set()
        for block in sorted(blocks, key=lambda item: ((item.get("bbox") or [0, 0, 0, 0])[1], (item.get("bbox") or [0, 0, 0, 0])[0])):
            key = (
                str(block.get("type") or ""),
                str(block.get("text") or ""),
                tuple(float(value) for value in block.get("bbox") or []),
            )
            if key in seen:
                continue
            seen.add(key)
            deduped.append(block)
        return deduped

    def _ppstructure_node_to_block(
        self,
        node: Dict[str, Any],
        *,
        image_width: int,
        image_height: int,
    ) -> Optional[Dict[str, Any]]:
        bbox = self._extract_ppstructure_bbox(node, image_width=image_width, image_height=image_height)
        text = self.lowmem_pdf_service.normalize_text(
            str(
                node.get("text")
                or node.get("block_content")
                or node.get("content")
                or node.get("rec_text")
                or node.get("markdown")
                or ""
            )
        )
        rows = node.get("rows")
        html = str(node.get("html") or node.get("table_html") or "")
        block_type = str(
            node.get("type")
            or node.get("label")
            or node.get("block_label")
            or node.get("category")
            or ""
        ).strip().lower()
        if "table" in block_type:
            normalized_rows = self._normalize_table_rows(rows)
            if normalized_rows:
                return {"type": "table", "rows": normalized_rows, "bbox": bbox or [0.08, 0.0, 0.92, 0.02]}
            if text:
                table_rows = self._markdown_table_rows(text)
                if table_rows:
                    return {"type": "table", "rows": table_rows, "bbox": bbox or [0.08, 0.0, 0.92, 0.02]}
            if html:
                table_rows = self._html_table_rows(html)
                if table_rows:
                    return {"type": "table", "rows": table_rows, "bbox": bbox or [0.08, 0.0, 0.92, 0.02]}
        if not bbox or not text:
            return None
        normalized_type = "title" if block_type in {"title", "header", "doc_title"} else "line"
        if block_type in {"paragraph", "text", "content", "body"}:
            normalized_type = "paragraph"
        return {
            "type": normalized_type,
            "text": text,
            "align": self._alignment_from_bbox(bbox),
            "bbox": bbox,
        }

    def _extract_ppstructure_bbox(
        self,
        node: Dict[str, Any],
        *,
        image_width: int,
        image_height: int,
    ) -> Optional[List[float]]:
        for key in ("bbox", "box", "coordinate", "block_bbox", "layout_bbox", "dt_poly", "dt_polys"):
            if key not in node:
                continue
            bbox = self._coerce_ppstructure_bbox(node.get(key), image_width=image_width, image_height=image_height)
            if bbox is not None:
                return bbox
        return None

    def _coerce_ppstructure_bbox(
        self,
        value: Any,
        *,
        image_width: int,
        image_height: int,
    ) -> Optional[List[float]]:
        if value is None:
            return None
        try:
            raw = value.tolist() if hasattr(value, "tolist") else value
        except Exception:
            raw = value
        if isinstance(raw, list) and len(raw) == 1 and isinstance(raw[0], (list, tuple)):
            raw = raw[0]
        coords: List[float] = []
        if isinstance(raw, (list, tuple)) and len(raw) == 4 and not isinstance(raw[0], (list, tuple)):
            coords = [float(item) for item in raw]
        elif isinstance(raw, (list, tuple)):
            points = self._collect_ppstructure_points(raw)
            if points:
                xs = [point[0] for point in points]
                ys = [point[1] for point in points]
                coords = [min(xs), min(ys), max(xs), max(ys)]
        if len(coords) != 4:
            return None
        max_value = max(abs(value) for value in coords)
        if max_value <= 1.5:
            return self._normalize_bbox(coords)
        left, top, right, bottom = coords
        width = max(1.0, float(image_width or 1))
        height = max(1.0, float(image_height or 1))
        return [
            round(max(0.0, min(1.0, left / width)), 5),
            round(max(0.0, min(1.0, top / height)), 5),
            round(max(0.0, min(1.0, right / width)), 5),
            round(max(0.0, min(1.0, bottom / height)), 5),
        ]

    def _collect_ppstructure_points(self, value: Any) -> List[tuple[float, float]]:
        points: List[tuple[float, float]] = []

        def is_number(item: Any) -> bool:
            try:
                float(item)
            except (TypeError, ValueError):
                return False
            return True

        def collect(node: Any) -> None:
            try:
                raw_node = node.tolist() if hasattr(node, "tolist") else node
            except Exception:
                raw_node = node
            if not isinstance(raw_node, (list, tuple)):
                return
            if len(raw_node) >= 2 and is_number(raw_node[0]) and is_number(raw_node[1]):
                points.append((float(raw_node[0]), float(raw_node[1])))
                return
            for child in raw_node:
                collect(child)

        collect(value)
        return points

    def _normalize_table_rows(self, rows: Any) -> List[List[str]]:
        if not isinstance(rows, list):
            return []
        normalized: List[List[str]] = []
        for row in rows:
            if isinstance(row, list):
                values = [self.lowmem_pdf_service.normalize_text(str(cell)) for cell in row if str(cell).strip()]
                if values:
                    normalized.append(values)
        return normalized

    def _markdown_table_rows(self, text: str) -> List[List[str]]:
        rows: List[List[str]] = []
        for line in str(text or "").splitlines():
            stripped = line.strip()
            if "|" not in stripped:
                continue
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if cells and not all(re.fullmatch(r":?-{2,}:?", cell or "") for cell in cells):
                rows.append(cells)
        return rows

    def _html_table_rows(self, html_text: str) -> List[List[str]]:
        from html.parser import HTMLParser
        from html import unescape

        class TableParser(HTMLParser):
            def __init__(self) -> None:
                super().__init__()
                self.rows: List[List[str]] = []
                self._row: Optional[List[str]] = None
                self._cell_parts: Optional[List[str]] = None

            def handle_starttag(self, tag: str, attrs: List[tuple[str, Optional[str]]]) -> None:
                tag_name = tag.lower()
                if tag_name == "tr":
                    self._row = []
                elif tag_name in {"td", "th"} and self._row is not None:
                    self._cell_parts = []

            def handle_data(self, data: str) -> None:
                if self._cell_parts is not None:
                    self._cell_parts.append(data)

            def handle_endtag(self, tag: str) -> None:
                tag_name = tag.lower()
                if tag_name in {"td", "th"} and self._row is not None and self._cell_parts is not None:
                    value = self.lowmem_normalize("".join(self._cell_parts))
                    self._row.append(value)
                    self._cell_parts = None
                elif tag_name == "tr" and self._row is not None:
                    if any(cell for cell in self._row):
                        self.rows.append(self._row)
                    self._row = None

            @staticmethod
            def lowmem_normalize(value: str) -> str:
                return re.sub(r"\s+", " ", unescape(str(value or ""))).strip()

        parser = TableParser()
        try:
            parser.feed(str(html_text or ""))
            parser.close()
        except Exception:
            return []
        return parser.rows

    def _merge_glm_text_with_layout_blocks(
        self,
        page_result: LocalPdfPageResult,
        layout_blocks: List[Dict[str, Any]],
    ) -> LocalPdfPageResult:
        text_lines = [
            self.lowmem_pdf_service.normalize_text(line)
            for line in str(page_result.text or "").splitlines()
            if self.lowmem_pdf_service.normalize_text(line)
        ]
        text_blocks = [block for block in layout_blocks if block.get("type") != "table"]
        table_blocks = [block for block in layout_blocks if block.get("type") == "table"]
        if not text_lines and not table_blocks:
            return page_result

        assigned_blocks: List[Dict[str, Any]] = []
        if text_blocks:
            count_delta = abs(len(text_blocks) - len(text_lines))
            if count_delta <= max(2, int(max(len(text_blocks), len(text_lines)) * 0.35)):
                for index, line in enumerate(text_lines):
                    layout_block = text_blocks[index] if index < len(text_blocks) else {}
                    assigned_blocks.append(
                        {
                            **layout_block,
                            "type": layout_block.get("type") or "line",
                            "text": line,
                            "align": layout_block.get("align") or self._alignment_from_bbox(layout_block.get("bbox")),
                            "bbox": layout_block.get("bbox") or [0.08, 0.0, 0.92, 0.02],
                        }
                    )

        if not assigned_blocks and text_lines:
            return page_result

        blocks = [*assigned_blocks, *table_blocks]
        blocks = sorted(blocks, key=lambda item: ((item.get("bbox") or [0, 0, 0, 0])[1], (item.get("bbox") or [0, 0, 0, 0])[0]))
        lines = [
            {
                "text": block.get("text", ""),
                "bbox": block.get("bbox") or [0.08, 0.0, 0.92, 0.02],
                "confidence": page_result.confidence,
                "source": "glm_text_ppstructurev3_layout",
            }
            for block in blocks
            if block.get("type") != "table" and str(block.get("text") or "").strip()
        ]
        return LocalPdfPageResult(
            page_number=page_result.page_number,
            width=page_result.width,
            height=page_result.height,
            text=page_result.text,
            source=page_result.source,
            quality="medium" if page_result.quality in {"failed", "low"} else page_result.quality,
            confidence=page_result.confidence,
            lines=lines or page_result.lines,
            blocks=blocks,
            warnings=sorted(set([*(page_result.warnings or []), "ppstructurev3_layout_used"])),
            engine=page_result.engine,
            used_vl_fallback=page_result.used_vl_fallback,
        )

    def _merge_candidate_text_with_existing_layout(
        self,
        candidate: LocalPdfPageResult,
        layout_page: LocalPdfPageResult,
    ) -> LocalPdfPageResult:
        if candidate.source != "paddleocr_vl" or not candidate.text.strip():
            return candidate
        if not self._page_has_layout_signal(layout_page):
            return candidate
        merged = self._merge_text_with_existing_layout(
            candidate,
            layout_page,
            warning="paddleocr_vl_text_existing_layout",
            engine_suffix="layout_reuse",
        )
        if merged is candidate:
            return candidate
        return LocalPdfPageResult(
            page_number=merged.page_number,
            width=merged.width,
            height=merged.height,
            text=merged.text,
            source=merged.source,
            quality=merged.quality,
            confidence=max(merged.confidence, candidate.confidence),
            lines=merged.lines,
            blocks=merged.blocks,
            warnings=sorted(set(merged.warnings or [])),
            engine=merged.engine,
            used_vl_fallback=True,
        )

    def _try_write_recovered_docx_with_final_text(
        self,
        recovery_probe: Dict[str, Any],
        result_pages: List[LocalPdfPageResult],
        normalized_path: Path,
    ) -> bool:
        if not recovery_probe.get("available"):
            return False
        recovery_path = Path(str(recovery_probe.get("output_docx_path") or "")).expanduser()
        if not recovery_path.exists():
            return False
        target_texts = self._result_page_text_units(result_pages)
        if not target_texts:
            return False
        try:
            document = Document(str(recovery_path))
            containers = self._docx_text_containers(document)
            if not containers:
                return False
            ratio = len(target_texts) / max(len(containers), 1)
            if ratio < 0.35 or ratio > 3.0:
                return False
            for index, text in enumerate(target_texts):
                if index < len(containers):
                    self._replace_paragraph_text(containers[index], text)
                else:
                    document.add_paragraph(text)
            document.save(normalized_path)
            return normalized_path.exists() and normalized_path.stat().st_size > 0
        except Exception:
            logger.debug("Failed to rewrite PP-Structure recovery DOCX text.", exc_info=True)
            return False

    def _result_page_text_units(self, result_pages: List[LocalPdfPageResult]) -> List[str]:
        units: List[str] = []
        for page in result_pages:
            blocks = page.blocks or self._blocks_from_lines(page.lines, page.text)
            if blocks:
                for block in blocks:
                    block_type = str(block.get("type") or "").strip().lower()
                    if block_type == "table":
                        for row in block.get("rows") or []:
                            row_text = "\t".join(
                                self.lowmem_pdf_service.normalize_text(str(cell))
                                for cell in row
                                if str(cell).strip()
                            )
                            if row_text:
                                units.append(row_text)
                    elif block_type != "spacer":
                        text = self.lowmem_pdf_service.normalize_text(str(block.get("text") or ""))
                        if text:
                            units.append(text)
            else:
                for line in self._lines_from_text(page.text):
                    text = self.lowmem_pdf_service.normalize_text(str(line.get("text") or ""))
                    if text:
                        units.append(text)
        return units

    def _docx_text_containers(self, document: Any) -> List[Any]:
        containers: List[Any] = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                containers.append(paragraph)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = [paragraph for paragraph in cell.paragraphs if paragraph.text.strip()]
                    containers.extend(paragraphs)
        return containers

    @staticmethod
    def _replace_paragraph_text(paragraph: Any, text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    def _run_ppstructure_recovery_sidecar(self, source_path: Path, output_docx_path: Path) -> Dict[str, Any]:
        if str(settings.LOCAL_PDF_DOCX_RECOVERY_ENGINE or "").strip().lower() != "ppstructure_recovery_to_doc":
            return {"available": False, "disabled": True, "warning": ""}
        python_path = self._resolve_recovery_python()
        if not python_path:
            return {"available": False, "warning": "ppstructure_docx_recovery_missing", "error": "ocr_python_missing"}

        input_path = self._temp_json_file(output_docx_path.parent, "ppstructure-in-")
        output_path = self._temp_json_file(output_docx_path.parent, "ppstructure-out-")
        payload = {
            "file_path": str(source_path),
            "output_docx_path": str(output_docx_path),
            "timeout": max(30, int(settings.LOCAL_PDF_DOCX_RECOVERY_TIMEOUT)),
            "table_enabled": bool(settings.LOCAL_PDF_DOCX_RECOVERY_TABLE),
            "formula_enabled": bool(settings.LOCAL_PDF_DOCX_RECOVERY_FORMULA),
        }
        input_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
        ensure_private_file(input_path)
        env = os.environ.copy()
        existing_pythonpath = str(env.get("PYTHONPATH") or "").strip()
        env["PYTHONPATH"] = (
            str(SOURCE_BACKEND_ROOT)
            if not existing_pythonpath
            else f"{SOURCE_BACKEND_ROOT}{os.pathsep}{existing_pythonpath}"
        )
        try:
            result = subprocess.run(
                [
                    python_path,
                    "-m",
                    "app.workers.local_pdf_recovery_sidecar",
                    str(input_path),
                    str(output_path),
                ],
                cwd=str(SOURCE_BACKEND_ROOT),
                env=env,
                capture_output=True,
                text=True,
                timeout=max(30, int(settings.LOCAL_PDF_DOCX_RECOVERY_TIMEOUT)),
                check=False,
            )
            output_text = output_path.read_text(encoding="utf-8") if output_path.exists() else ""
            if output_text.strip():
                try:
                    parsed = json.loads(output_text)
                except Exception:
                    parsed = None
                if isinstance(parsed, dict):
                    return parsed
            if result.returncode != 0:
                return {
                    "available": False,
                    "warning": "ppstructure_docx_recovery_failed",
                    "error": (result.stderr or result.stdout or output_text)[-1000:],
                }
            return {
                "available": False,
                "warning": "ppstructure_docx_recovery_failed",
                "error": "ppstructure_sidecar_empty_or_invalid_json",
            }
        except Exception as exc:
            return {"available": False, "warning": "ppstructure_docx_recovery_failed", "error": str(exc)}
        finally:
            for path in (input_path, output_path):
                try:
                    path.unlink(missing_ok=True)
                except Exception:
                    pass

    def _resolve_ocr_python(self) -> Optional[str]:
        configured = str(settings.LOCAL_PDF_OCR_PYTHON or "").strip()
        candidates = [
            configured,
            str(SOURCE_BACKEND_ROOT / "ocr_venv" / "bin" / "python"),
            sys.executable,
        ]
        for candidate in candidates:
            if candidate and Path(candidate).exists():
                return candidate
        return None

    def _resolve_recovery_python(self) -> Optional[str]:
        configured = str(settings.LOCAL_PDF_DOCX_RECOVERY_PYTHON or "").strip()
        candidates = [
            configured,
            str(SOURCE_BACKEND_ROOT / "ppstructure_recovery_venv" / "bin" / "python"),
            self._resolve_ocr_python(),
        ]
        for candidate in candidates:
            if candidate and Path(candidate).exists():
                return candidate
        return None

    @staticmethod
    def _temp_json_file(directory: Path, prefix: str) -> Path:
        handle = tempfile.NamedTemporaryFile(
            mode="w",
            encoding="utf-8",
            prefix=prefix,
            suffix=".json",
            dir=str(directory),
            delete=False,
        )
        handle.close()
        return Path(handle.name)

    def _render_pdf_page_image(
        self,
        pdf_document: pdfium.PdfDocument,
        page_index: int,
        page: Dict[str, Any],
        *,
        dpi: int,
    ) -> Image.Image:
        pdf_page = pdf_document[page_index]
        scale = max(1.0, float(dpi) / 72.0)
        bitmap = pdf_page.render(scale=scale, rev_byteorder=True)
        try:
            return bitmap.to_pil().convert("RGB")
        finally:
            bitmap.close()
            pdf_page.close()

    def _enhance_for_ocr(self, image: Image.Image) -> Image.Image:
        try:
            import cv2
            import numpy as np

            rgb = np.array(image.convert("RGB"))
            gray = cv2.cvtColor(rgb, cv2.COLOR_RGB2GRAY)
            gray = self._crop_border(gray, cv2, np)
            denoised = cv2.fastNlMeansDenoising(gray, None, 7, 7, 21)
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(denoised)
            enhanced = self._deskew(enhanced, cv2, np)
            return Image.fromarray(enhanced).convert("RGB")
        except Exception:
            logger.debug("Local high-quality image enhancement failed; using original image.", exc_info=True)
            return ImageOps.exif_transpose(image.convert("RGB"))

    def _crop_border(self, gray: Any, cv2, np) -> Any:
        try:
            _, thresholded = cv2.threshold(gray, 245, 255, cv2.THRESH_BINARY_INV)
            contours, _ = cv2.findContours(thresholded, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            if not contours:
                return gray
            x, y, width, height = cv2.boundingRect(np.vstack(contours))
            page_height, page_width = gray.shape[:2]
            if width * height < page_width * page_height * 0.2:
                return gray
            pad = 8
            return gray[
                max(0, y - pad) : min(page_height, y + height + pad),
                max(0, x - pad) : min(page_width, x + width + pad),
            ]
        except Exception:
            return gray

    def _deskew(self, gray: Any, cv2, np) -> Any:
        try:
            binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
            coords = np.column_stack(np.where(binary > 0))
            if coords.size < 200:
                return gray
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = 90 + angle
            if abs(angle) < 0.3 or abs(angle) > 8:
                return gray
            height, width = gray.shape[:2]
            matrix = cv2.getRotationMatrix2D((width / 2, height / 2), angle, 1.0)
            return cv2.warpAffine(gray, matrix, (width, height), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
        except Exception:
            return gray

    @staticmethod
    def _image_to_jpeg_bytes(image: Image.Image) -> bytes:
        buffer = BytesIO()
        image.convert("RGB").save(buffer, format="JPEG", quality=92, optimize=True)
        return buffer.getvalue()

    def _render_layout_ir_docx(self, pages: List[LocalPdfPageResult], output_path: Path) -> None:
        document = Document()
        self._remove_default_paragraph(document)
        section = document.sections[0]
        if pages:
            section.page_width = Pt(max(360.0, min(936.0, float(pages[0].width or 595.0))))
            section.page_height = Pt(max(360.0, min(1224.0, float(pages[0].height or 842.0))))
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        for page_index, page in enumerate(pages):
            if page_index > 0:
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            self._append_page_to_docx(document, page)

        document.save(output_path)

    def _append_page_to_docx(self, document: Any, page: LocalPdfPageResult) -> None:
        blocks = page.blocks or self._blocks_from_lines(page.lines, page.text)
        if not blocks:
            for line in self._lines_from_text(page.text):
                self._append_text_block(document, {"type": "line", **line}, page=page, previous_bottom=None)
            return
        previous_bottom: Optional[float] = None
        for block in blocks:
            block_type = str(block.get("type") or "line").strip().lower()
            if block_type == "spacer":
                document.add_paragraph()
                continue
            if block_type == "table":
                rows = block.get("rows") or []
                if rows:
                    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
                    table.style = "Table Grid"
                    for row_index, row in enumerate(rows):
                        for col_index, cell_text in enumerate(row):
                            table.rows[row_index].cells[col_index].text = str(cell_text)
                continue
            self._append_text_block(document, block, page=page, previous_bottom=previous_bottom)
            bbox = self._normalize_bbox(block.get("bbox"))
            previous_bottom = max(previous_bottom or 0.0, bbox[3])

    def _append_text_block(
        self,
        document: Any,
        block: Dict[str, Any],
        *,
        page: Optional[LocalPdfPageResult] = None,
        previous_bottom: Optional[float] = None,
    ) -> None:
        text = self.lowmem_pdf_service.normalize_text(str(block.get("text") or ""))
        if not text:
            return
        paragraph = document.add_paragraph()
        bbox = self._normalize_bbox(block.get("bbox"))
        page_width = float(page.width if page else 595.0)
        page_height = float(page.height if page else 842.0)
        available_width = max(240.0, page_width - (0.7 + 0.7) * 72.0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        if isinstance(bbox, list) and len(bbox) == 4:
            try:
                left_indent = max(0.0, min(180.0, float(bbox[0]) * available_width))
                right_indent = max(0.0, min(180.0, (1.0 - float(bbox[2])) * available_width))
                paragraph.paragraph_format.left_indent = Pt(left_indent)
                paragraph.paragraph_format.right_indent = Pt(right_indent)
                top = float(bbox[1])
                bottom = float(bbox[3])
                gap = top if previous_bottom is None else max(0.0, top - previous_bottom)
                paragraph.paragraph_format.space_before = Pt(max(0.0, min(18.0, gap * page_height * 0.55)))
            except Exception:
                pass
        align = str(block.get("align") or "").strip().lower()
        block_type = str(block.get("type") or "").strip().lower()
        if align == "center" or block_type == "title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(text)
        bbox_height = max(0.0, float(bbox[3]) - float(bbox[1])) if isinstance(bbox, list) and len(bbox) == 4 else 0.0
        inferred_size = max(8.5, min(14.0, bbox_height * page_height * 0.72)) if bbox_height else 10.5
        run.font.size = Pt(12 if block_type == "title" else inferred_size)
        run.bold = block_type == "title"

    def _normalize_blocks(self, blocks: Any) -> List[Dict[str, Any]]:
        if not isinstance(blocks, list):
            return []
        normalized: List[Dict[str, Any]] = []
        for item in blocks:
            if not isinstance(item, dict):
                continue
            block_type = str(item.get("type") or "line").strip().lower()
            if block_type not in {"title", "paragraph", "line", "table", "spacer"}:
                block_type = "line"
            if block_type == "table":
                rows = item.get("rows") or []
                normalized_rows = [
                    [self.lowmem_pdf_service.normalize_text(str(cell)) for cell in row if str(cell).strip()]
                    for row in rows
                    if isinstance(row, list)
                ]
                normalized_rows = [row for row in normalized_rows if row]
                if normalized_rows:
                    normalized.append({"type": "table", "rows": normalized_rows, "bbox": self._normalize_bbox(item.get("bbox"))})
                continue
            if block_type == "spacer":
                normalized.append({"type": "spacer"})
                continue
            text = self.lowmem_pdf_service.normalize_text(str(item.get("text") or ""))
            if not text:
                continue
            align = str(item.get("align") or "left").strip().lower()
            if align not in {"left", "center", "right"}:
                align = "left"
            normalized.append(
                {
                    "type": block_type,
                    "text": text,
                    "align": align,
                    "bbox": self._normalize_bbox(item.get("bbox")),
                }
            )
        return normalized

    def _normalize_lines(self, lines: Any) -> List[Dict[str, Any]]:
        if not isinstance(lines, list):
            return []
        normalized: List[Dict[str, Any]] = []
        for item in lines:
            if not isinstance(item, dict):
                continue
            text = self.lowmem_pdf_service.normalize_text(str(item.get("text") or ""))
            if not text:
                continue
            line = {
                "text": text,
                "bbox": self._normalize_bbox(item.get("bbox")),
                "confidence": self._safe_float(item.get("confidence"), default=0.0),
                "source": str(item.get("source") or "ocr"),
            }
            normalized.append(line)
        return sorted(normalized, key=lambda item: ((item.get("bbox") or [0, 0, 0, 0])[1], (item.get("bbox") or [0, 0, 0, 0])[0]))

    def _normalize_bbox(self, bbox: Any) -> List[float]:
        if not isinstance(bbox, (list, tuple)) or len(bbox) != 4:
            return [0.08, 0.0, 0.92, 0.0]
        try:
            values = [float(value) for value in bbox]
        except Exception:
            return [0.08, 0.0, 0.92, 0.0]
        max_value = max(abs(value) for value in values)
        if max_value > 1.5:
            scale = 1000.0 if max_value > 10 else 100.0
            values = [value / scale for value in values]
        left, top, right, bottom = [max(0.0, min(1.0, value)) for value in values]
        if right <= left:
            right = min(1.0, left + 0.1)
        if bottom <= top:
            bottom = min(1.0, top + 0.02)
        return [round(left, 5), round(top, 5), round(right, 5), round(bottom, 5)]

    def _lines_from_text(self, text: str) -> List[Dict[str, Any]]:
        lines = []
        for index, line in enumerate(self.lowmem_pdf_service.normalize_text(text).splitlines()):
            stripped = line.strip()
            if not stripped:
                continue
            top = min(0.96, 0.07 + index * 0.035)
            lines.append(
                {
                    "text": stripped,
                    "bbox": [0.08, round(top, 5), 0.92, round(min(0.99, top + 0.025), 5)],
                    "confidence": 0.0,
                    "source": "ocr",
                }
            )
        return lines

    def _blocks_from_lines(self, lines: List[Dict[str, Any]], fallback_text: str) -> List[Dict[str, Any]]:
        if lines:
            return [
                {
                    "type": "line",
                    "text": line.get("text", ""),
                    "align": self._alignment_from_bbox(line.get("bbox")),
                    "bbox": line.get("bbox") or [0.08, 0.0, 0.92, 0.0],
                }
                for line in lines
                if str(line.get("text") or "").strip()
            ]
        return [{"type": "line", **line, "align": "left"} for line in self._lines_from_text(fallback_text)]

    @staticmethod
    def _alignment_from_bbox(bbox: Any) -> str:
        if not isinstance(bbox, list) or len(bbox) < 4:
            return "left"
        try:
            left, _top, right, _bottom = [float(value) for value in bbox[:4]]
        except Exception:
            return "left"
        width = right - left
        center = (left + right) / 2.0
        if width < 0.55 and abs(center - 0.5) < 0.08:
            return "center"
        if left > 0.58:
            return "right"
        return "left"

    def _blocks_to_text(self, blocks: List[Dict[str, Any]]) -> str:
        parts: List[str] = []
        for block in blocks:
            if block.get("type") == "table":
                for row in block.get("rows") or []:
                    parts.append("\t".join(str(cell).strip() for cell in row if str(cell).strip()))
            elif block.get("type") != "spacer":
                text = self.lowmem_pdf_service.normalize_text(str(block.get("text") or ""))
                if text:
                    parts.append(text)
        return "\n".join(parts).strip()

    def _should_use_vl_fallback(self, page_result: LocalPdfPageResult, classified_page: Dict[str, Any]) -> bool:
        if not settings.LOCAL_PDF_VL_FALLBACK_ENABLED:
            return False
        policy = str(getattr(settings, "LOCAL_PDF_VL_FALLBACK_POLICY", "quality_verify") or "quality_verify").strip().lower()
        if policy in {"off", "disabled", "none", "false"}:
            return False
        if policy == "always":
            return True
        if page_result.quality in {"low", "failed"}:
            return True
        if policy in {"quality", "quality_gate", "quality_verify", "verify", "aggressive"}:
            if page_result.quality != "high":
                return True
            if page_result.confidence < 0.88:
                return True
            if not self._page_has_layout_signal(page_result):
                return True
            if (
                "glm_ocr_non_json_response" in set(page_result.warnings or [])
                and (
                    page_result.quality != "high"
                    or page_result.confidence < 0.88
                    or not self._page_has_layout_signal(page_result)
                    or self._abnormal_char_ratio(page_result.text) >= 0.04
                    or "[UNREADABLE]" in str(page_result.text or "").upper()
                )
            ):
                return True
            if "[UNREADABLE]" in str(page_result.text or "").upper():
                return True
            if self._abnormal_char_ratio(page_result.text) >= 0.04:
                return True
        if "table" in {str(block.get("type") or "").lower() for block in page_result.blocks} and page_result.confidence < 0.82:
            return True
        if float(classified_page.get("image_area_ratio") or 0.0) >= 0.85 and page_result.confidence < 0.78:
            return True
        return False

    def _should_high_res_retry(self, page_result: LocalPdfPageResult) -> bool:
        if self._glm_high_res_dpi() <= self._glm_render_dpi():
            return False
        if page_result.quality in {"low", "failed"}:
            return True
        if page_result.confidence < 0.72:
            return True
        return False

    def _run_high_res_glm_retry(
        self,
        *,
        pdf_document: pdfium.PdfDocument,
        page_index: int,
        classified_page: Dict[str, Any],
        glm_client: GlmOcrOllamaClient,
        page_number: int,
        total_pages: int,
    ) -> Optional[LocalPdfPageResult]:
        try:
            high_res_image = self._render_pdf_page_image(
                pdf_document,
                page_index,
                classified_page,
                dpi=self._glm_high_res_dpi(),
            )
            try:
                enhanced = self._enhance_for_ocr(high_res_image)
                image_bytes = self._image_to_jpeg_bytes(enhanced)
            finally:
                if "enhanced" in locals():
                    del enhanced
            raw_retry = glm_client.recognize(image_bytes, page_number=page_number, total_pages=total_pages)
            retry_result = self._page_result_from_ocr_payload(
                raw_retry,
                classified_page=classified_page,
                source="glm_ocr",
                engine="glm_ocr_ollama",
            )
            retry_warnings = list(retry_result.warnings or [])
            retry_warnings.append("glm_high_res_retry")
            retry_result.warnings = retry_warnings
            return retry_result
        except Exception as exc:
            logger.warning("High-resolution GLM-OCR retry failed on page %s: %s", page_number, exc, exc_info=True)
            return None
        finally:
            if "high_res_image" in locals():
                del high_res_image
            gc.collect()

    @staticmethod
    def _glm_render_dpi() -> int:
        configured = int(getattr(settings, "LOCAL_PDF_GLM_RENDER_DPI", 0) or 0)
        if configured <= 0:
            configured = min(int(settings.LOCAL_PDF_RENDER_DPI or 160), 160)
        return max(96, min(configured, 180))

    @staticmethod
    def _glm_high_res_dpi() -> int:
        configured = int(getattr(settings, "LOCAL_PDF_GLM_HIGH_RES_DPI", 0) or 0)
        if configured <= 0:
            configured = max(LocalHighQualityPdfNormalizationService._glm_render_dpi(), 180)
        return max(96, min(configured, 200))

    def _is_better_page_result(self, current: LocalPdfPageResult, candidate: LocalPdfPageResult) -> bool:
        quality_rank = {"failed": 0, "low": 1, "medium": 2, "high": 3}
        current_score = self._page_result_score(current)
        candidate_score = self._page_result_score(candidate)
        if quality_rank.get(candidate.quality, 0) > quality_rank.get(current.quality, 0):
            return True
        if candidate_score >= current_score + 4.0:
            return True
        current_readable = self.lowmem_pdf_service.count_readable_chars(current.text)
        candidate_readable = self.lowmem_pdf_service.count_readable_chars(candidate.text)
        if candidate.quality == current.quality and candidate.confidence >= current.confidence + 0.05:
            return True
        if (
            candidate_readable >= max(12, int(current_readable * 1.12))
            and self._abnormal_char_ratio(candidate.text) <= self._abnormal_char_ratio(current.text) + 0.02
            and candidate_score >= current_score - 1.0
        ):
            return True
        if current.confidence <= 0.05 and candidate_readable >= current_readable and candidate_score >= current_score:
            return True
        return False

    def _page_result_score(self, page_result: LocalPdfPageResult) -> float:
        quality_rank = {"failed": 0, "low": 1, "medium": 2, "high": 3}
        readable_chars = self.lowmem_pdf_service.count_readable_chars(page_result.text)
        confidence = max(0.0, min(1.0, float(page_result.confidence or 0.0)))
        score = float(quality_rank.get(page_result.quality, 0)) * 20.0
        score += confidence * 18.0
        score += min(readable_chars / 500.0, 1.0) * 16.0
        if self._page_has_layout_signal(page_result):
            score += 4.0
        if any(str(block.get("type") or "").lower() == "table" for block in page_result.blocks):
            score += 2.0
        score -= self._abnormal_char_ratio(page_result.text) * 30.0
        score -= str(page_result.text or "").upper().count("[UNREADABLE]") * 1.5
        return score

    @staticmethod
    def _abnormal_char_ratio(text: str) -> float:
        value = str(text or "")
        if not value:
            return 1.0
        compact = re.sub(r"\s+", "", value)
        if not compact:
            return 1.0
        normal_count = 0
        for char in compact:
            if (
                "\u4e00" <= char <= "\u9fff"
                or char.isalnum()
                or char in "，。；：、！？（）()《》<>[]【】“”\"'‘’.-—_/%‰＋+＝=￥$@#&*·…"
            ):
                normal_count += 1
        return max(0.0, min(1.0, 1.0 - normal_count / max(len(compact), 1)))

    @staticmethod
    def _page_has_layout_signal(page_result: LocalPdfPageResult) -> bool:
        warnings = set(page_result.warnings or [])
        if (
            "glm_ocr_non_json_response" in warnings
            and "rapidocr_layout_assist_used" not in warnings
            and "ppstructurev3_layout_used" not in warnings
        ):
            return False
        return bool(page_result.blocks or any(line.get("bbox") for line in page_result.lines))

    @staticmethod
    def _avg_confidence(lines: List[Dict[str, Any]]) -> float:
        values = [float(line.get("confidence") or 0.0) for line in lines if isinstance(line.get("confidence"), (int, float))]
        return mean(values) if values else 0.0

    def _normalize_quality(self, raw_quality: str, *, text: str, confidence: float) -> str:
        normalized = str(raw_quality or "").strip().lower()
        if normalized in {"high", "medium", "low", "failed"}:
            if normalized == "high" and confidence and confidence < 0.75:
                return "medium"
            return normalized
        readable_chars = self.lowmem_pdf_service.count_readable_chars(text)
        if readable_chars <= 0:
            return "failed"
        if confidence >= 0.86:
            return "high"
        if confidence >= 0.65 or readable_chars >= 80:
            return "medium"
        return "low"

    @staticmethod
    def _safe_float(value: Any, *, default: float) -> float:
        try:
            return float(value)
        except Exception:
            return default

    def _remove_default_paragraph(self, document: Any) -> None:
        if document.paragraphs and not document.paragraphs[0].text:
            element = document.paragraphs[0]._element
            element.getparent().remove(element)

    def _build_metadata(
        self,
        *,
        source_path: str,
        normalized_path: str,
        classification: Dict[str, Any],
        warnings: Iterable[str],
        docx_recovery_engine: str,
        quality_counts: Dict[str, int],
        vl_fallback_pages: List[int],
        alignment_summary: Dict[str, Any],
        vl_fallback_attempted_pages: Optional[List[int]] = None,
        vl_fallback_failed_pages: Optional[List[int]] = None,
        vl_fallback_rejected_pages: Optional[List[int]] = None,
        ocr_candidate_scores: Optional[List[Dict[str, Any]]] = None,
        page_quality_metrics: Optional[List[Dict[str, Any]]] = None,
        vl_fallback_candidate_pages: Optional[List[int]] = None,
        vl_fallback_budget: Optional[int] = None,
        checkpoint_id: Optional[str] = None,
        checkpoint_resume_pages: Optional[List[int]] = None,
        frontline_page_timings: Optional[List[Dict[str, Any]]] = None,
        anchor_only_pages: Optional[List[int]] = None,
        glm_page_pages: Optional[List[int]] = None,
        glm_block_pages: Optional[List[int]] = None,
        quality_gate_rejected_pages: Optional[List[int]] = None,
    ) -> Dict[str, Any]:
        pages = classification.get("pages") or []
        native_pages = [int(page.get("page_number") or 0) for page in pages if page.get("kind") == "native_text"]
        image_pages = [int(page.get("page_number") or 0) for page in pages if page.get("kind") != "native_text"]
        warning_values = sorted(set(str(item).strip() for item in warnings if str(item).strip()))
        return {
            "format": "docx",
            "source_format": "pdf",
            "normalized_format": "docx",
            "file_path": normalized_path,
            "source_file_path": source_path,
            "normalized_file_path": normalized_path,
            "pages": len(pages),
            "pdf_kind": classification.get("pdf_kind") or "scanned",
            "pdf_frontline_engine": "local_high_quality_pdf_frontline",
            "pdf_frontline_profile": settings.LOCAL_PDF_FRONTLINE_PROFILE,
            "pdf_normalization_engine": "local_high_quality_pdf_frontline",
            "final_text_engine": "glm_ocr_quality_gate",
            "ocr_engine": "glm_ocr_ollama",
            "ocr_model": settings.LOCAL_PDF_GLM_MODEL,
            "glm_policy": settings.LOCAL_PDF_GLM_POLICY,
            "glm_scope": settings.LOCAL_PDF_GLM_SCOPE,
            "layout_engine": "ppstructurev3",
            "anchor_ocr_engine": "ppocrv5",
            "docx_recovery_engine": docx_recovery_engine,
            "vl_fallback_policy": settings.LOCAL_PDF_VL_FALLBACK_POLICY,
            "vl_fallback_budget": vl_fallback_budget
            if vl_fallback_budget is not None
            else self._vl_fallback_budget(total_pages=len(pages), image_page_count=len(image_pages)),
            "vl_fallback_candidate_pages": sorted(set(vl_fallback_candidate_pages or [])),
            "vl_fallback_engine": "paddleocr_vl_1_5",
            "vl_fallback_model": settings.LOCAL_PDF_VL_MODEL,
            "vl_fallback_pages": sorted(set(vl_fallback_pages)),
            "vl_fallback_used_pages": sorted(set(vl_fallback_pages)),
            "vl_fallback_attempted_pages": sorted(set(vl_fallback_attempted_pages or [])),
            "vl_fallback_failed_pages": sorted(set(vl_fallback_failed_pages or [])),
            "vl_fallback_rejected_pages": sorted(set(vl_fallback_rejected_pages or [])),
            "ocr_candidate_scores": list(ocr_candidate_scores or []),
            "ocr_quality_summary": quality_counts,
            "page_quality_metrics": list(page_quality_metrics or []),
            "layout_alignment_summary": alignment_summary,
            "checkpoint_id": checkpoint_id,
            "checkpoint_resume_pages": sorted(set(checkpoint_resume_pages or [])),
            "frontline_page_timings": list(frontline_page_timings or []),
            "anchor_only_pages": sorted(set(anchor_only_pages or [])),
            "glm_page_pages": sorted(set(glm_page_pages or [])),
            "glm_block_pages": sorted(set(glm_block_pages or [])),
            "quality_gate_rejected_pages": sorted(set(quality_gate_rejected_pages or [])),
            "native_text_pages": len(native_pages),
            "native_text_page_numbers": native_pages,
            "image_pages": image_pages,
            "image_page_count": len(image_pages),
            "warnings": warning_values,
        }

    def _default_output_dir(self) -> Path:
        return Path(settings.RUNTIME_ROOT) / "pdf_normalized"
