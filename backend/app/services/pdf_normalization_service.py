"""PDF-to-DOCX normalization for the high_quality_lowmem workflow."""

from __future__ import annotations

import gc
import logging
import re
import uuid
from dataclasses import dataclass
from pathlib import Path
from statistics import mean
from typing import Any, Dict, Iterable, List, Optional

import pdfplumber
import pypdfium2 as pdfium
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from app.core.config import settings
from app.core.runtime_security import ensure_private_directory, ensure_private_file
from app.services.lowmem_model_assets import ocr_model_paths, ocr_models_ready

logger = logging.getLogger(__name__)


class PdfNormalizationError(RuntimeError):
    """Error with a stable code and metadata payload for UI/reporting."""

    def __init__(self, code: str, message: str, *, metadata: Optional[Dict[str, Any]] = None) -> None:
        super().__init__(f"{code}: {message}")
        self.code = code
        self.metadata = metadata or {}


@dataclass
class ImageCandidate:
    name: str
    image: Any


@dataclass(frozen=True)
class OcrAccuracyProfile:
    name: str
    max_candidates: int
    low_quality_max_candidates: int
    local_retry_max_regions: int
    line_retry_upscale: float
    low_confidence_threshold: float
    min_avg_confidence: float
    target_avg_confidence: float
    high_quality_confidence: float
    quality_gate_min_confidence: float
    quality_gate_max_low_ratio: float
    quality_gate_max_abnormal_ratio: float
    page_risk_threshold: float
    high_res_retry_enabled: bool


CONTRACT_OCR_TERMS = [
    "合同",
    "协议",
    "甲方",
    "乙方",
    "丙方",
    "签订",
    "履行",
    "付款",
    "支付",
    "人民币",
    "金额",
    "违约",
    "责任",
    "争议",
    "仲裁",
    "法院",
    "有限公司",
    "法定代表人",
    "授权代表",
    "开户行",
    "账号",
    "期限",
    "生效",
    "解除",
    "终止",
    "约定",
    "义务",
    "权利",
    "赔偿",
    "损失",
    "管辖",
    "诉讼",
    "案件",
    "判决",
    "裁定",
    "执行",
    "申请",
    "被告",
    "原告",
    "上诉",
    "代理人",
    "证据",
    "事实",
    "理由",
    "签章",
    "盖章",
]


class ImageEnhancementService:
    """Generate OCR candidates without keeping page-scale arrays alive."""

    def __init__(self, *, max_candidates: int, low_quality_max_candidates: int) -> None:
        self.max_candidates = max(1, int(max_candidates or 1))
        self.low_quality_max_candidates = max(self.max_candidates, int(low_quality_max_candidates or self.max_candidates))

    def generate_candidates(
        self,
        image: Any,
        *,
        aggressive: bool = False,
        exclude_names: Optional[set[str]] = None,
    ) -> Iterable[ImageCandidate]:
        cv2, np = self._cv2_np()
        gray = self._to_gray(image, cv2)
        gray = self._crop_border(gray, cv2, np)
        exclude_names = exclude_names or set()
        limit = self.low_quality_max_candidates if aggressive else self.max_candidates
        emitted = 0

        def _emit(name: str, candidate: Any) -> Optional[ImageCandidate]:
            nonlocal emitted
            if emitted >= limit or name in exclude_names:
                return None
            emitted += 1
            return ImageCandidate(name, candidate)

        candidate = _emit("grayscale_crop", cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR))
        if candidate is not None:
            yield candidate

        if emitted < limit:
            enhanced = self._enhance_gray(gray, cv2, np)
            candidate = _emit("clahe_denoise_deskew", cv2.cvtColor(enhanced, cv2.COLOR_GRAY2BGR))
            if candidate is not None:
                yield candidate

        if emitted < limit:
            sharpened = self._unsharp(enhanced if "enhanced" in locals() else gray, cv2)
            candidate = _emit("clahe_unsharp", cv2.cvtColor(sharpened, cv2.COLOR_GRAY2BGR))
            if candidate is not None:
                yield candidate

        if aggressive and emitted < limit:
            thresholded = cv2.adaptiveThreshold(
                gray,
                255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY,
                35,
                11,
            )
            candidate = _emit("adaptive_threshold", cv2.cvtColor(thresholded, cv2.COLOR_GRAY2BGR))
            if candidate is not None:
                yield candidate

        if aggressive and emitted < limit:
            _, otsu = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
            morph = cv2.morphologyEx(otsu, cv2.MORPH_OPEN, kernel)
            candidate = _emit("otsu_morph_open", cv2.cvtColor(morph, cv2.COLOR_GRAY2BGR))
            if candidate is not None:
                yield candidate

    def _cv2_np(self):
        import cv2
        import numpy as np

        return cv2, np

    def _to_gray(self, image: Any, cv2) -> Any:
        if len(image.shape) == 2:
            return image
        if image.shape[2] == 4:
            return cv2.cvtColor(image, cv2.COLOR_RGBA2GRAY)
        return cv2.cvtColor(image, cv2.COLOR_RGB2GRAY)

    def _enhance_gray(self, gray: Any, cv2, np) -> Any:
        denoised = cv2.fastNlMeansDenoising(gray, None, 7, 7, 21)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(denoised)
        return self._deskew(enhanced, cv2, np)

    def _unsharp(self, gray: Any, cv2) -> Any:
        blurred = cv2.GaussianBlur(gray, (0, 0), 1.0)
        return cv2.addWeighted(gray, 1.45, blurred, -0.45, 0)

    def _deskew(self, gray: Any, cv2, np) -> Any:
        try:
            binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
            coords = np.column_stack(np.where(binary > 0))
            if coords.size < 200:
                return gray
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = 90 + angle
            if abs(angle) < 0.3 or abs(angle) > 8:
                return gray
            height, width = gray.shape[:2]
            matrix = cv2.getRotationMatrix2D((width / 2, height / 2), angle, 1.0)
            return cv2.warpAffine(
                gray,
                matrix,
                (width, height),
                flags=cv2.INTER_CUBIC,
                borderMode=cv2.BORDER_REPLICATE,
            )
        except Exception:
            logger.debug("PDF OCR deskew failed; using non-rotated candidate.", exc_info=True)
            return gray

    def _crop_border(self, gray: Any, cv2, np) -> Any:
        try:
            _, thresholded = cv2.threshold(gray, 245, 255, cv2.THRESH_BINARY_INV)
            contours, _ = cv2.findContours(thresholded, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            if not contours:
                return gray
            x, y, width, height = cv2.boundingRect(np.vstack(contours))
            page_height, page_width = gray.shape[:2]
            if width * height < page_width * page_height * 0.2:
                return gray
            pad = 8
            left = max(0, x - pad)
            top = max(0, y - pad)
            right = min(page_width, x + width + pad)
            bottom = min(page_height, y + height + pad)
            if right - left < page_width * 0.5 or bottom - top < page_height * 0.5:
                return gray
            return gray[top:bottom, left:right]
        except Exception:
            logger.debug("PDF OCR border crop failed; using original candidate.", exc_info=True)
            return gray


class RapidOCROnNXService:
    """Thin adapter around RapidOCR with explicit local model paths."""

    def __init__(self, *, enable_angle_cls: bool = False, max_side_len: Optional[int] = None) -> None:
        missing = self.missing_requirements(require_cls=bool(enable_angle_cls))
        if missing:
            raise PdfNormalizationError(
                "ocr_model_missing",
                "RapidOCR runtime or local PP-OCR model assets are missing.",
                metadata={"warnings": ["ocr_model_missing"], "missing_ocr_requirements": missing},
            )

        from rapidocr import EngineType, LangCls, LangDet, LangRec, ModelType, OCRVersion, RapidOCR

        paths = ocr_model_paths()
        resolved_max_side_len = int(max_side_len or settings.PDF_NORMALIZE_MAX_IMAGE_EDGE)
        params: Dict[str, Any] = {
            "Global.log_level": "warning",
            "Global.use_det": True,
            "Global.use_cls": bool(enable_angle_cls),
            "Global.use_rec": True,
            "Global.max_side_len": resolved_max_side_len,
            "Det.engine_type": EngineType.ONNXRUNTIME,
            "Det.lang_type": LangDet.CH,
            "Det.model_type": ModelType.MOBILE,
            "Det.ocr_version": OCRVersion.PPOCRV5,
            "Det.model_path": str(paths["det"]),
            "Rec.engine_type": EngineType.ONNXRUNTIME,
            "Rec.lang_type": LangRec.CH,
            "Rec.model_type": ModelType.MOBILE,
            "Rec.ocr_version": OCRVersion.PPOCRV5,
            "Rec.model_path": str(paths["rec"]),
            "Rec.rec_batch_num": max(1, int(settings.PDF_OCR_REC_BATCH_NUM or 1)),
            "EngineConfig.onnxruntime.intra_op_num_threads": 1,
            "EngineConfig.onnxruntime.inter_op_num_threads": 1,
            "EngineConfig.onnxruntime.enable_cpu_mem_arena": False,
        }
        if paths.get("cls"):
            params.update(
                {
                    "Cls.engine_type": EngineType.ONNXRUNTIME,
                    "Cls.lang_type": LangCls.CH,
                    "Cls.model_type": ModelType.MOBILE,
                    "Cls.ocr_version": OCRVersion.PPOCRV4,
                    "Cls.model_path": str(paths["cls"]),
                    "Cls.cls_batch_num": 1,
                }
            )
        self.enable_angle_cls = bool(enable_angle_cls)
        self.cls_available = bool(paths.get("cls"))
        self.max_side_len = resolved_max_side_len
        self.engine = RapidOCR(params=params)

    @staticmethod
    def missing_requirements(*, require_cls: bool = True) -> List[str]:
        missing: List[str] = []
        try:
            import rapidocr  # noqa: F401
            import onnxruntime  # noqa: F401
        except Exception:
            missing.append("rapidocr_onnxruntime_dependency")
        if not ocr_models_ready(require_cls=require_cls):
            paths = ocr_model_paths()
            if not paths.get("det"):
                missing.append("ppocrv5_mobile_det")
            if not paths.get("rec"):
                missing.append("ppocrv5_mobile_rec")
            if require_cls and not paths.get("cls"):
                missing.append("ppocr_angle_cls")
        return sorted(set(missing))

    def recognize(self, image: Any, *, use_cls: Optional[bool] = None) -> Dict[str, Any]:
        result = self.engine(
            image,
            use_cls=self.enable_angle_cls if use_cls is None else bool(use_cls),
            return_word_box=False,
            text_score=0.0,
        )
        return self._normalize_result(result, image)

    def _normalize_result(self, result: Any, image: Any) -> Dict[str, Any]:
        image_height, image_width = image.shape[:2]
        txts = list(getattr(result, "txts", []) or [])
        scores = list(getattr(result, "scores", []) or [])
        boxes = getattr(result, "boxes", None)
        elapse = getattr(result, "elapse", None)
        elapse_list = list(getattr(result, "elapse_list", []) or [])

        if not txts and isinstance(result, tuple) and result:
            raw_lines = result[0] or []
            txts = []
            scores = []
            boxes = []
            for raw_line in raw_lines:
                if not isinstance(raw_line, (list, tuple)) or len(raw_line) < 2:
                    continue
                boxes.append(raw_line[0])
                text_score = raw_line[1]
                if isinstance(text_score, (list, tuple)) and len(text_score) >= 2:
                    txts.append(str(text_score[0] or ""))
                    scores.append(float(text_score[1] or 0.0))

        lines: List[Dict[str, Any]] = []
        for index, text in enumerate(txts):
            normalized_text = PdfNormalizationService.normalize_text(str(text or ""))
            if not normalized_text:
                continue
            confidence = self._safe_float(scores[index] if index < len(scores) else None, default=0.0)
            bbox_px = self._box_to_rect(boxes[index] if boxes is not None and index < len(boxes) else None)
            line: Dict[str, Any] = {
                "text": normalized_text,
                "confidence": round(confidence, 4),
                "bbox": self._normalize_bbox(bbox_px, image_width, image_height),
                "bbox_px": bbox_px,
                "source": "ocr",
            }
            lines.append(line)

        return {
            "text": "\n".join(line["text"] for line in lines),
            "lines": lines,
            "confidence": round(mean([line["confidence"] for line in lines]) if lines else 0.0, 4),
            "elapse": elapse,
            "elapse_list": elapse_list,
        }

    def _box_to_rect(self, box: Any) -> List[float]:
        if box is None:
            return [0.0, 0.0, 1.0, 1.0]
        try:
            points = list(box.tolist() if hasattr(box, "tolist") else box)
            xs = [float(point[0]) for point in points]
            ys = [float(point[1]) for point in points]
            return [min(xs), min(ys), max(xs), max(ys)]
        except Exception:
            return [0.0, 0.0, 1.0, 1.0]

    def _normalize_bbox(self, bbox: List[float], width: int, height: int) -> List[float]:
        if width <= 0 or height <= 0:
            return [0.0, 0.0, 0.0, 0.0]
        return [
            round(max(0.0, min(1.0, bbox[0] / width)), 4),
            round(max(0.0, min(1.0, bbox[1] / height)), 4),
            round(max(0.0, min(1.0, bbox[2] / width)), 4),
            round(max(0.0, min(1.0, bbox[3] / height)), 4),
        ]

    def _safe_float(self, value: Any, *, default: float) -> float:
        try:
            return float(value)
        except (TypeError, ValueError):
            return default


class PdfNormalizationService:
    """Normalize PDFs into temporary DOCX files before desensitization."""

    def __init__(self) -> None:
        self.accuracy_profile = self._resolve_accuracy_profile()
        self.image_enhancer = ImageEnhancementService(
            max_candidates=self.accuracy_profile.max_candidates,
            low_quality_max_candidates=self.accuracy_profile.low_quality_max_candidates,
        )

    def _resolve_accuracy_profile(self) -> OcrAccuracyProfile:
        requested = str(getattr(settings, "PDF_OCR_ACCURACY_PROFILE", "wps_challenger") or "wps_challenger").strip()
        normalized = requested.lower()
        if normalized not in {"balanced", "wps_challenger"}:
            normalized = "wps_challenger"

        base = OcrAccuracyProfile(
            name=normalized,
            max_candidates=max(1, int(settings.PDF_OCR_MAX_CANDIDATES or 1)),
            low_quality_max_candidates=max(1, int(settings.PDF_OCR_LOW_QUALITY_MAX_CANDIDATES or 1)),
            local_retry_max_regions=max(0, int(settings.PDF_OCR_LOCAL_RETRY_MAX_REGIONS or 0)),
            line_retry_upscale=max(1.0, float(settings.PDF_OCR_LINE_RETRY_UPSCALE or 1.0)),
            low_confidence_threshold=float(settings.PDF_OCR_LOW_CONFIDENCE_THRESHOLD),
            min_avg_confidence=float(settings.PDF_OCR_MIN_AVG_CONFIDENCE),
            target_avg_confidence=float(settings.PDF_OCR_TARGET_AVG_CONFIDENCE),
            high_quality_confidence=0.86,
            quality_gate_min_confidence=float(settings.PDF_OCR_QUALITY_GATE_MIN_CONFIDENCE),
            quality_gate_max_low_ratio=float(settings.PDF_OCR_QUALITY_GATE_MAX_LOW_RATIO),
            quality_gate_max_abnormal_ratio=float(settings.PDF_OCR_QUALITY_GATE_MAX_ABNORMAL_RATIO),
            page_risk_threshold=float(settings.PDF_OCR_PAGE_RISK_THRESHOLD),
            high_res_retry_enabled=bool(settings.PDF_OCR_HIGH_RES_RETRY_ENABLED),
        )
        if normalized == "balanced":
            return base

        return OcrAccuracyProfile(
            name="wps_challenger",
            max_candidates=max(base.max_candidates, 3),
            low_quality_max_candidates=max(base.low_quality_max_candidates, 5),
            local_retry_max_regions=max(base.local_retry_max_regions, 12),
            line_retry_upscale=max(base.line_retry_upscale, 1.6),
            low_confidence_threshold=min(base.low_confidence_threshold, 0.65),
            min_avg_confidence=max(base.min_avg_confidence, 0.72),
            target_avg_confidence=max(base.target_avg_confidence, 0.86),
            high_quality_confidence=max(base.high_quality_confidence, 0.86),
            quality_gate_min_confidence=max(base.quality_gate_min_confidence, 0.78),
            quality_gate_max_low_ratio=min(base.quality_gate_max_low_ratio, 0.25),
            quality_gate_max_abnormal_ratio=min(base.quality_gate_max_abnormal_ratio, 0.10),
            page_risk_threshold=min(base.page_risk_threshold, 0.45),
            high_res_retry_enabled=base.high_res_retry_enabled,
        )

    def normalize_pdf(self, file_path: str, *, output_dir: Optional[str] = None) -> Dict[str, Any]:
        source_path = Path(file_path).expanduser().resolve()
        if source_path.suffix.lower() != ".pdf":
            raise ValueError(f"PdfNormalizationService only accepts PDF files: {source_path}")

        output_root = Path(output_dir).expanduser().resolve() if output_dir else self._default_output_dir()
        ensure_private_directory(output_root)
        normalized_path = output_root / f"{source_path.stem}_{uuid.uuid4().hex[:10]}_normalized.docx"

        classification = self.classify_pdf(str(source_path))
        pdf_kind = classification["pdf_kind"]
        warnings: List[str] = []

        if pdf_kind == "native_text":
            try:
                self._convert_text_pdf_with_pdf2docx(str(source_path), normalized_path)
                engine = "pdf2docx"
            except Exception as exc:
                logger.warning("pdf2docx failed; using native DOCX rebuild: %s", exc)
                warnings.append("pdf2docx_failed_fallback_used")
                self._rebuild_native_docx(classification["pages"], normalized_path)
                engine = "native_rebuild"
            metadata = self._build_base_metadata(
                source_path=str(source_path),
                normalized_path=str(normalized_path),
                classification=classification,
                engine=engine,
                warnings=warnings,
            )
            ensure_private_file(normalized_path)
            return {
                "normalized_file_path": str(normalized_path),
                "metadata": metadata,
                "structure": {"pages": self._public_pages(classification["pages"])},
            }

        result = self._normalize_with_ocr(str(source_path), normalized_path, classification, warnings=warnings)
        ensure_private_file(normalized_path)
        return result

    def classify_pdf(self, file_path: str) -> Dict[str, Any]:
        pages: List[Dict[str, Any]] = []
        with pdfplumber.open(file_path) as pdf:
            for page_index, page in enumerate(pdf.pages):
                native_text = self.normalize_text(page.extract_text() or "")
                compact = re.sub(r"\s+", "", native_text)
                readable_chars = self.count_readable_chars(compact)
                readable_ratio = readable_chars / max(len(compact), 1)
                lines = self._extract_native_lines(page, native_text)
                image_area_ratio = self._image_area_ratio(page)
                is_native = (
                    readable_chars >= int(settings.PDF_OCR_TEXT_THRESHOLD)
                    and readable_ratio >= float(settings.PDF_OCR_MIN_READABLE_RATIO)
                )
                pages.append(
                    {
                        "page_number": page_index + 1,
                        "width": float(getattr(page, "width", 595.0) or 595.0),
                        "height": float(getattr(page, "height", 842.0) or 842.0),
                        "text": native_text,
                        "char_count": readable_chars,
                        "readable_ratio": round(readable_ratio, 4),
                        "image_area_ratio": round(image_area_ratio, 4),
                        "kind": "native_text" if is_native else "scan_or_image",
                        "source": "native" if is_native else "scan_or_image",
                        "lines": lines,
                        "blocks": self._build_blocks_from_lines(lines, native_text),
                    }
                )

        native_pages = sum(1 for page in pages if page["kind"] == "native_text")
        image_pages = len(pages) - native_pages
        if not pages:
            pdf_kind = "scanned"
        elif native_pages == len(pages):
            pdf_kind = "native_text"
        elif native_pages > 0 and image_pages > 0:
            pdf_kind = "mixed"
        elif any(float(page.get("image_area_ratio") or 0.0) >= 0.5 for page in pages):
            pdf_kind = "image_pdf"
        else:
            pdf_kind = "scanned"

        return {
            "pdf_kind": pdf_kind,
            "pages": pages,
            "native_text_pages": native_pages,
            "image_pages": image_pages,
        }

    @staticmethod
    def normalize_text(text: str) -> str:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        normalized = normalized.replace("\u3000", " ")
        normalized = re.sub(r"[ \t]+\n", "\n", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()

    @staticmethod
    def count_readable_chars(text: str) -> int:
        return len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", text or ""))

    def _default_output_dir(self) -> Path:
        return Path(settings.RUNTIME_ROOT) / "pdf_normalized"

    def _convert_text_pdf_with_pdf2docx(self, source_path: str, output_path: Path) -> None:
        from pdf2docx import Converter

        converter = Converter(source_path)
        try:
            converter.convert(str(output_path), start=0, end=None)
        finally:
            converter.close()
        if not output_path.exists() or output_path.stat().st_size <= 0:
            raise RuntimeError("pdf2docx did not produce a DOCX file.")

    def _rebuild_native_docx(self, pages: List[Dict[str, Any]], output_path: Path) -> None:
        document = Document()
        self._remove_default_paragraph(document)
        for page_index, page in enumerate(pages):
            if page_index > 0:
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            self._append_page_lines(document, page.get("lines") or [], fallback_text=str(page.get("text") or ""))
        document.save(output_path)

    def _normalize_with_ocr(
        self,
        source_path: str,
        normalized_path: Path,
        classification: Dict[str, Any],
        *,
        warnings: List[str],
    ) -> Dict[str, Any]:
        missing = RapidOCROnNXService.missing_requirements(require_cls=bool(settings.PDF_OCR_ENABLE_ANGLE_CLS))
        if missing:
            metadata = self._build_base_metadata(
                source_path=source_path,
                normalized_path=str(normalized_path),
                classification=classification,
                engine="rapidocr_ppocrv5",
                warnings=[*warnings, "ocr_model_missing"],
            )
            metadata["missing_ocr_requirements"] = missing
            raise PdfNormalizationError(
                "ocr_model_missing",
                "RapidOCR PP-OCRv5 mobile det/rec local model assets are missing.",
                metadata=metadata,
            )

        ocr_service = RapidOCROnNXService(
            enable_angle_cls=bool(settings.PDF_OCR_ENABLE_ANGLE_CLS),
            max_side_len=max(
                int(settings.PDF_NORMALIZE_MAX_IMAGE_EDGE),
                int(settings.PDF_OCR_HIGH_RES_RETRY_MAX_IMAGE_EDGE),
            ),
        )
        document = Document()
        self._remove_default_paragraph(document)
        pdf_document = pdfium.PdfDocument(source_path)
        result_pages: List[Dict[str, Any]] = []
        quality_counts = {"high": 0, "medium": 0, "low": 0, "failed": 0}
        ocr_failed_pages: List[int] = []
        ocr_low_quality_pages: List[int] = []
        high_res_retry_pages: List[int] = []
        high_res_retry_used_pages: List[int] = []
        enhancement_warnings: List[str] = []
        low_confidence_lines: List[Dict[str, Any]] = []
        page_candidate_scores: List[Dict[str, Any]] = []
        page_risk_scores: List[Dict[str, Any]] = []
        quality_gate_failed_pages: List[int] = []
        local_retry_regions = 0

        try:
            for page_index, classified_page in enumerate(classification["pages"]):
                if page_index > 0:
                    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

                if classified_page["kind"] == "native_text":
                    native_page = dict(classified_page)
                    native_page["source"] = "native"
                    result_pages.append(native_page)
                    self._append_page_lines(
                        document,
                        native_page.get("lines") or [],
                        fallback_text=str(native_page.get("text") or ""),
                    )
                    continue

                try:
                    page_number = int(classified_page["page_number"])
                    rendered = self._render_pdf_page(pdf_document, page_index, classified_page)
                    page_result = self._ocr_page(rendered, ocr_service)
                    page_risk_score = self._page_risk_score(classified_page, page_result, rendered)
                    if self._should_high_res_retry(page_result, rendered, high_res_retry_pages, page_risk_score):
                        high_res_retry_pages.append(page_number)
                        retry_rendered = self._render_pdf_page(
                            pdf_document,
                            page_index,
                            classified_page,
                            max_edge=int(settings.PDF_OCR_HIGH_RES_RETRY_MAX_IMAGE_EDGE),
                            render_scale=float(settings.PDF_OCR_HIGH_RES_RETRY_SCALE),
                        )
                        try:
                            retry_result = self._ocr_page(retry_rendered, ocr_service)
                            retry_result["high_res_retry"] = True
                            retry_result["high_res_shape"] = list(retry_rendered.shape[:2])
                            selected_result, retry_used = self._select_better_ocr_result(page_result, retry_result)
                            page_result = selected_result
                            if retry_used:
                                high_res_retry_used_pages.append(page_number)
                        finally:
                            del retry_rendered
                            gc.collect()
                    page_risk_score = self._page_risk_score(classified_page, page_result, rendered)
                    quality_gate = self._page_quality_gate(page_result)
                    page_quality = str(page_result.get("quality") or "failed")
                    quality_counts[page_quality] = int(quality_counts.get(page_quality, 0)) + 1
                    if page_quality == "failed" or not str(page_result.get("text") or "").strip():
                        ocr_failed_pages.append(page_number)
                    elif page_quality == "low":
                        ocr_low_quality_pages.append(page_number)
                    if quality_gate == "failed":
                        quality_gate_failed_pages.append(page_number)
                    local_retry_regions += int(page_result.get("local_retry_regions") or 0)
                    page_candidate_scores.append(
                        {
                            "page_number": page_number,
                            "selected_candidate": page_result.get("candidate"),
                            "candidate_scores": page_result.get("candidate_scores") or [],
                        }
                    )
                    page_risk_scores.append(
                        {
                            "page_number": page_number,
                            "page_risk_score": page_risk_score,
                            "quality_gate": quality_gate,
                        }
                    )
                    page_low_confidence_lines = self._low_confidence_line_entries(page_number, page_result)
                    low_confidence_lines.extend(page_low_confidence_lines)
                    page_entry = {
                        "page_number": classified_page["page_number"],
                        "width": classified_page["width"],
                        "height": classified_page["height"],
                        "text": page_result.get("text") or "",
                        "source": "ocr" if page_result.get("text") else "ocr_empty",
                        "char_count": self.count_readable_chars(str(page_result.get("text") or "")),
                        "ocr_quality": page_quality,
                        "ocr_candidate": page_result.get("candidate"),
                        "ocr_candidate_scores": page_result.get("candidate_scores") or [],
                        "ocr_score": page_result.get("score"),
                        "ocr_confidence": page_result.get("confidence"),
                        "ocr_low_confidence_line_ratio": page_result.get("low_confidence_line_ratio"),
                        "ocr_low_confidence_lines": page_low_confidence_lines,
                        "ocr_abnormal_char_ratio": page_result.get("abnormal_char_ratio"),
                        "ocr_domain_term_hits": page_result.get("domain_term_hits"),
                        "ocr_local_retry_regions": int(page_result.get("local_retry_regions") or 0),
                        "ocr_high_res_retry": bool(page_result.get("high_res_retry")),
                        "ocr_quality_gate": quality_gate,
                        "page_risk_score": page_risk_score,
                        "lines": page_result.get("lines") or [],
                        "blocks": self._build_blocks_from_lines(page_result.get("lines") or [], str(page_result.get("text") or "")),
                    }
                    result_pages.append(page_entry)
                    self._append_page_lines(document, page_entry["lines"], fallback_text=page_entry["text"])
                except Exception as exc:
                    logger.warning("RapidOCR failed on PDF page %s: %s", page_index + 1, exc, exc_info=True)
                    ocr_failed_pages.append(int(classified_page["page_number"]))
                    quality_counts["failed"] += 1
                    result_pages.append(
                        {
                            "page_number": classified_page["page_number"],
                            "width": classified_page["width"],
                            "height": classified_page["height"],
                            "text": "",
                            "source": "ocr_failed",
                            "char_count": 0,
                            "ocr_quality": "failed",
                            "ocr_quality_gate": "failed",
                            "page_risk_score": 1.0,
                            "lines": [],
                            "blocks": [],
                            "warnings": ["ocr_failed"],
                        }
                    )
                    quality_gate_failed_pages.append(int(classified_page["page_number"]))
                finally:
                    if "rendered" in locals():
                        del rendered
                    gc.collect()
        finally:
            pdf_document.close()

        document.save(normalized_path)
        if ocr_failed_pages:
            warnings.append("ocr_failed_pages")
        if ocr_low_quality_pages:
            warnings.append("ocr_low_quality_pages")
        if high_res_retry_used_pages:
            warnings.append("ocr_high_res_retry_used")
        if quality_gate_failed_pages:
            warnings.append("ocr_quality_gate_failed")
        warnings.extend(enhancement_warnings)
        ocr_quality_gate = self._document_quality_gate(
            quality_counts=quality_counts,
            failed_pages=ocr_failed_pages,
            low_quality_pages=ocr_low_quality_pages,
            quality_gate_failed_pages=quality_gate_failed_pages,
            low_confidence_lines=low_confidence_lines,
        )

        metadata = self._build_base_metadata(
            source_path=source_path,
            normalized_path=str(normalized_path),
            classification=classification,
            engine="rapidocr_ppocrv5",
            warnings=warnings,
        )
        metadata.update(
            {
                "ocr_engine": "rapidocr_onnxruntime",
                "ocr_model": "ppocrv5_mobile",
                "ocr_pages": sum(1 for page in result_pages if page.get("source") == "ocr"),
                "ocr_failed_pages": ocr_failed_pages,
                "ocr_failed_page_count": len(ocr_failed_pages),
                "ocr_candidate_scores": page_candidate_scores,
                "ocr_low_confidence_lines": low_confidence_lines,
                "ocr_low_confidence_line_count": len(low_confidence_lines),
                "ocr_local_retry_regions": local_retry_regions,
                "ocr_high_res_retry_pages": high_res_retry_pages,
                "ocr_high_res_retry_used_pages": high_res_retry_used_pages,
                "ocr_quality_summary": quality_counts,
                "ocr_quality_gate": ocr_quality_gate,
                "ocr_quality_gate_failed_pages": quality_gate_failed_pages,
                "page_risk_scores": page_risk_scores,
            }
        )
        if ocr_low_quality_pages:
            metadata["ocr_low_quality_pages"] = ocr_low_quality_pages

        return {
            "normalized_file_path": str(normalized_path),
            "metadata": metadata,
            "structure": {"pages": self._public_pages(result_pages)},
        }

    def _ocr_page(self, image: Any, ocr_service: RapidOCROnNXService) -> Dict[str, Any]:
        candidate_results: List[Dict[str, Any]] = []
        seen_candidates: set[str] = set()
        for candidate in self.image_enhancer.generate_candidates(image, aggressive=False):
            try:
                seen_candidates.add(candidate.name)
                candidate_results.append(self._run_candidate_ocr(candidate, ocr_service))
            finally:
                del candidate
                gc.collect()

        best = self._best_candidate(candidate_results)
        if self._requires_aggressive_candidates(best):
            for candidate in self.image_enhancer.generate_candidates(
                image,
                aggressive=True,
                exclude_names=seen_candidates,
            ):
                try:
                    seen_candidates.add(candidate.name)
                    candidate_results.append(self._run_candidate_ocr(candidate, ocr_service))
                finally:
                    del candidate
                    gc.collect()

        if not candidate_results:
            return {"text": "", "lines": [], "quality": "failed", "score": 0.0}
        selected = self._best_candidate(candidate_results)
        selected["candidate_scores"] = self._candidate_score_summary(candidate_results)
        for line in selected.get("lines") or []:
            line.pop("bbox_px", None)
            line["quality"] = self._line_quality(float(line.get("confidence") or 0.0))
            line["candidate"] = selected.get("candidate")
            line.setdefault("line_error_score", round(self._line_error_score(str(line.get("text") or "")), 4))
        return selected

    def _run_candidate_ocr(self, candidate: ImageCandidate, ocr_service: RapidOCROnNXService) -> Dict[str, Any]:
        raw_result = ocr_service.recognize(candidate.image)
        raw_result["candidate"] = candidate.name
        raw_result = self._score_ocr_result(raw_result)
        if (
            self.accuracy_profile.name == "wps_challenger"
            and not bool(settings.PDF_OCR_ENABLE_ANGLE_CLS)
            and bool(getattr(ocr_service, "cls_available", False))
            and self._should_retry_with_angle_cls(raw_result)
        ):
            cls_retry = ocr_service.recognize(candidate.image, use_cls=True)
            cls_retry["candidate"] = f"{candidate.name}_angle_cls"
            cls_retry["angle_cls_retry"] = True
            cls_retry = self._score_ocr_result(cls_retry)
            if self._is_meaningfully_better_ocr_result(raw_result, cls_retry, min_gain=2.0):
                raw_result = cls_retry
        if raw_result.get("quality") in {"low", "failed"} or self._has_suspicious_lines(raw_result):
            raw_result = self._retry_low_confidence_regions(candidate.image, raw_result, ocr_service)
            raw_result = self._score_ocr_result(raw_result)
        return raw_result

    def _best_candidate(self, candidate_results: List[Dict[str, Any]]) -> Dict[str, Any]:
        if not candidate_results:
            return {"text": "", "lines": [], "quality": "failed", "score": 0.0}
        return max(candidate_results, key=lambda item: float(item.get("score") or 0.0))

    def _candidate_score_summary(self, candidate_results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        summary = []
        for result in candidate_results:
            summary.append(
                {
                    "candidate": result.get("candidate"),
                    "score": result.get("score"),
                    "quality": result.get("quality"),
                    "confidence": result.get("confidence"),
                    "readable_chars": result.get("readable_chars"),
                    "low_confidence_line_ratio": result.get("low_confidence_line_ratio"),
                    "low_confidence_line_count": result.get("low_confidence_line_count"),
                    "abnormal_char_ratio": result.get("abnormal_char_ratio"),
                    "line_error_penalty": result.get("line_error_penalty"),
                    "max_line_error_score": result.get("max_line_error_score"),
                    "domain_term_hits": result.get("domain_term_hits"),
                    "local_retry_regions": result.get("local_retry_regions"),
                }
            )
        return sorted(summary, key=lambda item: float(item.get("score") or 0.0), reverse=True)

    def _requires_aggressive_candidates(self, result: Dict[str, Any]) -> bool:
        if not result:
            return True
        if str(result.get("quality") or "") in {"low", "failed"}:
            return True
        if float(result.get("confidence") or 0.0) < self.accuracy_profile.target_avg_confidence:
            return True
        if float(result.get("low_confidence_line_ratio") or 0.0) > min(0.15, self.accuracy_profile.quality_gate_max_low_ratio):
            return True
        if float(result.get("abnormal_char_ratio") or 0.0) > min(0.08, self.accuracy_profile.quality_gate_max_abnormal_ratio):
            return True
        if float(result.get("max_line_error_score") or 0.0) >= 0.55:
            return True
        return False

    def _should_high_res_retry(
        self,
        result: Dict[str, Any],
        image: Any,
        retry_pages: List[int],
        page_risk_score: float = 0.0,
    ) -> bool:
        if not self.accuracy_profile.high_res_retry_enabled:
            return False
        if len(retry_pages) >= int(settings.PDF_OCR_HIGH_RES_RETRY_MAX_PAGES):
            return False
        try:
            current_max_edge = max(int(image.shape[0]), int(image.shape[1]))
        except Exception:
            current_max_edge = 0
        if current_max_edge >= int(settings.PDF_OCR_HIGH_RES_RETRY_MAX_IMAGE_EDGE) * 0.95:
            return False
        if self._requires_aggressive_candidates(result):
            return True
        return self.accuracy_profile.name == "wps_challenger" and page_risk_score >= self.accuracy_profile.page_risk_threshold

    def _select_better_ocr_result(self, primary: Dict[str, Any], retry: Dict[str, Any]) -> tuple[Dict[str, Any], bool]:
        min_gain = float(settings.PDF_OCR_HIGH_RES_RETRY_MIN_SCORE_GAIN)
        if self._is_meaningfully_better_ocr_result(primary, retry, min_gain=min_gain):
            retry["high_res_retry"] = True
            retry["retry_score_gain"] = round(float(retry.get("score") or 0.0) - float(primary.get("score") or 0.0), 4)
            return retry, True
        primary["high_res_retry"] = False
        primary["retry_score_gain"] = round(float(retry.get("score") or 0.0) - float(primary.get("score") or 0.0), 4)
        return primary, False

    def _is_meaningfully_better_ocr_result(
        self,
        primary: Dict[str, Any],
        retry: Dict[str, Any],
        *,
        min_gain: float,
    ) -> bool:
        primary_score = float(primary.get("score") or 0.0)
        retry_score = float(retry.get("score") or 0.0)
        if str(primary.get("quality") or "") in {"low", "failed"} and str(retry.get("quality") or "") in {"high", "medium"}:
            return True
        return retry_score >= primary_score + float(min_gain)

    def _should_retry_with_angle_cls(self, result: Dict[str, Any]) -> bool:
        if str(result.get("quality") or "") == "failed":
            return True
        if float(result.get("confidence") or 0.0) < max(0.55, self.accuracy_profile.min_avg_confidence - 0.1):
            return True
        if float(result.get("low_confidence_line_ratio") or 0.0) >= 0.45:
            return True
        return False

    def _score_ocr_result(self, result: Dict[str, Any]) -> Dict[str, Any]:
        lines = self._sort_and_merge_lines(
            [line for line in result.get("lines") or [] if str(line.get("text") or "").strip()]
        )
        text = "\n".join(str(line.get("text") or "").strip() for line in lines)
        confidences = [float(line.get("confidence") or 0.0) for line in lines]
        avg_confidence = mean(confidences) if confidences else 0.0
        low_ratio = (
            sum(1 for value in confidences if value < self.accuracy_profile.low_confidence_threshold) / len(confidences)
            if confidences
            else 1.0
        )
        readable_chars = self.count_readable_chars(text)
        abnormal_ratio = self._abnormal_char_ratio(text)
        clause_score = self._clause_sequence_score(text)
        domain_term_hits = self._domain_term_hits(text)
        line_error_scores = [self._line_error_score(str(line.get("text") or "")) for line in lines]
        for line, line_error_score in zip(lines, line_error_scores):
            line["line_error_score"] = round(line_error_score, 4)
        line_error_penalty = mean(line_error_scores) if line_error_scores else 1.0
        max_line_error_score = max(line_error_scores) if line_error_scores else 1.0
        score = (
            avg_confidence * 100.0
            + min(readable_chars, 1200) / 20.0
            + clause_score * 3.0
            + min(domain_term_hits, 12) * 1.2
            - low_ratio * 25.0
            - abnormal_ratio * 22.0
            - line_error_penalty * 12.0
        )
        quality = self._page_quality(avg_confidence, low_ratio, readable_chars)
        result.update(
            {
                "lines": lines,
                "text": text,
                "confidence": round(avg_confidence, 4),
                "low_confidence_line_count": sum(
                    1 for value in confidences if value < self.accuracy_profile.low_confidence_threshold
                ),
                "low_confidence_line_ratio": round(low_ratio, 4),
                "abnormal_char_ratio": round(abnormal_ratio, 4),
                "line_error_penalty": round(line_error_penalty, 4),
                "max_line_error_score": round(max_line_error_score, 4),
                "domain_term_hits": domain_term_hits,
                "readable_chars": readable_chars,
                "score": round(score, 4),
                "quality": quality,
            }
        )
        return result

    def _sort_and_merge_lines(self, lines: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        if not lines:
            return []

        def _sort_key(line: Dict[str, Any]) -> tuple[float, float]:
            bbox = line.get("bbox") or [0.0, 0.0, 0.0, 0.0]
            try:
                return (float(bbox[1]), float(bbox[0]))
            except Exception:
                return (0.0, 0.0)

        sorted_lines = sorted((dict(line) for line in lines), key=_sort_key)
        merged: List[Dict[str, Any]] = []
        for line in sorted_lines:
            if not merged or not self._same_visual_line(merged[-1], line):
                merged.append(line)
                continue
            previous = merged[-1]
            previous_text = str(previous.get("text") or "").strip()
            current_text = str(line.get("text") or "").strip()
            separator = "" if self._mostly_cjk(previous_text + current_text) else " "
            previous["text"] = f"{previous_text}{separator}{current_text}".strip()
            previous["confidence"] = round(
                mean([float(previous.get("confidence") or 0.0), float(line.get("confidence") or 0.0)]),
                4,
            )
            previous["bbox"] = self._merge_bbox(previous.get("bbox"), line.get("bbox"))
            previous["bbox_px"] = self._merge_bbox(previous.get("bbox_px"), line.get("bbox_px"))
        return merged

    def _same_visual_line(self, previous: Dict[str, Any], current: Dict[str, Any]) -> bool:
        previous_bbox = previous.get("bbox") or []
        current_bbox = current.get("bbox") or []
        if len(previous_bbox) < 4 or len(current_bbox) < 4:
            return False
        try:
            prev_top, prev_bottom = float(previous_bbox[1]), float(previous_bbox[3])
            cur_top, cur_bottom = float(current_bbox[1]), float(current_bbox[3])
            prev_left, prev_right = float(previous_bbox[0]), float(previous_bbox[2])
            cur_left, cur_right = float(current_bbox[0]), float(current_bbox[2])
        except Exception:
            return False
        prev_height = max(prev_bottom - prev_top, 0.001)
        cur_height = max(cur_bottom - cur_top, 0.001)
        center_delta = abs(((prev_top + prev_bottom) / 2.0) - ((cur_top + cur_bottom) / 2.0))
        horizontal_gap = cur_left - prev_right
        if center_delta > max(prev_height, cur_height) * 0.6:
            return False
        if horizontal_gap < -0.01 or horizontal_gap > 0.18:
            return False
        previous_text = str(previous.get("text") or "").strip()
        current_text = str(current.get("text") or "").strip()
        if len(previous_text) > 28 and len(current_text) > 28:
            return False
        if prev_left > cur_left or prev_right > cur_right:
            return False
        return True

    def _merge_bbox(self, first: Any, second: Any) -> List[float]:
        if not isinstance(first, list) or not isinstance(second, list) or len(first) < 4 or len(second) < 4:
            return first if isinstance(first, list) else [0.0, 0.0, 0.0, 0.0]
        try:
            return [
                min(float(first[0]), float(second[0])),
                min(float(first[1]), float(second[1])),
                max(float(first[2]), float(second[2])),
                max(float(first[3]), float(second[3])),
            ]
        except Exception:
            return first

    def _mostly_cjk(self, text: str) -> bool:
        compact = re.sub(r"\s+", "", text or "")
        if not compact:
            return True
        cjk_count = len(re.findall(r"[\u4e00-\u9fff]", compact))
        return cjk_count / max(len(compact), 1) >= 0.45

    def _has_suspicious_lines(self, result: Dict[str, Any]) -> bool:
        for line in result.get("lines") or []:
            if self._line_error_score(str(line.get("text") or "")) >= 0.55:
                return True
        return False

    def _domain_term_hits(self, text: str) -> int:
        return sum(1 for term in CONTRACT_OCR_TERMS if term in (text or ""))

    def _line_error_score(self, text: str) -> float:
        compact = re.sub(r"\s+", "", text or "")
        if not compact:
            return 1.0
        abnormal = self._abnormal_char_ratio(compact)
        repeated = 0.0
        if re.search(r"(.)(\1{4,})", compact):
            repeated = 0.35
        replacement = 0.35 if "�" in compact or "□" in compact or "▯" in compact else 0.0
        symbol_run = 0.25 if re.search(r"[~`|·•]{3,}", compact) else 0.0
        broken_cjk = 0.2 if re.search(r"[\u4e00-\u9fff][A-Za-z]{1,2}[\u4e00-\u9fff]", compact) else 0.0
        bracket_penalty = 0.15 if self._has_unbalanced_brackets(compact) else 0.0
        missing_clause_ordinal = 0.45 if re.search(r"第[条款项章节]", compact) else 0.0
        broken_contract_amount = 0.2 if "合同额" in compact and "合同金额" not in compact else 0.0
        broken_party_phrase = 0.25 if re.search(r"双(应|按|均|同意|协商|履行|签订|确认)", compact) else 0.0
        missing_party_suffix = 0.6 if re.search(r"^[甲乙丙丁][：:]", compact) else 0.0
        return min(
            1.0,
            abnormal
            + repeated
            + replacement
            + symbol_run
            + broken_cjk
            + bracket_penalty
            + missing_clause_ordinal
            + broken_contract_amount
            + broken_party_phrase
            + missing_party_suffix,
        )

    def _has_unbalanced_brackets(self, text: str) -> bool:
        pairs = [("(", ")"), ("（", "）"), ("[", "]"), ("【", "】"), ("《", "》")]
        return any(abs(text.count(left) - text.count(right)) >= 2 for left, right in pairs)

    def _retry_low_confidence_regions(
        self,
        image: Any,
        result: Dict[str, Any],
        ocr_service: RapidOCROnNXService,
    ) -> Dict[str, Any]:
        lines = list(result.get("lines") or [])
        retry_budget = self.accuracy_profile.local_retry_max_regions
        if not lines or retry_budget <= 0:
            return result

        cv2, _np = self.image_enhancer._cv2_np()
        try:
            image_height, image_width = image.shape[:2]
        except Exception:
            return result
        retried = 0
        for line in lines:
            if retried >= retry_budget:
                break
            confidence = float(line.get("confidence") or 0.0)
            line_error_score = self._line_error_score(str(line.get("text") or ""))
            if confidence >= self.accuracy_profile.low_confidence_threshold and line_error_score < 0.35:
                continue
            bbox = line.get("bbox_px") or [0, 0, image_width, image_height]
            x0, y0, x1, y1 = [int(round(float(value))) for value in bbox]
            pad = 10
            left = max(0, x0 - pad)
            top = max(0, y0 - pad)
            right = min(image_width, x1 + pad)
            bottom = min(image_height, y1 + pad)
            if right - left < 12 or bottom - top < 12:
                continue
            crop = image[top:bottom, left:right]
            try:
                crop = self._upscale_line_crop(crop, cv2)
                retry = ocr_service.recognize(crop, use_cls=bool(settings.PDF_OCR_ENABLE_ANGLE_CLS))
                retry_lines = retry.get("lines") or []
                retry_text = self.normalize_text("".join(str(item.get("text") or "") for item in retry_lines))
                retry_conf = mean([float(item.get("confidence") or 0.0) for item in retry_lines]) if retry_lines else 0.0
                retry_error_score = self._line_error_score(retry_text)
                if retry_text and (
                    retry_conf > confidence + 0.05
                    or (retry_conf >= confidence and retry_error_score + 0.2 < line_error_score)
                ):
                    line["text"] = retry_text
                    line["confidence"] = round(retry_conf, 4)
                    line["was_local_retry"] = True
                    line["retry_confidence_delta"] = round(retry_conf - confidence, 4)
                    line["line_error_score"] = round(retry_error_score, 4)
                retried += 1
            finally:
                del crop
                gc.collect()
        result["lines"] = lines
        result["local_retry_regions"] = retried
        return result

    def _upscale_line_crop(self, crop: Any, cv2) -> Any:
        try:
            upscale = self.accuracy_profile.line_retry_upscale
            if upscale <= 1.0:
                return crop
            height, width = crop.shape[:2]
            if height >= 72:
                return crop
            return cv2.resize(
                crop,
                (max(width + 1, int(width * upscale)), max(height + 1, int(height * upscale))),
                interpolation=cv2.INTER_CUBIC,
            )
        except Exception:
            return crop

    def _render_pdf_page(
        self,
        pdf_document: pdfium.PdfDocument,
        page_index: int,
        page: Dict[str, Any],
        *,
        max_edge: Optional[int] = None,
        render_scale: Optional[float] = None,
    ) -> Any:
        import numpy as np

        pdf_page = pdf_document[page_index]
        width = float(page.get("width") or 595.0)
        height = float(page.get("height") or 842.0)
        resolved_max_edge = float(max_edge or settings.PDF_NORMALIZE_MAX_IMAGE_EDGE)
        resolved_render_scale = float(render_scale or settings.PDF_OCR_RENDER_SCALE)
        scale = max(1.0, min(resolved_render_scale, resolved_max_edge / max(width, height, 1.0)))
        bitmap = pdf_page.render(scale=scale, rev_byteorder=True)
        try:
            pil_image = bitmap.to_pil().convert("RGB")
            return np.array(pil_image)
        finally:
            bitmap.close()
            pdf_page.close()

    def _page_quality(self, avg_confidence: float, low_ratio: float, readable_chars: int) -> str:
        if readable_chars <= 0:
            return "failed"
        if avg_confidence >= self.accuracy_profile.high_quality_confidence and low_ratio <= 0.12:
            return "high"
        if avg_confidence >= self.accuracy_profile.min_avg_confidence:
            return "medium"
        return "low"

    def _line_quality(self, confidence: float) -> str:
        if confidence >= self.accuracy_profile.high_quality_confidence:
            return "high"
        if confidence >= self.accuracy_profile.low_confidence_threshold:
            return "medium"
        return "low"

    def _page_quality_gate(self, result: Dict[str, Any]) -> str:
        if str(result.get("quality") or "") in {"failed", "low"}:
            return "failed"
        if float(result.get("confidence") or 0.0) < self.accuracy_profile.quality_gate_min_confidence:
            return "review_required"
        if float(result.get("low_confidence_line_ratio") or 0.0) > self.accuracy_profile.quality_gate_max_low_ratio:
            return "review_required"
        if float(result.get("abnormal_char_ratio") or 0.0) > self.accuracy_profile.quality_gate_max_abnormal_ratio:
            return "review_required"
        if float(result.get("max_line_error_score") or 0.0) >= 0.55:
            return "review_required"
        if str(result.get("quality") or "") == "medium":
            return "review_required"
        return "passed"

    def _document_quality_gate(
        self,
        *,
        quality_counts: Dict[str, int],
        failed_pages: List[int],
        low_quality_pages: List[int],
        quality_gate_failed_pages: List[int],
        low_confidence_lines: List[Dict[str, Any]],
    ) -> str:
        if failed_pages or low_quality_pages or quality_gate_failed_pages:
            return "failed"
        if int(quality_counts.get("medium") or 0) > 0 or low_confidence_lines:
            return "review_required"
        return "passed"

    def _low_confidence_line_entries(self, page_number: int, result: Dict[str, Any]) -> List[Dict[str, Any]]:
        entries: List[Dict[str, Any]] = []
        for line_index, line in enumerate(result.get("lines") or [], start=1):
            confidence = float(line.get("confidence") or 0.0)
            line_error_score = float(line.get("line_error_score") or self._line_error_score(str(line.get("text") or "")))
            if confidence >= self.accuracy_profile.low_confidence_threshold and line_error_score < 0.55:
                continue
            entries.append(
                {
                    "page_number": page_number,
                    "line_index": line_index,
                    "text": str(line.get("text") or ""),
                    "confidence": round(confidence, 4),
                    "line_error_score": round(line_error_score, 4),
                    "bbox": line.get("bbox") or [0.0, 0.0, 0.0, 0.0],
                }
            )
        return entries

    def _page_risk_score(self, classified_page: Dict[str, Any], result: Dict[str, Any], image: Any = None) -> float:
        risk = 0.0
        if float(classified_page.get("image_area_ratio") or 0.0) >= 0.75:
            risk += 0.12
        if int(classified_page.get("char_count") or 0) <= 0 and classified_page.get("kind") != "native_text":
            risk += 0.08
        if float(result.get("confidence") or 0.0) < self.accuracy_profile.quality_gate_min_confidence:
            risk += 0.28
        if float(result.get("low_confidence_line_ratio") or 0.0) > self.accuracy_profile.quality_gate_max_low_ratio:
            risk += 0.22
        if float(result.get("abnormal_char_ratio") or 0.0) > self.accuracy_profile.quality_gate_max_abnormal_ratio:
            risk += 0.18
        if float(result.get("line_error_penalty") or 0.0) > 0.22:
            risk += 0.12
        if float(result.get("max_line_error_score") or 0.0) >= 0.55:
            risk += 0.22
        if int(result.get("readable_chars") or 0) < 20:
            risk += 0.12
        try:
            height, width = image.shape[:2]
            if min(height, width) < 900:
                risk += 0.08
        except Exception:
            pass
        return round(min(1.0, risk), 4)

    def _abnormal_char_ratio(self, text: str) -> float:
        compact = re.sub(r"\s+", "", text or "")
        if not compact:
            return 1.0
        normal = re.findall(r"[\u4e00-\u9fffA-Za-z0-9，。；：、,.()（）【】\[\]《》<>/\\+=_%￥$#&*\"'“”‘’!?！？-]", compact)
        return max(0.0, 1.0 - len(normal) / max(len(compact), 1))

    def _clause_sequence_score(self, text: str) -> float:
        matches = re.findall(r"(?:第[一二三四五六七八九十百]+条|\\b\\d+(?:\\.\\d+)*[、.])", text or "")
        if not matches:
            return 0.0
        return min(float(len(matches)), 10.0) / 10.0

    def _extract_native_lines(self, page: Any, native_text: str) -> List[Dict[str, Any]]:
        try:
            raw_lines = page.extract_text_lines(strip=False, return_chars=True, layout=True)
        except TypeError:
            try:
                raw_lines = page.extract_text_lines(strip=False, return_chars=True)
            except Exception:
                raw_lines = []
        except Exception:
            raw_lines = []

        if not isinstance(raw_lines, list) or not raw_lines:
            return self._lines_from_text(native_text)

        page_width = float(getattr(page, "width", 595.0) or 595.0)
        page_height = float(getattr(page, "height", 842.0) or 842.0)
        lines: List[Dict[str, Any]] = []
        for raw_line in raw_lines:
            if not isinstance(raw_line, dict):
                continue
            text = self.normalize_text(str(raw_line.get("text") or ""))
            if not text:
                continue
            x0 = float(raw_line.get("x0", 0.0) or 0.0)
            x1 = float(raw_line.get("x1", page_width) or page_width)
            top = float(raw_line.get("top", 0.0) or 0.0)
            bottom = float(raw_line.get("bottom", top + 12.0) or (top + 12.0))
            confidence = 1.0
            lines.append(
                {
                    "text": text,
                    "bbox": [
                        round(max(0.0, min(1.0, x0 / page_width)), 4),
                        round(max(0.0, min(1.0, top / page_height)), 4),
                        round(max(0.0, min(1.0, x1 / page_width)), 4),
                        round(max(0.0, min(1.0, bottom / page_height)), 4),
                    ],
                    "confidence": confidence,
                    "source": "native",
                    "quality": "high",
                }
            )
        return lines or self._lines_from_text(native_text)

    def _lines_from_text(self, text: str) -> List[Dict[str, Any]]:
        lines: List[Dict[str, Any]] = []
        for index, line in enumerate(self.normalize_text(text).splitlines()):
            stripped = line.strip()
            if not stripped:
                continue
            top = min(0.95, 0.06 + index * 0.03)
            lines.append(
                {
                    "text": stripped,
                    "bbox": [0.08, round(top, 4), 0.92, round(min(0.98, top + 0.025), 4)],
                    "confidence": 1.0,
                    "source": "native",
                    "quality": "high",
                }
            )
        return lines

    def _image_area_ratio(self, page: Any) -> float:
        page_area = float(getattr(page, "width", 0.0) or 0.0) * float(getattr(page, "height", 0.0) or 0.0)
        if page_area <= 0:
            return 0.0
        total = 0.0
        for image in getattr(page, "images", []) or []:
            try:
                width = abs(float(image.get("x1", 0.0) or 0.0) - float(image.get("x0", 0.0) or 0.0))
                height = abs(float(image.get("bottom", 0.0) or 0.0) - float(image.get("top", 0.0) or 0.0))
                total += max(0.0, width * height)
            except Exception:
                continue
        return min(1.0, total / page_area)

    def _build_blocks_from_lines(self, lines: List[Dict[str, Any]], fallback_text: str) -> List[Dict[str, Any]]:
        if lines:
            return [{"type": "line", "text": line.get("text", ""), "align": "left"} for line in lines if line.get("text")]
        text = self.normalize_text(fallback_text)
        return [{"type": "line", "text": line, "align": "left"} for line in text.splitlines() if line.strip()]

    def _append_page_lines(self, document: Document, lines: List[Dict[str, Any]], *, fallback_text: str) -> None:
        emitted = False
        for line in lines:
            text = self.normalize_text(str(line.get("text") or ""))
            if not text:
                continue
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(10.5)
            emitted = True
        if emitted:
            return
        for raw_line in self.normalize_text(fallback_text).splitlines():
            text = raw_line.strip()
            if text:
                document.add_paragraph(text)

    def _remove_default_paragraph(self, document: Document) -> None:
        if document.paragraphs and not document.paragraphs[0].text:
            element = document.paragraphs[0]._element
            element.getparent().remove(element)

    def _build_base_metadata(
        self,
        *,
        source_path: str,
        normalized_path: str,
        classification: Dict[str, Any],
        engine: str,
        warnings: List[str],
    ) -> Dict[str, Any]:
        pages = classification.get("pages") or []
        image_pages = [int(page.get("page_number") or 0) for page in pages if page.get("kind") != "native_text"]
        native_pages = [int(page.get("page_number") or 0) for page in pages if page.get("kind") == "native_text"]
        return {
            "format": "docx",
            "source_format": "pdf",
            "normalized_format": "docx",
            "file_path": normalized_path,
            "source_file_path": source_path,
            "normalized_file_path": normalized_path,
            "pages": len(pages),
            "pdf_kind": classification.get("pdf_kind") or "scanned",
            "pdf_normalization_engine": engine,
            "ocr_accuracy_profile": self.accuracy_profile.name,
            "native_text_pages": len(native_pages),
            "native_text_page_numbers": native_pages,
            "image_pages": image_pages,
            "image_page_count": len(image_pages),
            "ocr_pages": 0,
            "ocr_failed_pages": [],
            "ocr_low_confidence_lines": [],
            "ocr_low_confidence_line_count": 0,
            "ocr_local_retry_regions": 0,
            "ocr_quality_summary": {"high": 0, "medium": 0, "low": 0, "failed": 0},
            "ocr_quality_gate": "not_applicable" if not image_pages else "pending",
            "ocr_quality_gate_failed_pages": [],
            "warnings": sorted(set(item for item in warnings if item)),
        }

    def _public_pages(self, pages: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        public_pages: List[Dict[str, Any]] = []
        for page in pages:
            public_page = dict(page)
            public_lines = []
            for line in page.get("lines") or []:
                public_line = dict(line)
                public_line.pop("bbox_px", None)
                public_lines.append(public_line)
            public_page["lines"] = public_lines
            public_pages.append(public_page)
        return public_pages
